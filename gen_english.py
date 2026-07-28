"""
六年级英语（PEP人教版）单词表 + 词组句 + 整句翻译 生成器
输出：3个Word文档（单词表、词组句题目+答案、翻译题题目+答案）
依赖：pip install python-docx
"""

import random
import time
import csv
import os
from collections import defaultdict

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("请先安装 python-docx：pip install python-docx")
    exit(1)

# 使用时间戳作为随机种子，每次运行结果不同
random.seed(int(time.time()))

# ==================== 单词数据（按主题分类） ====================

VOCAB_DATA = {
    "Unit 1 My Day": [
        ("morning", "/ˈmɔːnɪŋ/", "n.", "早晨，上午"),
        ("afternoon", "/ˌɑːftəˈnuːn/", "n.", "下午"),
        ("evening", "/ˈiːvnɪŋ/", "n.", "傍晚，晚上"),
        ("night", "/naɪt/", "n.", "夜晚"),
        ("breakfast", "/ˈbrekfəst/", "n.", "早餐"),
        ("lunch", "/lʌntʃ/", "n.", "午餐"),
        ("dinner", "/ˈdɪnə(r)/", "n.", "晚餐"),
        ("get up", "/ɡet ʌp/", "v.", "起床"),
        ("go to school", "/ɡəʊ tə skuːl/", "v.", "去上学"),
        ("go to bed", "/ɡəʊ tə bed/", "v.", "去睡觉"),
        ("have class", "/hæv klɑːs/", "v.", "上课"),
        ("do homework", "/duː ˈhəʊmwɜːk/", "v.", "做作业"),
        ("watch TV", "/wɒtʃ ˌtiːˈviː/", "v.", "看电视"),
        ("read a book", "/riːd ə bʊk/", "v.", "看书"),
        ("play sports", "/pleɪ spɔːts/", "v.", "做运动"),
        ("eat breakfast", "/iːt ˈbrekfəst/", "v.", "吃早餐"),
        ("eat lunch", "/iːt lʌntʃ/", "v.", "吃午餐"),
        ("eat dinner", "/iːt ˈdɪnə(r)/", "v.", "吃晚餐"),
        ("usually", "/ˈjuːʒuəli/", "adv.", "通常"),
        ("often", "/ˈɒfn/", "adv.", "经常"),
        ("sometimes", "/ˈsʌmtaɪmz/", "adv.", "有时"),
        ("always", "/ˈɔːlweɪz/", "adv.", "总是"),
        ("early", "/ˈɜːli/", "adv./adj.", "早的（地）"),
        ("late", "/leɪt/", "adv./adj.", "晚的（地）"),
        ("o'clock", "/əˈklɒk/", "adv.", "…点钟"),
        ("minute", "/ˈmɪnɪt/", "n.", "分钟"),
        ("hour", "/ˈaʊə(r)/", "n.", "小时"),
        ("day", "/deɪ/", "n.", "天，日子"),
        ("week", "/wiːk/", "n.", "周，星期"),
        ("weekend", "/ˌwiːkˈend/", "n.", "周末"),
    ],
    "Unit 2 Last Weekend": [
        ("yesterday", "/ˈjestədeɪ/", "adv./n.", "昨天"),
        ("last", "/lɑːst/", "adj.", "上一个的，刚过去的"),
        ("weekend", "/ˌwiːkˈend/", "n.", "周末"),
        ("cleaned", "/kliːnd/", "v.", "打扫（过去式）"),
        ("cooked", "/kʊkt/", "v.", "烹饪（过去式）"),
        ("washed", "/wɒʃt/", "v.", "洗（过去式）"),
        ("watched", "/wɒtʃt/", "v.", "观看（过去式）"),
        ("played", "/pleɪd/", "v.", "玩耍（过去式）"),
        ("visited", "/ˈvɪzɪtɪd/", "v.", "拜访（过去式）"),
        ("read", "/red/", "v.", "读（过去式）"),
        ("saw", "/sɔː/", "v.", "看见（过去式）"),
        ("went", "/went/", "v.", "去（过去式）"),
        ("had", "/hæd/", "v.", "有，吃（过去式）"),
        ("was", "/wɒz/", "v.", "是（过去式）"),
        ("did", "/dɪd/", "v.", "做（过去式）"),
        ("stayed", "/steɪd/", "v.", "停留（过去式）"),
        ("studied", "/ˈstʌdid/", "v.", "学习（过去式）"),
        ("fished", "/fɪʃt/", "v.", "钓鱼（过去式）"),
        ("swam", "/swæm/", "v.", "游泳（过去式）"),
        ("climbed", "/klaɪmd/", "v.", "攀爬（过去式）"),
        ("rowed", "/rəʊd/", "v.", "划（过去式）"),
        ("boat", "/bəʊt/", "n.", "小船"),
        ("park", "/pɑːk/", "n.", "公园"),
        ("museum", "/mjuˈziːəm/", "n.", "博物馆"),
        ("cinema", "/ˈsɪnəmə/", "n.", "电影院"),
        ("library", "/ˈlaɪbrəri/", "n.", "图书馆"),
        ("hospital", "/ˈhɒspɪtl/", "n.", "医院"),
        ("supermarket", "/ˈsuːpəmɑːkɪt/", "n.", "超市"),
        ("bookstore", "/ˈbʊkstɔː(r)/", "n.", "书店"),
        ("happy", "/ˈhæpi/", "adj.", "开心的"),
    ],
    "Unit 3 Where Did You Go?": [
        ("where", "/weə(r)/", "adv.", "在哪里"),
        ("went", "/went/", "v.", "去（过去式）"),
        ("camp", "/kæmp/", "n./v.", "野营；野营"),
        ("rode", "/rəʊd/", "v.", "骑（过去式）"),
        ("hurt", "/hɜːt/", "v.", "受伤（过去式）"),
        ("ate", "/eɪt/", "v.", "吃（过去式）"),
        ("took", "/tʊk/", "v.", "拍照（过去式）"),
        ("bought", "/bɔːt/", "v.", "买（过去式）"),
        ("gift", "/ɡɪft/", "n.", "礼物"),
        ("fell", "/fel/", "v.", "摔倒（过去式）"),
        ("off", "/ɒf/", "prep.", "从…掉下"),
        ("licked", "/lɪkt/", "v.", "舔（过去式）"),
        ("beach", "/biːtʃ/", "n.", "海滩"),
        ("basket", "/ˈbɑːskɪt/", "n.", "篮子"),
        ("part", "/pɑːt/", "n.", "部分"),
        ("labour", "/ˈleɪbə(r)/", "n.", "劳动"),
        ("holiday", "/ˈhɒlədeɪ/", "n.", "假日"),
        ("travelled", "/ˈtrævld/", "v.", "旅行（过去式）"),
        ("mule", "/mjuːl/", "n.", "骡子"),
        ("Turpan", "/ˈtʊərpæn/", "n.", "吐鲁番"),
        ("could", "/kʊd/", "v.", "能（过去式）"),
        ("till", "/tɪl/", "prep./conj.", "直到"),
        ("favourite", "/ˈfeɪvərɪt/", "adj.", "最喜欢的"),
        ("look", "/lʊk/", "v.", "看"),
        ("like", "/laɪk/", "v./prep.", "像；喜欢"),
    ],
    "Unit 4 Then and Now": [
        ("then", "/ðen/", "adv.", "那时"),
        ("now", "/naʊ/", "adv.", "现在"),
        ("ago", "/əˈɡəʊ/", "adv.", "…以前"),
        ("before", "/bɪˈfɔː(r)/", "adv./prep.", "以前；在…之前"),
        ("there was", "/ðeə wɒz/", "v.", "有（过去式）"),
        ("there were", "/ðeə wɜː(r)/", "v.", "有（过去式复数）"),
        ("wasn't", "/ˈwɒznt/", "v.", "不是（过去式否定）"),
        ("weren't", "/wɜːnt/", "v.", "不是（过去式否定复数）"),
        ("building", "/ˈbɪldɪŋ/", "n.", "建筑物"),
        ("house", "/haʊs/", "n.", "房屋"),
        ("village", "/ˈvɪlɪdʒ/", "n.", "村庄"),
        ("city", "/ˈsɪti/", "n.", "城市"),
        ("town", "/taʊn/", "n.", "小镇"),
        ("old", "/əʊld/", "adj.", "旧的，老的"),
        ("new", "/njuː/", "adj.", "新的"),
        ("young", "/jʌŋ/", "adj.", "年轻的"),
        ("tall", "/tɔːl/", "adj.", "高的"),
        ("short", "/ʃɔːt/", "adj.", "矮的，短的"),
        ("big", "/bɪɡ/", "adj.", "大的"),
        ("small", "/smɔːl/", "adj.", "小的"),
        ("different", "/ˈdɪfrənt/", "adj.", "不同的"),
        ("same", "/seɪm/", "adj.", "相同的"),
        ("change", "/tʃeɪndʒ/", "v./n.", "改变"),
        ("dining hall", "/ˈdaɪnɪŋ hɔːl/", "n.", "餐厅"),
        ("gym", "/dʒɪm/", "n.", "体育馆"),
        ("grass", "/ɡrɑːs/", "n.", "草坪"),
        ("year", "/jɪə(r)/", "n.", "年"),
        ("month", "/mʌnθ/", "n.", "月份"),
        ("day", "/deɪ/", "n.", "天"),
    ],
    "Unit 5 What Does He Do?": [
        ("doctor", "/ˈdɒktə(r)/", "n.", "医生"),
        ("nurse", "/nɜːs/", "n.", "护士"),
        ("teacher", "/ˈtiːtʃə(r)/", "n.", "教师"),
        ("student", "/ˈstjuːdnt/", "n.", "学生"),
        ("police officer", "/pəˈliːs ˈɒfɪsə(r)/", "n.", "警察"),
        ("driver", "/ˈdraɪvə(r)/", "n.", "司机"),
        ("farmer", "/ˈfɑːmə(r)/", "n.", "农民"),
        ("worker", "/ˈwɜːkə(r)/", "n.", "工人"),
        ("singer", "/ˈsɪŋə(r)/", "n.", "歌手"),
        ("writer", "/ˈraɪtə(r)/", "n.", "作家"),
        ("actor", "/ˈæktə(r)/", "n.", "男演员"),
        ("actress", "/ˈæktrəs/", "n.", "女演员"),
        ("artist", "/ˈɑːtɪst/", "n.", "画家，艺术家"),
        ("engineer", "/ˌendʒɪˈnɪə(r)/", "n.", "工程师"),
        ("accountant", "/əˈkaʊntənt/", "n.", "会计"),
        ("pilot", "/ˈpaɪlət/", "n.", "飞行员"),
        ("coach", "/kəʊtʃ/", "n.", "教练"),
        ("scientist", "/ˈsaɪəntɪst/", "n.", "科学家"),
        ("postman", "/ˈpəʊstmən/", "n.", "邮递员"),
        ("businessman", "/ˈbɪznəsmæn/", "n.", "商人"),
        ("fisherman", "/ˈfɪʃəmən/", "n.", "渔民"),
        ("factory", "/ˈfæktri/", "n.", "工厂"),
        ("hospital", "/ˈhɒspɪtl/", "n.", "医院"),
        ("school", "/skuːl/", "n.", "学校"),
        ("company", "/ˈkʌmpəni/", "n.", "公司"),
        ("university", "/ˌjuːnɪˈvɜːsəti/", "n.", "大学"),
        ("gym", "/dʒɪm/", "n.", "体育馆"),
        ("office", "/ˈɒfɪs/", "n.", "办公室"),
        ("restaurant", "/ˈrestrɒnt/", "n.", "餐馆"),
        ("job", "/dʒɒb/", "n.", "工作"),
    ],
    "Unit 6 How Do You Feel?": [
        ("happy", "/ˈhæpi/", "adj.", "高兴的"),
        ("sad", "/sæd/", "adj.", "伤心的"),
        ("angry", "/ˈæŋɡri/", "adj.", "生气的"),
        ("afraid", "/əˈfreɪd/", "adj.", "害怕的"),
        ("worried", "/ˈwʌrid/", "adj.", "担心的"),
        ("tired", "/ˈtaɪəd/", "adj.", "疲倦的"),
        ("hungry", "/ˈhʌŋɡri/", "adj.", "饥饿的"),
        ("thirsty", "/ˈθɜːsti/", "adj.", "口渴的"),
        ("cold", "/kəʊld/", "adj.", "冷的"),
        ("hot", "/hɒt/", "adj.", "热的"),
        ("warm", "/wɔːm/", "adj.", "温暖的"),
        ("cool", "/kuːl/", "adj.", "凉爽的"),
        ("ill", "/ɪl/", "adj.", "生病的"),
        ("sick", "/sɪk/", "adj.", "不舒服的"),
        ("bored", "/bɔːd/", "adj.", "无聊的"),
        ("excited", "/ɪkˈsaɪtɪd/", "adj.", "兴奋的"),
        ("surprised", "/səˈpraɪzd/", "adj.", "惊讶的"),
        ("scared", "/skeəd/", "adj.", "恐惧的"),
        ("proud", "/praʊd/", "adj.", "骄傲的"),
        ("shy", "/ʃaɪ/", "adj.", "害羞的"),
        ("kind", "/kaɪnd/", "adj.", "友善的"),
        ("friendly", "/ˈfrendli/", "adj.", "友好的"),
        ("polite", "/pəˈlaɪt/", "adj.", "有礼貌的"),
        ("helpful", "/ˈhelpfl/", "adj.", "乐于助人的"),
        ("clever", "/ˈklevə(r)/", "adj.", "聪明的"),
        ("hard-working", "/ˌhɑːd ˈwɜːkɪŋ/", "adj.", "勤劳的"),
        ("funny", "/ˈfʌni/", "adj.", "滑稽的，有趣的"),
        ("strict", "/strɪkt/", "adj.", "严格的"),
        ("quiet", "/ˈkwaɪət/", "adj.", "安静的"),
        ("active", "/ˈæktɪv/", "adj.", "活跃的"),
    ],
    "Food & Drink": [
        ("rice", "/raɪs/", "n.", "米饭"),
        ("noodles", "/ˈnuːdlz/", "n.", "面条"),
        ("bread", "/bred/", "n.", "面包"),
        ("egg", "/eɡ/", "n.", "鸡蛋"),
        ("milk", "/mɪlk/", "n.", "牛奶"),
        ("juice", "/dʒuːs/", "n.", "果汁"),
        ("water", "/ˈwɔːtə(r)/", "n.", "水"),
        ("tea", "/tiː/", "n.", "茶"),
        ("coffee", "/ˈkɒfi/", "n.", "咖啡"),
        ("cake", "/keɪk/", "n.", "蛋糕"),
        ("apple", "/ˈæpl/", "n.", "苹果"),
        ("banana", "/bəˈnɑːnə/", "n.", "香蕉"),
        ("orange", "/ˈɒrɪndʒ/", "n.", "橙子"),
        ("grape", "/ɡreɪp/", "n.", "葡萄"),
        ("watermelon", "/ˈwɔːtəmelən/", "n.", "西瓜"),
        ("strawberry", "/ˈstrɔːbəri/", "n.", "草莓"),
        ("pear", "/peə(r)/", "n.", "梨"),
        ("peach", "/piːtʃ/", "n.", "桃子"),
        ("chicken", "/ˈtʃɪkɪn/", "n.", "鸡肉"),
        ("fish", "/fɪʃ/", "n.", "鱼肉"),
        ("beef", "/biːf/", "n.", "牛肉"),
        ("pork", "/pɔːk/", "n.", "猪肉"),
        ("vegetable", "/ˈvedʒtəbl/", "n.", "蔬菜"),
        ("tomato", "/təˈmɑːtəʊ/", "n.", "西红柿"),
        ("potato", "/pəˈteɪtəʊ/", "n.", "土豆"),
        ("carrot", "/ˈkærət/", "n.", "胡萝卜"),
        ("onion", "/ˈʌnjən/", "n.", "洋葱"),
        ("cabbage", "/ˈkæbɪdʒ/", "n.", "卷心菜"),
        ("ice cream", "/aɪs kriːm/", "n.", "冰淇淋"),
        ("candy", "/ˈkændi/", "n.", "糖果"),
        ("chocolate", "/ˈtʃɒklət/", "n.", "巧克力"),
        ("sandwich", "/ˈsænwɪtʃ/", "n.", "三明治"),
        ("hamburger", "/ˈhæmbɜːɡə(r)/", "n.", "汉堡包"),
        ("salad", "/ˈsæləd/", "n.", "沙拉"),
        ("soup", "/suːp/", "n.", "汤"),
        ("healthy", "/ˈhelθi/", "adj.", "健康的"),
        ("delicious", "/dɪˈlɪʃəs/", "adj.", "美味的"),
        ("fresh", "/freʃ/", "adj.", "新鲜的"),
        ("sweet", "/swiːt/", "adj.", "甜的"),
        ("sour", "/ˈsaʊə(r)/", "adj.", "酸的"),
    ],
    "Animals & Nature": [
        ("cat", "/kæt/", "n.", "猫"),
        ("dog", "/dɒɡ/", "n.", "狗"),
        ("bird", "/bɜːd/", "n.", "鸟"),
        ("fish", "/fɪʃ/", "n.", "鱼"),
        ("rabbit", "/ˈræbɪt/", "n.", "兔子"),
        ("monkey", "/ˈmʌŋki/", "n.", "猴子"),
        ("elephant", "/ˈelɪfənt/", "n.", "大象"),
        ("tiger", "/ˈtaɪɡə(r)/", "n.", "老虎"),
        ("lion", "/ˈlaɪən/", "n.", "狮子"),
        ("panda", "/ˈpændə/", "n.", "熊猫"),
        ("bear", "/beə(r)/", "n.", "熊"),
        ("horse", "/hɔːs/", "n.", "马"),
        ("sheep", "/ʃiːp/", "n.", "绵羊"),
        ("cow", "/kaʊ/", "n.", "奶牛"),
        ("pig", "/pɪɡ/", "n.", "猪"),
        ("duck", "/dʌk/", "n.", "鸭子"),
        ("hen", "/hen/", "n.", "母鸡"),
        ("goat", "/ɡəʊt/", "n.", "山羊"),
        ("mouse", "/maʊs/", "n.", "老鼠"),
        ("snake", "/sneɪk/", "n.", "蛇"),
        ("frog", "/frɒɡ/", "n.", "青蛙"),
        ("insect", "/ˈɪnsekt/", "n.", "昆虫"),
        ("butterfly", "/ˈbʌtəflaɪ/", "n.", "蝴蝶"),
        ("bee", "/biː/", "n.", "蜜蜂"),
        ("ant", "/ænt/", "n.", "蚂蚁"),
        ("tree", "/triː/", "n.", "树"),
        ("flower", "/ˈflaʊə(r)/", "n.", "花"),
        ("grass", "/ɡrɑːs/", "n.", "草"),
        ("leaf", "/liːf/", "n.", "叶子"),
        ("river", "/ˈrɪvə(r)/", "n.", "河流"),
        ("lake", "/leɪk/", "n.", "湖"),
        ("mountain", "/ˈmaʊntən/", "n.", "山"),
        ("forest", "/ˈfɒrɪst/", "n.", "森林"),
        ("sky", "/skaɪ/", "n.", "天空"),
        ("cloud", "/klaʊd/", "n.", "云"),
        ("rain", "/reɪn/", "n.", "雨"),
        ("snow", "/snəʊ/", "n.", "雪"),
        ("wind", "/wɪnd/", "n.", "风"),
        ("sun", "/sʌn/", "n.", "太阳"),
        ("moon", "/muːn/", "n.", "月亮"),
        ("star", "/stɑː(r)/", "n.", "星星"),
    ],
    "Weather & Seasons": [
        ("spring", "/sprɪŋ/", "n.", "春天"),
        ("summer", "/ˈsʌmə(r)/", "n.", "夏天"),
        ("autumn", "/ˈɔːtəm/", "n.", "秋天"),
        ("winter", "/ˈwɪntə(r)/", "n.", "冬天"),
        ("sunny", "/ˈsʌni/", "adj.", "晴朗的"),
        ("cloudy", "/ˈklaʊdi/", "adj.", "多云的"),
        ("rainy", "/ˈreɪni/", "adj.", "下雨的"),
        ("snowy", "/ˈsnəʊi/", "adj.", "下雪的"),
        ("windy", "/ˈwɪndi/", "adj.", "有风的"),
        ("foggy", "/ˈfɒɡi/", "adj.", "有雾的"),
        ("weather", "/ˈweðə(r)/", "n.", "天气"),
        ("temperature", "/ˈtemprətʃə(r)/", "n.", "温度"),
        ("degree", "/dɪˈɡriː/", "n.", "度"),
        ("season", "/ˈsiːzn/", "n.", "季节"),
        ("swim", "/swɪm/", "v.", "游泳"),
        ("skate", "/skeɪt/", "v.", "滑冰"),
        ("ski", "/skiː/", "v.", "滑雪"),
        ("fly a kite", "/flaɪ ə kaɪt/", "v.", "放风筝"),
        ("plant trees", "/plɑːnt triːz/", "v.", "种树"),
        ("pick apples", "/pɪk ˈæplz/", "v.", "摘苹果"),
        ("make a snowman", "/meɪk ə ˈsnəʊmæn/", "v.", "堆雪人"),
        ("go on a picnic", "/ɡəʊ ɒn ə ˈpɪknɪk/", "v.", "去野餐"),
    ],
    "Clothing & Shopping": [
        ("shirt", "/ʃɜːt/", "n.", "衬衫"),
        ("T-shirt", "/ˈtiː ʃɜːt/", "n.", "T恤衫"),
        ("dress", "/dres/", "n.", "连衣裙"),
        ("skirt", "/skɜːt/", "n.", "裙子"),
        ("pants", "/pænts/", "n.", "裤子"),
        ("jeans", "/dʒiːnz/", "n.", "牛仔裤"),
        ("shorts", "/ʃɔːts/", "n.", "短裤"),
        ("coat", "/kəʊt/", "n.", "外套"),
        ("jacket", "/ˈdʒækɪt/", "n.", "夹克"),
        ("sweater", "/ˈswetə(r)/", "n.", "毛衣"),
        ("socks", "/sɒks/", "n.", "短袜"),
        ("shoes", "/ʃuːz/", "n.", "鞋子"),
        ("boots", "/buːts/", "n.", "靴子"),
        ("hat", "/hæt/", "n.", "帽子"),
        ("cap", "/kæp/", "n.", "鸭舌帽"),
        ("scarf", "/skɑːf/", "n.", "围巾"),
        ("gloves", "/ɡlʌvz/", "n.", "手套"),
        ("umbrella", "/ʌmˈbrelə/", "n.", "雨伞"),
        ("sunglasses", "/ˈsʌnɡlɑːsɪz/", "n.", "太阳镜"),
        ("size", "/saɪz/", "n.", "尺码"),
        ("colour", "/ˈkʌlə(r)/", "n.", "颜色"),
        ("red", "/red/", "adj.", "红色的"),
        ("blue", "/bluː/", "adj.", "蓝色的"),
        ("green", "/ɡriːn/", "adj.", "绿色的"),
        ("yellow", "/ˈjeləʊ/", "adj.", "黄色的"),
        ("white", "/waɪt/", "adj.", "白色的"),
        ("black", "/blæk/", "adj.", "黑色的"),
        ("pink", "/pɪŋk/", "adj.", "粉色的"),
        ("purple", "/ˈpɜːpl/", "adj.", "紫色的"),
        ("orange", "/ˈɒrɪndʒ/", "adj.", "橙色的"),
        ("cheap", "/tʃiːp/", "adj.", "便宜的"),
        ("expensive", "/ɪkˈspensɪv/", "adj.", "昂贵的"),
        ("pretty", "/ˈprɪti/", "adj.", "漂亮的"),
        ("try on", "/traɪ ɒn/", "v.", "试穿"),
        ("how much", "/haʊ mʌtʃ/", "adv.", "多少钱"),
    ],
    "Family & People": [
        ("father", "/ˈfɑːðə(r)/", "n.", "父亲"),
        ("mother", "/ˈmʌðə(r)/", "n.", "母亲"),
        ("brother", "/ˈbrʌðə(r)/", "n.", "兄弟"),
        ("sister", "/ˈsɪstə(r)/", "n.", "姐妹"),
        ("grandfather", "/ˈɡrænfɑːðə(r)/", "n.", "祖父"),
        ("grandmother", "/ˈɡrænmʌðə(r)/", "n.", "祖母"),
        ("uncle", "/ˈʌŋkl/", "n.", "叔叔"),
        ("aunt", "/ɑːnt/", "n.", "阿姨"),
        ("cousin", "/ˈkʌzn/", "n.", "表兄弟姐妹"),
        ("son", "/sʌn/", "n.", "儿子"),
        ("daughter", "/ˈdɔːtə(r)/", "n.", "女儿"),
        ("family", "/ˈfæməli/", "n.", "家庭"),
        ("parent", "/ˈpeərənt/", "n.", "父（母）亲"),
        ("friend", "/frend/", "n.", "朋友"),
        ("classmate", "/ˈklɑːsmeɪt/", "n.", "同学"),
        ("neighbour", "/ˈneɪbə(r)/", "n.", "邻居"),
        ("people", "/ˈpiːpl/", "n.", "人们"),
        ("child", "/tʃaɪld/", "n.", "小孩"),
        ("baby", "/ˈbeɪbi/", "n.", "婴儿"),
        ("man", "/mæn/", "n.", "男人"),
        ("woman", "/ˈwʊmən/", "n.", "女人"),
        ("boy", "/bɔɪ/", "n.", "男孩"),
        ("girl", "/ɡɜːl/", "n.", "女孩"),
        ("name", "/neɪm/", "n.", "名字"),
        ("age", "/eɪdʒ/", "n.", "年龄"),
    ],
    "Body & Health": [
        ("head", "/hed/", "n.", "头"),
        ("face", "/feɪs/", "n.", "脸"),
        ("eye", "/aɪ/", "n.", "眼睛"),
        ("ear", "/ɪə(r)/", "n.", "耳朵"),
        ("nose", "/nəʊz/", "n.", "鼻子"),
        ("mouth", "/maʊθ/", "n.", "嘴巴"),
        ("tooth", "/tuːθ/", "n.", "牙齿"),
        ("hair", "/heə(r)/", "n.", "头发"),
        ("neck", "/nek/", "n.", "脖子"),
        ("shoulder", "/ˈʃəʊldə(r)/", "n.", "肩膀"),
        ("arm", "/ɑːm/", "n.", "手臂"),
        ("hand", "/hænd/", "n.", "手"),
        ("finger", "/ˈfɪŋɡə(r)/", "n.", "手指"),
        ("leg", "/leɡ/", "n.", "腿"),
        ("foot", "/fʊt/", "n.", "脚"),
        ("knee", "/niː/", "n.", "膝盖"),
        ("back", "/bæk/", "n.", "背部"),
        ("stomach", "/ˈstʌmək/", "n.", "胃，肚子"),
        ("heart", "/hɑːt/", "n.", "心脏"),
        ("body", "/ˈbɒdi/", "n.", "身体"),
        ("medicine", "/ˈmedsn/", "n.", "药"),
        ("rest", "/rest/", "v./n.", "休息"),
        ("exercise", "/ˈeksəsaɪz/", "n./v.", "锻炼"),
    ],
    "School & Study": [
        ("school", "/skuːl/", "n.", "学校"),
        ("classroom", "/ˈklɑːsruːm/", "n.", "教室"),
        ("desk", "/desk/", "n.", "书桌"),
        ("chair", "/tʃeə(r)/", "n.", "椅子"),
        ("blackboard", "/ˈblækbɔːd/", "n.", "黑板"),
        ("book", "/bʊk/", "n.", "书"),
        ("pen", "/pen/", "n.", "钢笔"),
        ("pencil", "/ˈpensl/", "n.", "铅笔"),
        ("ruler", "/ˈruːlə(r)/", "n.", "尺子"),
        ("eraser", "/ɪˈreɪzə(r)/", "n.", "橡皮"),
        ("bag", "/bæɡ/", "n.", "包"),
        ("computer", "/kəmˈpjuːtə(r)/", "n.", "电脑"),
        ("phone", "/fəʊn/", "n.", "电话"),
        ("lesson", "/ˈlesn/", "n.", "课"),
        ("subject", "/ˈsʌbdʒɪkt/", "n.", "科目"),
        ("Chinese", "/ˌtʃaɪˈniːz/", "n.", "语文"),
        ("English", "/ˈɪŋɡlɪʃ/", "n.", "英语"),
        ("math", "/mæθ/", "n.", "数学"),
        ("music", "/ˈmjuːzɪk/", "n.", "音乐"),
        ("art", "/ɑːt/", "n.", "美术"),
        ("science", "/ˈsaɪəns/", "n.", "科学"),
        ("PE", "/ˌpiː ˈiː/", "n.", "体育"),
        ("homework", "/ˈhəʊmwɜːk/", "n.", "家庭作业"),
        ("exam", "/ɪɡˈzæm/", "n.", "考试"),
        ("test", "/test/", "n.", "测试"),
        ("question", "/ˈkwestʃən/", "n.", "问题"),
        ("answer", "/ˈɑːnsə(r)/", "n./v.", "答案；回答"),
        ("write", "/raɪt/", "v.", "写"),
        ("draw", "/drɔː/", "v.", "画"),
        ("sing", "/sɪŋ/", "v.", "唱歌"),
        ("dance", "/dɑːns/", "v.", "跳舞"),
        ("speak", "/spiːk/", "v.", "说"),
        ("listen", "/ˈlɪsn/", "v.", "听"),
        ("learn", "/lɜːn/", "v.", "学习"),
        ("study", "/ˈstʌdi/", "v.", "学习"),
        ("teach", "/tiːtʃ/", "v.", "教"),
    ],
    "Places & Transport": [
        ("home", "/həʊm/", "n.", "家"),
        ("house", "/haʊs/", "n.", "房子"),
        ("room", "/ruːm/", "n.", "房间"),
        ("bedroom", "/ˈbedruːm/", "n.", "卧室"),
        ("living room", "/ˈlɪvɪŋ ruːm/", "n.", "客厅"),
        ("kitchen", "/ˈkɪtʃɪn/", "n.", "厨房"),
        ("bathroom", "/ˈbɑːθruːm/", "n.", "浴室"),
        ("garden", "/ˈɡɑːdn/", "n.", "花园"),
        ("door", "/dɔː(r)/", "n.", "门"),
        ("window", "/ˈwɪndəʊ/", "n.", "窗户"),
        ("floor", "/flɔː(r)/", "n.", "地板"),
        ("wall", "/wɔːl/", "n.", "墙"),
        ("table", "/ˈteɪbl/", "n.", "桌子"),
        ("bed", "/bed/", "n.", "床"),
        ("sofa", "/ˈsəʊfə/", "n.", "沙发"),
        ("lamp", "/læmp/", "n.", "灯"),
        ("clock", "/klɒk/", "n.", "钟"),
        ("phone", "/fəʊn/", "n.", "电话"),
        ("bus", "/bʌs/", "n.", "公共汽车"),
        ("car", "/kɑː(r)/", "n.", "小汽车"),
        ("taxi", "/ˈtæksi/", "n.", "出租车"),
        ("bike", "/baɪk/", "n.", "自行车"),
        ("train", "/treɪn/", "n.", "火车"),
        ("plane", "/pleɪn/", "n.", "飞机"),
        ("ship", "/ʃɪp/", "n.", "轮船"),
        ("subway", "/ˈsʌbweɪ/", "n.", "地铁"),
        ("station", "/ˈsteɪʃn/", "n.", "车站"),
        ("airport", "/ˈeəpɔːt/", "n.", "机场"),
        ("road", "/rəʊd/", "n.", "路"),
        ("street", "/striːt/", "n.", "街道"),
        ("turn", "/tɜːn/", "v./n.", "转弯"),
        ("left", "/left/", "n./adj.", "左边"),
        ("right", "/raɪt/", "n./adj.", "右边"),
        ("straight", "/streɪt/", "adv.", "直走"),
        ("next to", "/nekst tuː/", "prep.", "在…旁边"),
        ("behind", "/bɪˈhaɪnd/", "prep.", "在…后面"),
        ("in front of", "/ɪn frʌnt əv/", "prep.", "在…前面"),
        ("between", "/bɪˈtwiːn/", "prep.", "在…之间"),
        ("near", "/nɪə(r)/", "prep./adj.", "在…附近"),
        ("far", "/fɑː(r)/", "adj./adv.", "远的（地）"),
    ],
    "Actions & Verbs": [
        ("run", "/rʌn/", "v.", "跑"),
        ("walk", "/wɔːk/", "v.", "走"),
        ("jump", "/dʒʌmp/", "v.", "跳"),
        ("sit", "/sɪt/", "v.", "坐"),
        ("stand", "/stænd/", "v.", "站"),
        ("open", "/ˈəʊpən/", "v.", "打开"),
        ("close", "/kləʊz/", "v.", "关"),
        ("give", "/ɡɪv/", "v.", "给"),
        ("take", "/teɪk/", "v.", "拿"),
        ("come", "/kʌm/", "v.", "来"),
        ("go", "/ɡəʊ/", "v.", "去"),
        ("see", "/siː/", "v.", "看见"),
        ("look", "/lʊk/", "v.", "看"),
        ("hear", "/hɪə(r)/", "v.", "听见"),
        ("say", "/seɪ/", "v.", "说"),
        ("tell", "/tel/", "v.", "告诉"),
        ("ask", "/ɑːsk/", "v.", "问"),
        ("think", "/θɪŋk/", "v.", "想"),
        ("know", "/nəʊ/", "v.", "知道"),
        ("want", "/wɒnt/", "v.", "想要"),
        ("like", "/laɪk/", "v.", "喜欢"),
        ("love", "/lʌv/", "v.", "爱"),
        ("help", "/help/", "v.", "帮助"),
        ("make", "/meɪk/", "v.", "制作"),
        ("use", "/juːz/", "v.", "使用"),
        ("find", "/faɪnd/", "v.", "找到"),
        ("put", "/pʊt/", "v.", "放"),
        ("get", "/ɡet/", "v.", "得到"),
        ("bring", "/brɪŋ/", "v.", "带来"),
        ("send", "/send/", "v.", "发送"),
        ("start", "/stɑːt/", "v.", "开始"),
        ("stop", "/stɒp/", "v.", "停止"),
        ("finish", "/ˈfɪnɪʃ/", "v.", "完成"),
        ("try", "/traɪ/", "v.", "尝试"),
        ("need", "/niːd/", "v.", "需要"),
        ("feel", "/fiːl/", "v.", "感觉"),
        ("play", "/pleɪ/", "v.", "玩"),
        ("work", "/wɜːk/", "v.", "工作"),
        ("live", "/lɪv/", "v.", "住"),
        ("travel", "/ˈtrævl/", "v.", "旅行"),
    ],
    "Numbers & Time": [
        ("one", "/wʌn/", "num.", "一"),
        ("two", "/tuː/", "num.", "二"),
        ("three", "/θriː/", "num.", "三"),
        ("four", "/fɔː(r)/", "num.", "四"),
        ("five", "/faɪv/", "num.", "五"),
        ("six", "/sɪks/", "num.", "六"),
        ("seven", "/ˈsevn/", "num.", "七"),
        ("eight", "/eɪt/", "num.", "八"),
        ("nine", "/naɪn/", "num.", "九"),
        ("ten", "/ten/", "num.", "十"),
        ("eleven", "/ɪˈlevn/", "num.", "十一"),
        ("twelve", "/twelv/", "num.", "十二"),
        ("thirteen", "/ˌθɜːˈtiːn/", "num.", "十三"),
        ("fourteen", "/ˌfɔːˈtiːn/", "num.", "十四"),
        ("fifteen", "/ˌfɪfˈtiːn/", "num.", "十五"),
        ("twenty", "/ˈtwenti/", "num.", "二十"),
        ("thirty", "/ˈθɜːti/", "num.", "三十"),
        ("forty", "/ˈfɔːti/", "num.", "四十"),
        ("fifty", "/ˈfɪfti/", "num.", "五十"),
        ("sixty", "/ˈsɪksti/", "num.", "六十"),
        ("seventy", "/ˈsevnti/", "num.", "七十"),
        ("eighty", "/ˈeɪti/", "num.", "八十"),
        ("ninety", "/ˈnaɪnti/", "num.", "九十"),
        ("hundred", "/ˈhʌndrəd/", "num.", "百"),
        ("thousand", "/ˈθaʊznd/", "num.", "千"),
        ("first", "/fɜːst/", "num.", "第一"),
        ("second", "/ˈsekənd/", "num.", "第二"),
        ("third", "/θɜːd/", "num.", "第三"),
        ("Monday", "/ˈmʌndeɪ/", "n.", "星期一"),
        ("Tuesday", "/ˈtjuːzdeɪ/", "n.", "星期二"),
        ("Wednesday", "/ˈwenzdeɪ/", "n.", "星期三"),
        ("Thursday", "/ˈθɜːzdeɪ/", "n.", "星期四"),
        ("Friday", "/ˈfraɪdeɪ/", "n.", "星期五"),
        ("Saturday", "/ˈsætədeɪ/", "n.", "星期六"),
        ("Sunday", "/ˈsʌndeɪ/", "n.", "星期日"),
        ("January", "/ˈdʒænjuəri/", "n.", "一月"),
        ("February", "/ˈfebruəri/", "n.", "二月"),
        ("March", "/mɑːtʃ/", "n.", "三月"),
        ("April", "/ˈeɪprəl/", "n.", "四月"),
        ("May", "/meɪ/", "n.", "五月"),
        ("June", "/dʒuːn/", "n.", "六月"),
        ("July", "/dʒuˈlaɪ/", "n.", "七月"),
        ("August", "/ˈɔːɡəst/", "n.", "八月"),
        ("September", "/sepˈtembə(r)/", "n.", "九月"),
        ("October", "/ɒkˈtəʊbə(r)/", "n.", "十月"),
        ("November", "/nəʊˈvembə(r)/", "n.", "十一月"),
        ("December", "/dɪˈsembə(r)/", "n.", "十二月"),
    ],
    "Holidays & Festivals": [
        ("Spring Festival", "/sprɪŋ ˈfestɪvl/", "n.", "春节"),
        ("Lantern Festival", "/ˈlæntən ˈfestɪvl/", "n.", "元宵节"),
        ("Dragon Boat Festival", "/ˈdræɡən bəʊt ˈfestɪvl/", "n.", "端午节"),
        ("Mid-Autumn Festival", "/mɪd ˈɔːtəm ˈfestɪvl/", "n.", "中秋节"),
        ("Christmas", "/ˈkrɪsməs/", "n.", "圣诞节"),
        ("Thanksgiving", "/ˈθæŋksɡɪvɪŋ/", "n.", "感恩节"),
        ("Halloween", "/ˌhæləʊˈiːn/", "n.", "万圣节"),
        ("New Year", "/njuː jɪə(r)/", "n.", "新年"),
        ("birthday", "/ˈbɜːθdeɪ/", "n.", "生日"),
        ("party", "/ˈpɑːti/", "n.", "聚会"),
        ("present", "/ˈpreznt/", "n.", "礼物"),
        ("card", "/kɑːd/", "n.", "贺卡"),
        ("mooncake", "/ˈmuːnkeɪk/", "n.", "月饼"),
        ("dumpling", "/ˈdʌmplɪŋ/", "n.", "饺子"),
        ("dragon boat", "/ˈdræɡən bəʊt/", "n.", "龙舟"),
        ("festival", "/ˈfestɪvl/", "n.", "节日"),
        ("celebrate", "/ˈselɪbreɪt/", "v.", "庆祝"),
    ],
    "Daily Life": [
        ("wake up", "/weɪk ʌp/", "v.", "醒来"),
        ("brush teeth", "/brʌʃ tiːθ/", "v.", "刷牙"),
        ("wash face", "/wɒʃ feɪs/", "v.", "洗脸"),
        ("comb hair", "/kəʊm heə(r)/", "v.", "梳头"),
        ("get dressed", "/ɡet drest/", "v.", "穿衣服"),
        ("have a shower", "/hæv ə ˈʃaʊə(r)/", "v.", "洗澡"),
        ("go home", "/ɡəʊ həʊm/", "v.", "回家"),
        ("come home", "/kʌm həʊm/", "v.", "回家"),
        ("clean the room", "/kliːn ðə ruːm/", "v.", "打扫房间"),
        ("tidy up", "/ˈtaɪdi ʌp/", "v.", "整理"),
        ("do the dishes", "/duː ðə ˈdɪʃɪz/", "v.", "洗碗"),
        ("make the bed", "/meɪk ðə bed/", "v.", "整理床铺"),
        ("feed the dog", "/fiːd ðə dɒɡ/", "v.", "喂狗"),
        ("water the flowers", "/ˈwɔːtə(r) ðə ˈflaʊəz/", "v.", "浇花"),
        ("go shopping", "/ɡəʊ ˈʃɒpɪŋ/", "v.", "去购物"),
        ("cook dinner", "/kʊk ˈdɪnə(r)/", "v.", "做晚饭"),
        ("set the table", "/set ðə ˈteɪbl/", "v.", "摆餐具"),
        ("clear the table", "/klɪə(r) ðə ˈteɪbl/", "v.", "收拾餐具"),
        ("take a walk", "/teɪk ə wɔːk/", "v.", "散步"),
        ("read newspapers", "/riːd ˈnjuːzpeɪpəz/", "v.", "看报纸"),
        ("listen to music", "/ˈlɪsn tə ˈmjuːzɪk/", "v.", "听音乐"),
        ("play computer games", "/pleɪ kəmˈpjuːtə ɡeɪmz/", "v.", "玩电脑游戏"),
        ("surf the Internet", "/sɜːf ðə ˈɪntənet/", "v.", "上网"),
        ("chat online", "/tʃæt ˌɒnˈlaɪn/", "v.", "在线聊天"),
        ("go out", "/ɡəʊ aʊt/", "v.", "外出"),
        ("stay at home", "/steɪ æt həʊm/", "v.", "待在家里"),
        ("go to sleep", "/ɡəʊ tə sliːp/", "v.", "去睡觉"),
        ("fall asleep", "/fɔːl əˈsliːp/", "v.", "入睡"),
        ("dream", "/driːm/", "v./n.", "做梦；梦"),
    ],
    "Common Adjectives": [
        ("good", "/ɡʊd/", "adj.", "好的"),
        ("bad", "/bæd/", "adj.", "坏的"),
        ("great", "/ɡreɪt/", "adj.", "伟大的，极好的"),
        ("nice", "/naɪs/", "adj.", "美好的"),
        ("beautiful", "/ˈbjuːtɪfl/", "adj.", "美丽的"),
        ("ugly", "/ˈʌɡli/", "adj.", "丑陋的"),
        ("clean", "/kliːn/", "adj.", "干净的"),
        ("dirty", "/ˈdɜːti/", "adj.", "脏的"),
        ("fast", "/fɑːst/", "adj.", "快的"),
        ("slow", "/sləʊ/", "adj.", "慢的"),
        ("easy", "/ˈiːzi/", "adj.", "容易的"),
        ("difficult", "/ˈdɪfɪkəlt/", "adj.", "困难的"),
        ("hard", "/hɑːd/", "adj.", "硬的，困难的"),
        ("soft", "/sɒft/", "adj.", "软的"),
        ("heavy", "/ˈhevi/", "adj.", "重的"),
        ("light", "/laɪt/", "adj.", "轻的，明亮的"),
        ("dark", "/dɑːk/", "adj.", "黑暗的"),
        ("bright", "/braɪt/", "adj.", "明亮的"),
        ("loud", "/laʊd/", "adj.", "大声的"),
        ("quiet", "/ˈkwaɪət/", "adj.", "安静的"),
        ("full", "/fʊl/", "adj.", "满的"),
        ("empty", "/ˈempti/", "adj.", "空的"),
        ("rich", "/rɪtʃ/", "adj.", "富有的"),
        ("poor", "/pɔː(r)/", "adj.", "贫穷的"),
        ("strong", "/strɒŋ/", "adj.", "强壮的"),
        ("weak", "/wiːk/", "adj.", "虚弱的"),
        ("thick", "/θɪk/", "adj.", "厚的"),
        ("thin", "/θɪn/", "adj.", "薄的，瘦的"),
        ("wide", "/waɪd/", "adj.", "宽的"),
        ("narrow", "/ˈnærəʊ/", "adj.", "窄的"),
        ("deep", "/diːp/", "adj.", "深的"),
        ("shallow", "/ˈʃæləʊ/", "adj.", "浅的"),
        ("safe", "/seɪf/", "adj.", "安全的"),
        ("dangerous", "/ˈdeɪndʒərəs/", "adj.", "危险的"),
        ("important", "/ɪmˈpɔːtnt/", "adj.", "重要的"),
        ("interesting", "/ˈɪntrəstɪŋ/", "adj.", "有趣的"),
        ("boring", "/ˈbɔːrɪŋ/", "adj.", "无聊的"),
        ("wonderful", "/ˈwʌndəfl/", "adj.", "精彩的"),
        ("terrible", "/ˈterəbl/", "adj.", "糟糕的"),
        ("famous", "/ˈfeɪməs/", "adj.", "著名的"),
    ],
    "Common Verbs 2": [
        ("eat", "/iːt/", "v.", "吃"),
        ("drink", "/drɪŋk/", "v.", "喝"),
        ("sleep", "/sliːp/", "v.", "睡觉"),
        ("wake", "/weɪk/", "v.", "醒来"),
        ("fly", "/flaɪ/", "v.", "飞"),
        ("swim", "/swɪm/", "v.", "游泳"),
        ("climb", "/klaɪm/", "v.", "爬"),
        ("crawl", "/krɔːl/", "v.", "爬行"),
        ("pull", "/pʊl/", "v.", "拉"),
        ("push", "/pʊʃ/", "v.", "推"),
        ("throw", "/θrəʊ/", "v.", "扔"),
        ("catch", "/kætʃ/", "v.", "接住"),
        ("carry", "/ˈkæri/", "v.", "搬运"),
        ("hold", "/həʊld/", "v.", "握住"),
        ("drop", "/drɒp/", "v.", "掉落"),
        ("pick", "/pɪk/", "v.", "捡，摘"),
        ("cut", "/kʌt/", "v.", "切，剪"),
        ("break", "/breɪk/", "v.", "打破"),
        ("fix", "/fɪks/", "v.", "修理"),
        ("build", "/bɪld/", "v.", "建造"),
        ("dig", "/dɪɡ/", "v.", "挖"),
        ("hang", "/hæŋ/", "v.", "悬挂"),
        ("shake", "/ʃeɪk/", "v.", "摇晃"),
        ("knock", "/nɒk/", "v.", "敲"),
        ("touch", "/tʌtʃ/", "v.", "触摸"),
        ("smell", "/smel/", "v.", "闻"),
        ("taste", "/teɪst/", "v.", "品尝"),
        ("laugh", "/lɑːf/", "v.", "笑"),
        ("cry", "/kraɪ/", "v.", "哭"),
        ("shout", "/ʃaʊt/", "v.", "喊叫"),
        ("whisper", "/ˈwɪspə(r)/", "v.", "低语"),
        ("smile", "/smaɪl/", "v.", "微笑"),
        ("remember", "/rɪˈmembə(r)/", "v.", "记住"),
        ("forget", "/fəˈɡet/", "v.", "忘记"),
        ("decide", "/dɪˈsaɪd/", "v.", "决定"),
        ("choose", "/tʃuːz/", "v.", "选择"),
        ("promise", "/ˈprɒmɪs/", "v.", "承诺"),
        ("agree", "/əˈɡriː/", "v.", "同意"),
        ("refuse", "/rɪˈfjuːz/", "v.", "拒绝"),
        ("explain", "/ɪkˈspleɪn/", "v.", "解释"),
    ],
    "Prepositions & Conjunctions": [
        ("in", "/ɪn/", "prep.", "在…里面"),
        ("on", "/ɒn/", "prep.", "在…上面"),
        ("at", "/æt/", "prep.", "在"),
        ("to", "/tə/", "prep.", "到，向"),
        ("for", "/fɔː(r)/", "prep.", "为了"),
        ("with", "/wɪð/", "prep.", "和…一起"),
        ("without", "/wɪˈðaʊt/", "prep.", "没有"),
        ("about", "/əˈbaʊt/", "prep.", "关于"),
        ("from", "/frɒm/", "prep.", "从"),
        ("of", "/ɒv/", "prep.", "…的"),
        ("by", "/baɪ/", "prep.", "通过，被"),
        ("into", "/ˈɪntə/", "prep.", "进入"),
        ("out of", "/aʊt əv/", "prep.", "从…出来"),
        ("over", "/ˈəʊvə(r)/", "prep.", "在…上方"),
        ("under", "/ˈʌndə(r)/", "prep.", "在…下面"),
        ("above", "/əˈbʌv/", "prep.", "在…上面"),
        ("below", "/bɪˈləʊ/", "prep.", "在…下面"),
        ("around", "/əˈraʊnd/", "prep.", "围绕"),
        ("through", "/θruː/", "prep.", "穿过"),
        ("across", "/əˈkrɒs/", "prep.", "横穿"),
        ("along", "/əˈlɒŋ/", "prep.", "沿着"),
        ("up", "/ʌp/", "prep./adv.", "向上"),
        ("down", "/daʊn/", "prep./adv.", "向下"),
        ("and", "/ænd/", "conj.", "和"),
        ("but", "/bʌt/", "conj.", "但是"),
        ("or", "/ɔː(r)/", "conj.", "或者"),
        ("so", "/səʊ/", "conj.", "所以"),
        ("because", "/bɪˈkɒz/", "conj.", "因为"),
        ("if", "/ɪf/", "conj.", "如果"),
        ("when", "/wen/", "conj.", "当…时"),
        ("while", "/waɪl/", "conj.", "当…时"),
        ("before", "/bɪˈfɔː(r)/", "conj.", "在…之前"),
        ("after", "/ˈɑːftə(r)/", "conj.", "在…之后"),
        ("although", "/ɔːlˈðəʊ/", "conj.", "虽然"),
        ("however", "/haʊˈevə(r)/", "adv.", "然而"),
        ("therefore", "/ˈðeəfɔː(r)/", "adv.", "因此"),
    ],
    "Question Words": [
        ("what", "/wɒt/", "adv.", "什么"),
        ("who", "/huː/", "adv.", "谁"),
        ("where", "/weə(r)/", "adv.", "哪里"),
        ("when", "/wen/", "adv.", "什么时候"),
        ("why", "/waɪ/", "adv.", "为什么"),
        ("how", "/haʊ/", "adv.", "怎样"),
        ("which", "/wɪtʃ/", "adv.", "哪一个"),
        ("whose", "/huːz/", "adv.", "谁的"),
        ("whom", "/həʊm/", "adv.", "谁（宾格）"),
        ("how many", "/haʊ ˈmeni/", "adv.", "多少（可数）"),
        ("how much", "/haʊ mʌtʃ/", "adv.", "多少（不可数）"),
        ("how old", "/haʊ əʊld/", "adv.", "多大年龄"),
        ("how often", "/haʊ ˈɒfn/", "adv.", "多久一次"),
        ("how long", "/haʊ lɒŋ/", "adv.", "多长时间"),
        ("how far", "/haʊ fɑː(r)/", "adv.", "多远"),
    ],
    "Common Phrases": [
        ("good morning", "/ɡʊd ˈmɔːnɪŋ/", "n.", "早上好"),
        ("good afternoon", "/ɡʊd ˌɑːftəˈnuːn/", "n.", "下午好"),
        ("good evening", "/ɡʊd ˈiːvnɪŋ/", "n.", "晚上好"),
        ("good night", "/ɡʊd naɪt/", "n.", "晚安"),
        ("goodbye", "/ɡʊdˈbaɪ/", "n.", "再见"),
        ("hello", "/həˈləʊ/", "n.", "你好"),
        ("hi", "/haɪ/", "n.", "嗨"),
        ("please", "/pliːz/", "adv.", "请"),
        ("thank you", "/θæŋk juː/", "v.", "谢谢你"),
        ("thanks", "/θæŋks/", "n.", "谢谢"),
        ("you're welcome", "/jɔːˈ welkəm/", "n.", "不客气"),
        ("excuse me", "/ɪkˈskjuːz miː/", "v.", "打扰一下"),
        ("I'm sorry", "/aɪm ˈsɒri/", "n.", "对不起"),
        ("that's OK", "/ðæts əʊˈkeɪ/", "n.", "没关系"),
        ("of course", "/əv ˈkɔːs/", "adv.", "当然"),
        ("no problem", "/nəʊ ˈprɒbləm/", "n.", "没问题"),
        ("well done", "/wel dʌn/", "n.", "做得好"),
        ("come on", "/kʌm ɒn/", "v.", "加油"),
        ("hurry up", "/ˈhʌri ʌp/", "v.", "快点"),
        ("be careful", "/bi ˈkeəfl/", "v.", "小心"),
        ("look out", "/lʊk aʊt/", "v.", "当心"),
        ("wait a minute", "/weɪt ə ˈmɪnɪt/", "v.", "等一下"),
        ("by the way", "/baɪ ðə weɪ/", "adv.", "顺便说一下"),
        ("in fact", "/ɪn fækt/", "adv.", "事实上"),
        ("at first", "/æt fɜːst/", "adv.", "起初"),
        ("at last", "/æt lɑːst/", "adv.", "最后"),
        ("for example", "/fɔː(r) ɪɡˈzɑːmpl/", "adv.", "例如"),
        ("as well", "/æz wel/", "adv.", "也"),
        ("right now", "/raɪt naʊ/", "adv.", "现在"),
        ("just now", "/dʒʌst naʊ/", "adv.", "刚才"),
    ],
}

# ==================== 词组句题库 ====================

SENTENCE_BANK = [
    # Unit 1 My Day
    ("I / usually / get up / at / 6:30 / in the morning",
     "I usually get up at 6:30 in the morning.",
     "我通常早上六点半起床。"),
    ("eat breakfast / I / at / seven o'clock",
     "I eat breakfast at seven o'clock.",
     "我七点钟吃早餐。"),
    ("go to school / I / often / by bus",
     "I often go to school by bus.",
     "我经常坐公共汽车去上学。"),
    ("do homework / I / in the / evening",
     "I do homework in the evening.",
     "我在晚上做作业。"),
    ("play sports / I / sometimes / after school",
     "I sometimes play sports after school.",
     "我有时放学后做运动。"),
    ("go to bed / I / at / nine o'clock / at night",
     "I go to bed at nine o'clock at night.",
     "我晚上九点钟去睡觉。"),
    ("watch TV / I / often / on the / weekend",
     "I often watch TV on the weekend.",
     "我周末经常看电视。"),
    ("read a book / I / sometimes / in the / afternoon",
     "I sometimes read a book in the afternoon.",
     "我有时下午看书。"),
    ("always / eat dinner / we / together",
     "We always eat dinner together.",
     "我们总是一起吃晚餐。"),
    ("have class / I / in the / morning",
     "I have class in the morning.",
     "我在上午上课。"),
    # Unit 2 Last Weekend
    ("I / cleaned / my room / yesterday",
     "I cleaned my room yesterday.",
     "我昨天打扫了房间。"),
    ("cooked / my mother / lunch / last Sunday",
     "My mother cooked lunch last Sunday.",
     "我妈妈上周日做了午餐。"),
    ("washed / I / my clothes / yesterday",
     "I washed my clothes yesterday.",
     "我昨天洗了衣服。"),
    ("watched / we / a film / last night",
     "We watched a film last night.",
     "我们昨晚看了一部电影。"),
    ("played / I / football / with / my friends / last weekend",
     "I played football with my friends last weekend.",
     "我上周末和朋友们踢了足球。"),
    ("visited / I / my grandparents / last Saturday",
     "I visited my grandparents last Saturday.",
     "我上周六拜访了祖父母。"),
    ("read / I / a book / yesterday / evening",
     "I read a book yesterday evening.",
     "我昨天傍晚看了一本书。"),
    ("went / I / to the park / last Sunday",
     "I went to the park last Sunday.",
     "我上周日去了公园。"),
    ("had / I / a good time / last weekend",
     "I had a good time last weekend.",
     "我上周末玩得很开心。"),
    ("stayed / I / at home / and / studied / last Saturday",
     "I stayed at home and studied last Saturday.",
     "我上周六待在家里学习。"),
    # Unit 3 Where Did You Go?
    ("went / I / to / the beach / last holiday",
     "I went to the beach last holiday.",
     "我上个假期去了海滩。"),
    ("rode / I / a horse / in / the park",
     "I rode a horse in the park.",
     "我在公园里骑了马。"),
    ("took / I / many pictures / on / my holiday",
     "I took many pictures on my holiday.",
     "我在假期拍了很多照片。"),
    ("bought / I / a gift / for / my mother",
     "I bought a gift for my mother.",
     "我给我妈妈买了一个礼物。"),
    ("ate / I / fresh food / on / the beach",
     "I ate fresh food on the beach.",
     "我在海滩上吃了新鲜的食物。"),
    ("went / I / camping / with / my family / last weekend",
     "I went camping with my family last weekend.",
     "我上周末和家人去野营了。"),
    ("swam / I / in / the lake / yesterday",
     "I swam in the lake yesterday.",
     "我昨天在湖里游泳了。"),
    ("climbed / I / a mountain / last holiday",
     "I climbed a mountain last holiday.",
     "我上个假期爬了一座山。"),
    # Unit 4 Then and Now
    ("there was / no gym / in / my old school",
     "There was no gym in my old school.",
     "我以前的学校没有体育馆。"),
    ("there were / many trees / in / the village",
     "There were many trees in the village.",
     "村庄里有很多树。"),
    ("I / was / short / then",
     "I was short then.",
     "那时我很矮。"),
    ("I / am / tall / now",
     "I am tall now.",
     "我现在很高。"),
    ("there was / a small house / here / before",
     "There was a small house here before.",
     "以前这里有一间小房子。"),
    ("there are / many buildings / here / now",
     "There are many buildings here now.",
     "现在这里有很多建筑物。"),
    ("the city / is / different / from / the village",
     "The city is different from the village.",
     "城市和村庄不同。"),
    ("I / didn't like / maths / before / but / I like it now",
     "I didn't like maths before, but I like it now.",
     "我以前不喜欢数学，但现在喜欢了。"),
    # Unit 5 What Does He Do?
    ("my father / is / a doctor / in / the hospital",
     "My father is a doctor in the hospital.",
     "我爸爸是医院里的一名医生。"),
    ("she / wants to be / a teacher",
     "She wants to be a teacher.",
     "她想成为一名教师。"),
    ("he / works / in / a company",
     "He works in a company.",
     "他在一家公司工作。"),
    ("my mother / is / a nurse",
     "My mother is a nurse.",
     "我妈妈是一名护士。"),
    ("he / is / a police officer",
     "He is a police officer.",
     "他是一名警察。"),
    ("I / want to be / a scientist / in the future",
     "I want to be a scientist in the future.",
     "我将来想成为一名科学家。"),
    ("she / teaches / English / at / our school",
     "She teaches English at our school.",
     "她在我们学校教英语。"),
    ("he / goes to work / by car / every day",
     "He goes to work by car every day.",
     "他每天开车去上班。"),
    # Unit 6 How Do You Feel?
    ("I / am / happy / today",
     "I am happy today.",
     "我今天很高兴。"),
    ("she / is / worried / about / her exam",
     "She is worried about her exam.",
     "她担心她的考试。"),
    ("he / feels / tired / after / the race",
     "He feels tired after the race.",
     "比赛后他感到疲倦。"),
    ("I / am / hungry / and / thirsty",
     "I am hungry and thirsty.",
     "我又饿又渴。"),
    ("the children / are / excited / about / the trip",
     "The children are excited about the trip.",
     "孩子们对这次旅行感到兴奋。"),
    ("don't be / afraid / of / making mistakes",
     "Don't be afraid of making mistakes.",
     "不要害怕犯错。"),
    ("I / feel / cold / in / winter",
     "I feel cold in winter.",
     "冬天我感到冷。"),
    ("she / is / proud of / her son",
     "She is proud of her son.",
     "她为她的儿子感到骄傲。"),
    ("he / was / ill / yesterday",
     "He was ill yesterday.",
     "他昨天生病了。"),
    ("we / are / surprised / to / see / the gift",
     "We are surprised to see the gift.",
     "我们看到礼物很惊讶。"),
    # 综合日常
    ("I / like / eating / apples / very much",
     "I like eating apples very much.",
     "我非常喜欢吃苹果。"),
    ("my favourite / food / is / noodles",
     "My favourite food is noodles.",
     "我最喜欢的食物是面条。"),
    ("the weather / is / sunny / and / warm / today",
     "The weather is sunny and warm today.",
     "今天天气晴朗又温暖。"),
    ("we / often / fly kites / in / spring",
     "We often fly kites in spring.",
     "我们经常在春天放风筝。"),
    ("he / goes to bed / early / every / evening",
     "He goes to bed early every evening.",
     "他每天晚上很早睡觉。"),
    ("can / I / help / you",
     "Can I help you?",
     "我能帮助你吗？"),
    ("what / do / you / usually / do / on the weekend",
     "What do you usually do on the weekend?",
     "你周末通常做什么？"),
    ("how / do / you / go / to school",
     "How do you go to school?",
     "你怎么去上学？"),
    ("where / did / you / go / last weekend",
     "Where did you go last weekend?",
     "你上周末去了哪里？"),
    ("what / did / you / do / yesterday",
     "What did you do yesterday?",
     "你昨天做了什么？"),
    # Daily Life
    ("I / wake up / at / 7 o'clock / every morning",
     "I wake up at 7 o'clock every morning.",
     "我每天早上七点醒来。"),
    ("brush / I / my teeth / before / breakfast",
     "I brush my teeth before breakfast.",
     "我在早餐前刷牙。"),
    ("go / I / shopping / with / my mother / on Sundays",
     "I go shopping with my mother on Sundays.",
     "我周日和妈妈去购物。"),
    ("take / I / a walk / after / dinner",
     "I take a walk after dinner.",
     "我晚饭后散步。"),
    ("listen / I / to music / in / my free time",
     "I listen to music in my free time.",
     "我在空闲时间听音乐。"),
    ("stay / I / at home / and / read books / on rainy days",
     "I stay at home and read books on rainy days.",
     "下雨天我待在家里看书。"),
    ("help / I / my mother / do the dishes / every evening",
     "I help my mother do the dishes every evening.",
     "我每天晚上帮妈妈洗碗。"),
    ("feed / I / the dog / before / I go to school",
     "I feed the dog before I go to school.",
     "我上学前喂狗。"),
    # Common Adjectives
    ("this / book / is / very / interesting",
     "This book is very interesting.",
     "这本书很有趣。"),
    ("the / weather / is / beautiful / today",
     "The weather is beautiful today.",
     "今天天气很好。"),
    ("the / question / is / too / difficult / for me",
     "The question is too difficult for me.",
     "这个问题对我来说太难了。"),
    ("the / box / is / too / heavy / to carry",
     "The box is too heavy to carry.",
     "这个箱子太重了，搬不动。"),
    ("she / is / a / famous / singer / in China",
     "She is a famous singer in China.",
     "她是中国著名的歌手。"),
    ("the / river / is / deep / and / dangerous",
     "The river is deep and dangerous.",
     "这条河又深又危险。"),
    ("be / careful / when / you / cross the road",
     "Be careful when you cross the road.",
     "过马路时要小心。"),
    ("the / room / is / clean / and / bright",
     "The room is clean and bright.",
     "房间又干净又明亮。"),
    # Common Verbs
    ("please / close / the door / when / you leave",
     "Please close the door when you leave.",
     "离开时请关门。"),
    ("don't / forget / to / turn off / the light",
     "Don't forget to turn off the light.",
     "别忘了关灯。"),
    ("can / you / help / me / carry / this box",
     "Can you help me carry this box?",
     "你能帮我搬这个箱子吗？"),
    ("I / want / to / buy / a present / for / my friend",
     "I want to buy a present for my friend.",
     "我想给朋友买个礼物。"),
    ("remember / to / wash / your hands / before / meals",
     "Remember to wash your hands before meals.",
     "饭前记得洗手。"),
    ("throw / the / ball / to / me / please",
     "Throw the ball to me, please.",
     "请把球扔给我。"),
    ("I / need / to / fix / my bike / this weekend",
     "I need to fix my bike this weekend.",
     "我这个周末需要修理自行车。"),
    ("the / children / are / laughing / and / crying",
     "The children are laughing and crying.",
     "孩子们又在笑又在哭。"),
    # Prepositions & Conjunctions
    ("the / cat / is / under / the table",
     "The cat is under the table.",
     "猫在桌子下面。"),
    ("I / live / in / a / small / village",
     "I live in a small village.",
     "我住在一个小村庄里。"),
    ("she / is / reading / a book / in / the library",
     "She is reading a book in the library.",
     "她正在图书馆看书。"),
    ("I / like / apples / but / I don't like bananas",
     "I like apples, but I don't like bananas.",
     "我喜欢苹果，但不喜欢香蕉。"),
    ("he / was / tired / so / he went to bed early",
     "He was tired, so he went to bed early.",
     "他很累，所以很早就睡了。"),
    ("I / stay / at home / because / it is raining",
     "I stay at home because it is raining.",
     "我待在家里，因为在下雨。"),
    ("we / can / go / to the park / if / the weather is fine",
     "We can go to the park if the weather is fine.",
     "如果天气好，我们可以去公园。"),
    ("although / he / is / young / he / is / very / clever",
     "Although he is young, he is very clever.",
     "虽然他很小，但他很聪明。"),
    # Common Phrases
    ("good morning / teacher / how / are / you",
     "Good morning, teacher. How are you?",
     "老师早上好，您好吗？"),
    ("thank you / very much / for / your help",
     "Thank you very much for your help.",
     "非常感谢你的帮助。"),
    ("excuse me / where / is / the / nearest / hospital",
     "Excuse me, where is the nearest hospital?",
     "打扰一下，最近的医院在哪里？"),
    ("I'm sorry / I / am / late / for / class",
     "I'm sorry, I am late for class.",
     "对不起，我上课迟到了。"),
    ("hurry up / or / we / will / miss / the bus",
     "Hurry up, or we will miss the bus.",
     "快点，否则我们要错过公交车了。"),
    ("be careful / the / road / is / wet / and / slippery",
     "Be careful, the road is wet and slippery.",
     "小心，路又湿又滑。"),
    ("wait / a minute / I / will / be / right back",
     "Wait a minute, I will be right back.",
     "等一下，我马上回来。"),
    ("of course / I / can / help / you / with / your homework",
     "Of course, I can help you with your homework.",
     "当然，我可以帮你做作业。"),
]

# ==================== 整句翻译题库 ====================

TRANSLATION_BANK = [
    # 中译英
    ("zh2en", "我通常六点三十分起床。",
     "I usually get up at six thirty."),
    ("zh2en", "你昨天做了什么？",
     "What did you do yesterday?"),
    ("zh2en", "我上周末去了公园。",
     "I went to the park last weekend."),
    ("zh2en", "她是一名护士，在医院工作。",
     "She is a nurse. She works in a hospital."),
    ("zh2en", "我妈妈昨天做了美味的晚餐。",
     "My mother cooked a delicious dinner yesterday."),
    ("zh2en", "你周末通常做什么？",
     "What do you usually do on the weekend?"),
    ("zh2en", "我经常在周末看电视。",
     "I often watch TV on the weekend."),
    ("zh2en", "以前这里没有体育馆。",
     "There was no gym here before."),
    ("zh2en", "现在这里有很多高楼。",
     "There are many tall buildings here now."),
    ("zh2en", "我爸爸每天开车去上班。",
     "My father goes to work by car every day."),
    ("zh2en", "我想成为一名科学家。",
     "I want to be a scientist."),
    ("zh2en", "天气晴朗又温暖。",
     "The weather is sunny and warm."),
    ("zh2en", "我们经常在春天放风筝。",
     "We often fly kites in spring."),
    ("zh2en", "你感觉怎么样？",
     "How do you feel?"),
    ("zh2en", "我感到又累又饿。",
     "I feel tired and hungry."),
    ("zh2en", "不要担心你的考试。",
     "Don't worry about your exam."),
    ("zh2en", "他昨天生病了，今天好多了。",
     "He was ill yesterday. He feels much better today."),
    ("zh2en", "你怎么去上学？",
     "How do you go to school?"),
    ("zh2en", "我通常坐公共汽车去上学。",
     "I usually go to school by bus."),
    ("zh2en", "你上周末去了哪里？",
     "Where did you go last weekend?"),
    ("zh2en", "我去了海滩，拍了很多照片。",
     "I went to the beach and took many pictures."),
    ("zh2en", "我最喜欢的食物是面条。",
     "My favourite food is noodles."),
    ("zh2en", "她昨天打扫了房间，洗了衣服。",
     "She cleaned the room and washed the clothes yesterday."),
    ("zh2en", "我七点钟吃早餐，然后去上学。",
     "I eat breakfast at seven o'clock, then go to school."),
    ("zh2en", "孩子们正在公园里开心地玩耍。",
     "The children are playing happily in the park."),
    ("zh2en", "他比我高。",
     "He is taller than me."),
    ("zh2en", "我的学校以前很小，现在很大了。",
     "My school was small before, but it is big now."),
    ("zh2en", "你喜欢什么运动？",
     "What sport do you like?"),
    ("zh2en", "我喜欢踢足球和游泳。",
     "I like playing football and swimming."),
    ("zh2en", "春节是中国最重要的节日。",
     "The Spring Festival is the most important festival in China."),
    # 英译中
    ("en2zh", "I usually get up at six thirty in the morning.",
     "我通常早上六点三十分起床。"),
    ("en2zh", "What did you do last weekend?",
     "你上周末做了什么？"),
    ("en2zh", "I went camping with my family last holiday.",
     "我上个假期和家人去野营了。"),
    ("en2zh", "My father is a doctor. He works in a hospital.",
     "我爸爸是一名医生，他在医院工作。"),
    ("en2zh", "I want to be a teacher in the future.",
     "我将来想成为一名教师。"),
    ("en2zh", "The weather is sunny and warm today.",
     "今天天气晴朗又温暖。"),
    ("en2zh", "We often fly kites in spring.",
     "我们经常在春天放风筝。"),
    ("en2zh", "How do you feel? I feel tired and hungry.",
     "你感觉怎么样？我感到又累又饿。"),
    ("en2zh", "Don't worry about your exam. You can do it!",
     "不要担心你的考试。你能做到的！"),
    ("en2zh", "He was ill yesterday, but he feels better today.",
     "他昨天生病了，但今天好多了。"),
    ("en2zh", "I usually go to school by bus.",
     "我通常坐公共汽车去上学。"),
    ("en2zh", "Where did you go last weekend? I went to the beach.",
     "你上周末去了哪里？我去了海滩。"),
    ("en2zh", "My favourite food is noodles. I like eating noodles very much.",
     "我最喜欢的食物是面条。我非常喜欢吃面条。"),
    ("en2zh", "She cleaned the room and washed the clothes yesterday.",
     "她昨天打扫了房间，洗了衣服。"),
    ("en2zh", "I eat breakfast at seven o'clock, then go to school.",
     "我七点钟吃早餐，然后去上学。"),
    ("en2zh", "The children are playing happily in the park.",
     "孩子们正在公园里开心地玩耍。"),
    ("en2zh", "My school was small before, but it is big now.",
     "我的学校以前很小，但现在很大了。"),
    ("en2zh", "What sport do you like? I like playing football and swimming.",
     "你喜欢什么运动？我喜欢踢足球和游泳。"),
    ("en2zh", "The Spring Festival is the most important festival in China.",
     "春节是中国最重要的节日。"),
    ("en2zh", "I rode a horse and took many pictures on my holiday.",
     "我在假期骑了马，拍了很多照片。"),
    ("en2zh", "There was no gym in my old school.",
     "我以前的学校没有体育馆。"),
    ("en2zh", "She is a nurse. She works in a hospital.",
     "她是一名护士，她在医院工作。"),
    ("en2zh", "I often play sports after school.",
     "我经常放学后做运动。"),
    ("en2zh", "How do you go to school? I usually go by bus.",
     "你怎么去上学？我通常坐公共汽车去。"),
    ("en2zh", "We always eat dinner together.",
     "我们总是一起吃晚餐。"),
    ("en2zh", "I bought a gift for my mother.",
     "我给我妈妈买了一个礼物。"),
    ("en2zh", "He goes to work by car every day.",
     "他每天开车去上班。"),
    ("en2zh", "I feel cold in winter.",
     "冬天我感到冷。"),
    ("en2zh", "Don't be afraid of making mistakes.",
     "不要害怕犯错。"),
    ("en2zh", "The city is different from the village.",
     "城市和村庄不同。"),
    # Daily Life
    ("zh2en", "我每天早上七点醒来。",
     "I wake up at 7 o'clock every morning."),
    ("zh2en", "我在空闲时间听音乐。",
     "I listen to music in my free time."),
    ("zh2en", "下雨天我待在家里看书。",
     "I stay at home and read books on rainy days."),
    ("zh2en", "我每天晚上帮妈妈洗碗。",
     "I help my mother do the dishes every evening."),
    ("en2zh", "I wake up at 7 o'clock every morning.",
     "我每天早上七点醒来。"),
    ("en2zh", "I listen to music in my free time.",
     "我在空闲时间听音乐。"),
    ("en2zh", "I stay at home and read books on rainy days.",
     "下雨天我待在家里看书。"),
    ("en2zh", "I help my mother do the dishes every evening.",
     "我每天晚上帮妈妈洗碗。"),
    # Adjectives
    ("zh2en", "这本书很有趣。",
     "This book is very interesting."),
    ("zh2en", "这个问题对我来说太难了。",
     "The question is too difficult for me."),
    ("zh2en", "这个箱子太重了，搬不动。",
     "The box is too heavy to carry."),
    ("zh2en", "这条河又深又危险。",
     "The river is deep and dangerous."),
    ("zh2en", "过马路时要小心。",
     "Be careful when you cross the road."),
    ("en2zh", "This book is very interesting.",
     "这本书很有趣。"),
    ("en2zh", "The question is too difficult for me.",
     "这个问题对我来说太难了。"),
    ("en2zh", "The box is too heavy to carry.",
     "这个箱子太重了，搬不动。"),
    ("en2zh", "The river is deep and dangerous.",
     "这条河又深又危险。"),
    ("en2zh", "Be careful when you cross the road.",
     "过马路时要小心。"),
    # Verbs
    ("zh2en", "离开时请关门。",
     "Please close the door when you leave."),
    ("zh2en", "别忘了关灯。",
     "Don't forget to turn off the light."),
    ("zh2en", "你能帮我搬这个箱子吗？",
     "Can you help me carry this box?"),
    ("zh2en", "我想给朋友买个礼物。",
     "I want to buy a present for my friend."),
    ("zh2en", "饭前记得洗手。",
     "Remember to wash your hands before meals."),
    ("en2zh", "Please close the door when you leave.",
     "离开时请关门。"),
    ("en2zh", "Don't forget to turn off the light.",
     "别忘了关灯。"),
    ("en2zh", "Can you help me carry this box?",
     "你能帮我搬这个箱子吗？"),
    ("en2zh", "I want to buy a present for my friend.",
     "我想给朋友买个礼物。"),
    ("en2zh", "Remember to wash your hands before meals.",
     "饭前记得洗手。"),
    # Prepositions & Conjunctions
    ("zh2en", "猫在桌子下面。",
     "The cat is under the table."),
    ("zh2en", "我住在一个小村庄里。",
     "I live in a small village."),
    ("zh2en", "她正在图书馆看书。",
     "She is reading a book in the library."),
    ("zh2en", "我喜欢苹果，但不喜欢香蕉。",
     "I like apples, but I don't like bananas."),
    ("zh2en", "他很累，所以很早就睡了。",
     "He was tired, so he went to bed early."),
    ("zh2en", "我待在家里，因为在下雨。",
     "I stay at home because it is raining."),
    ("zh2en", "如果天气好，我们可以去公园。",
     "We can go to the park if the weather is fine."),
    ("en2zh", "The cat is under the table.",
     "猫在桌子下面。"),
    ("en2zh", "I live in a small village.",
     "我住在一个小村庄里。"),
    ("en2zh", "She is reading a book in the library.",
     "她正在图书馆看书。"),
    ("en2zh", "I like apples, but I don't like bananas.",
     "我喜欢苹果，但不喜欢香蕉。"),
    ("en2zh", "He was tired, so he went to bed early.",
     "他很累，所以很早就睡了。"),
    ("en2zh", "I stay at home because it is raining.",
     "我待在家里，因为在下雨。"),
    ("en2zh", "We can go to the park if the weather is fine.",
     "如果天气好，我们可以去公园。"),
    # Common Phrases
    ("zh2en", "老师早上好，您好吗？",
     "Good morning, teacher. How are you?"),
    ("zh2en", "非常感谢你的帮助。",
     "Thank you very much for your help."),
    ("zh2en", "打扰一下，最近的医院在哪里？",
     "Excuse me, where is the nearest hospital?"),
    ("zh2en", "对不起，我上课迟到了。",
     "I'm sorry, I am late for class."),
    ("zh2en", "快点，否则我们要错过公交车了。",
     "Hurry up, or we will miss the bus."),
    ("zh2en", "小心，路又湿又滑。",
     "Be careful, the road is wet and slippery."),
    ("zh2en", "等一下，我马上回来。",
     "Wait a minute, I will be right back."),
    ("en2zh", "Good morning, teacher. How are you?",
     "老师早上好，您好吗？"),
    ("en2zh", "Thank you very much for your help.",
     "非常感谢你的帮助。"),
    ("en2zh", "Excuse me, where is the nearest hospital?",
     "打扰一下，最近的医院在哪里？"),
    ("en2zh", "I'm sorry, I am late for class.",
     "对不起，我上课迟到了。"),
    ("en2zh", "Hurry up, or we will miss the bus.",
     "快点，否则我们要错过公交车了。"),
    ("en2zh", "Be careful, the road is wet and slippery.",
     "小心，路又湿又滑。"),
    ("en2zh", "Wait a minute, I will be right back.",
     "等一下，我马上回来。"),
]

# ==================== 外部题库加载 ====================

def load_sentence_bank_from_csv(csv_path):
    """
    从CSV文件加载词组句题库
    CSV格式: scrambled,correct,chinese
    例如: I / usually / get up / at / 6:30,I usually get up at 6:30.,我通常六点三十分起床。
    """
    bank = []
    if os.path.exists(csv_path):
        with open(csv_path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            for row in reader:
                if len(row) >= 3 and row[0].strip() and not row[0].startswith('#'):
                    bank.append((row[0].strip(), row[1].strip(), row[2].strip()))
        print(f"从 {csv_path} 加载了 {len(bank)} 道词组句题目")
    return bank

def load_translation_bank_from_csv(csv_path):
    """
    从CSV文件加载翻译题题库
    CSV格式: direction,source,target
    direction: zh2en(中译英) 或 en2zh(英译中)
    例如: zh2en,我通常六点三十分起床。,I usually get up at six thirty.
    """
    bank = []
    if os.path.exists(csv_path):
        with open(csv_path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            for row in reader:
                if len(row) >= 3 and row[0].strip() and not row[0].startswith('#'):
                    bank.append((row[0].strip(), row[1].strip(), row[2].strip()))
        print(f"从 {csv_path} 加载了 {len(bank)} 道翻译题")
    return bank

# 加载外部题库（如果存在），否则使用内置题库
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SENTENCE_CSV = os.path.join(SCRIPT_DIR, "sentence_bank.csv")
TRANSLATION_CSV = os.path.join(SCRIPT_DIR, "translation_bank.csv")

# 尝试从外部CSV加载，如果不存在则使用内置数据
external_sentences = load_sentence_bank_from_csv(SENTENCE_CSV)
SENTENCE_BANK_FINAL = external_sentences if external_sentences else SENTENCE_BANK

external_translations = load_translation_bank_from_csv(TRANSLATION_CSV)
TRANSLATION_BANK_FINAL = external_translations if external_translations else TRANSLATION_BANK

# ==================== 生成函数 ====================

def generate_vocab_doc():
    """生成单词表Word文档"""
    doc = Document()

    # 页面设置 A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PEP小学英语六年级词汇表")
    run.font.size = Pt(18)
    run.bold = True
    run.font.name = "Arial"

    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("（按主题分类）")
    run.font.size = Pt(12)
    run.font.name = "Arial"

    total_words = 0
    for unit_name, words in VOCAB_DATA.items():
        total_words += len(words)

        # 单元标题
        heading = doc.add_paragraph()
        heading.space_before = Pt(12)
        run = heading.add_run(unit_name)
        run.font.size = Pt(14)
        run.bold = True
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0, 51, 102)

        # 表格
        cols = 4
        rows = (len(words) + cols - 1) // cols
        table = doc.add_table(rows=rows, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, (word, phonetic, pos, meaning) in enumerate(words):
            row_idx = i % rows
            col_idx = i // rows
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(f"{word}")
            run.font.size = Pt(10)
            run.bold = True
            run.font.name = "Arial"
            run = p.add_run(f"  {phonetic}  {pos}  {meaning}")
            run.font.size = Pt(9)
            run.font.name = "Arial"

        # 表格边框
        from docx.oxml.ns import qn
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._new_tblPr()
        borders = tblPr.find(qn('w:tblBorders'))
        if borders is None:
            from lxml import etree
            borders = etree.SubElement(tblPr, qn('w:tblBorders'))
            for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                element = etree.SubElement(borders, qn(f'w:{edge}'))
                element.set(qn('w:val'), 'single')
                element.set(qn('w:sz'), '4')
                element.set(qn('w:space'), '0')
                element.set(qn('w:color'), '999999')

    # 统计信息
    doc.add_paragraph()
    stat = doc.add_paragraph()
    stat.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = stat.add_run(f"共 {len(VOCAB_DATA)} 个主题，{total_words} 个单词/词组")
    run.font.size = Pt(11)
    run.font.name = "Arial"

    return doc


def generate_sentence_doc():
    """生成词组句题目+答案"""
    doc_q = Document()  # 题目卷
    doc_a = Document()  # 答案卷

    for doc in (doc_q, doc_a):
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 题目卷标题
    title = doc_q.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("六年级英语 词组句练习")
    run.font.size = Pt(18)
    run.bold = True
    run.font.name = "Arial"

    inst = doc_q.add_paragraph()
    inst.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = inst.add_run("要求：将下列打乱顺序的词组重新排列，组成正确的英语句子。")
    run.font.size = Pt(11)
    run.font.name = "Arial"

    # 答案卷标题
    title_a = doc_a.add_paragraph()
    title_a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_a.add_run("六年级英语 词组句练习（答案）")
    run.font.size = Pt(18)
    run.bold = True
    run.font.name = "Arial"

    # 打乱题目顺序
    sentences = list(SENTENCE_BANK_FINAL)
    random.shuffle(sentences)

    for idx, (scrambled, correct, chinese) in enumerate(sentences, 1):
        # 题目卷
        p = doc_q.add_paragraph()
        p.space_before = Pt(4)
        p.space_after = Pt(4)
        run = p.add_run(f"{idx}. ")
        run.font.size = Pt(11)
        run.bold = True
        run.font.name = "Arial"
        run = p.add_run(f"[ {scrambled} ]")
        run.font.size = Pt(11)
        run.font.name = "Arial"
        # 答题线
        run = p.add_run("  ________________________________")
        run.font.size = Pt(11)
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(180, 180, 180)

        # 答案卷
        p = doc_a.add_paragraph()
        p.space_before = Pt(3)
        p.space_after = Pt(3)
        run = p.add_run(f"{idx}. ")
        run.font.size = Pt(11)
        run.bold = True
        run.font.name = "Arial"
        run = p.add_run(f"{correct}")
        run.font.size = Pt(11)
        run.font.name = "Arial"
        run = p.add_run(f"  ({chinese})")
        run.font.size = Pt(10)
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(100, 100, 100)

    return doc_q, doc_a


def generate_translation_doc():
    """生成整句翻译题目+答案"""
    doc_q = Document()
    doc_a = Document()

    for doc in (doc_q, doc_a):
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 题目卷
    title = doc_q.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("六年级英语 整句翻译练习")
    run.font.size = Pt(18)
    run.bold = True
    run.font.name = "Arial"

    inst = doc_q.add_paragraph()
    run = inst.add_run("要求：将下列句子翻译成另一种语言（中译英或英译中）。")
    run.font.size = Pt(11)
    run.font.name = "Arial"

    # 答案卷
    title_a = doc_a.add_paragraph()
    title_a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_a.add_run("六年级英语 整句翻译练习（答案）")
    run.font.size = Pt(18)
    run.bold = True
    run.font.name = "Arial"

    # 分离中译英和英译中
    zh2en = [(d, s, a) for d, s, a in TRANSLATION_BANK_FINAL if d == "zh2en"]
    en2zh = [(d, s, a) for d, s, a in TRANSLATION_BANK_FINAL if d == "en2zh"]
    random.shuffle(zh2en)
    random.shuffle(en2zh)

    idx = 1

    # 第一部分：中译英
    sec1 = doc_q.add_paragraph()
    sec1.space_before = Pt(8)
    run = sec1.add_run("一、中译英（共%d题）" % len(zh2en))
    run.font.size = Pt(13)
    run.bold = True
    run.font.name = "Arial"

    sec1a = doc_a.add_paragraph()
    sec1a.space_before = Pt(8)
    run = sec1a.add_run("一、中译英")
    run.font.size = Pt(13)
    run.bold = True
    run.font.name = "Arial"

    for _, chinese, english in zh2en:
        p = doc_q.add_paragraph()
        p.space_before = Pt(3)
        p.space_after = Pt(3)
        run = p.add_run(f"{idx}. {chinese}")
        run.font.size = Pt(11)
        run.font.name = "Arial"
        run = p.add_run("\n________________________________")
        run.font.size = Pt(11)
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(180, 180, 180)

        pa = doc_a.add_paragraph()
        pa.space_before = Pt(2)
        pa.space_after = Pt(2)
        run = pa.add_run(f"{idx}. {english}")
        run.font.size = Pt(11)
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0, 100, 0)

        idx += 1

    # 第二部分：英译中
    sec2 = doc_q.add_paragraph()
    sec2.space_before = Pt(8)
    run = sec2.add_run("二、英译中（共%d题）" % len(en2zh))
    run.font.size = Pt(13)
    run.bold = True
    run.font.name = "Arial"

    sec2a = doc_a.add_paragraph()
    sec2a.space_before = Pt(8)
    run = sec2a.add_run("二、英译中")
    run.font.size = Pt(13)
    run.bold = True
    run.font.name = "Arial"

    for _, english, chinese in en2zh:
        p = doc_q.add_paragraph()
        p.space_before = Pt(3)
        p.space_after = Pt(3)
        run = p.add_run(f"{idx}. {english}")
        run.font.size = Pt(11)
        run.font.name = "Arial"
        run = p.add_run("\n________________________________")
        run.font.size = Pt(11)
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(180, 180, 180)

        pa = doc_a.add_paragraph()
        pa.space_before = Pt(2)
        pa.space_after = Pt(2)
        run = pa.add_run(f"{idx}. {chinese}")
        run.font.size = Pt(11)
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0, 100, 0)

        idx += 1

    return doc_q, doc_a


# ==================== 主程序 ====================

def main():
    print("=" * 50)
    print("  PEP六年级英语 单词表+词组句+翻译题 生成器")
    print("=" * 50)

    # 统计
    total_words = sum(len(v) for v in VOCAB_DATA.values())
    print(f"\n单词表：{len(VOCAB_DATA)} 个主题，{total_words} 个单词/词组")
    print(f"词组句题库：{len(SENTENCE_BANK_FINAL)} 题")
    print(f"翻译题题库：{len(TRANSLATION_BANK_FINAL)} 题")

    # 时间戳后缀
    ts = time.strftime("%Y%m%d_%H%M%S")

    # 1. 单词表
    print("\n[1/3] 生成单词表...")
    doc_vocab = generate_vocab_doc()
    fname_vocab = f"PEP六年级英语词汇表_{ts}.docx"
    doc_vocab.save(fname_vocab)
    print(f"  -> {fname_vocab}")

    # 2. 词组句
    print("[2/3] 生成词组句练习...")
    doc_sq, doc_sa = generate_sentence_doc()
    fname_sq = f"六年级英语词组句练习_题目卷_{ts}.docx"
    fname_sa = f"六年级英语词组句练习_答案卷_{ts}.docx"
    doc_sq.save(fname_sq)
    doc_sa.save(fname_sa)
    print(f"  -> {fname_sq}")
    print(f"  -> {fname_sa}")

    # 3. 翻译题
    print("[3/3] 生成整句翻译练习...")
    doc_tq, doc_ta = generate_translation_doc()
    fname_tq = f"六年级英语整句翻译练习_题目卷_{ts}.docx"
    fname_ta = f"六年级英语整句翻译练习_答案卷_{ts}.docx"
    doc_tq.save(fname_tq)
    doc_ta.save(fname_ta)
    print(f"  -> {fname_tq}")
    print(f"  -> {fname_ta}")

    print("\n[OK] 全部完成！共生成5个文件。")


if __name__ == "__main__":
    main()
