"""
小学数学应用题500道生成器（子类型系统）
12个大类，每类10个子类型，每子类型生成多道题
总数可配置，默认500道
输出：题目卷 + 答案卷 两个Word文档
依赖：pip install python-docx
"""

import random
import time
import sys
from fractions import Fraction
from math import gcd

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("请先安装 python-docx：pip install python-docx")
    exit(1)

random.seed(int(time.time()))

# ==================== 工具函数 ====================
def ri(a, b):
    return random.randint(a, b)

def rc(lst):
    return random.choice(lst)

# ==================== 1. 行程问题（10个子类型） ====================

def xingcheng_1():
    """基本行程：已知速度时间求路程"""
    v, t = ri(30, 80), ri(2, 6)
    d = v * t
    return f"一辆汽车每小时行{v}千米，行了{t}小时，一共行了多少千米？", f"路程 = {v} × {t} = {d}（千米）"

def xingcheng_2():
    """已知路程速度求时间"""
    v, t = ri(40, 80), ri(2, 5)
    d = v * t
    return f"甲乙两地相距{d}千米，汽车每小时行{v}千米，需要几小时到达？", f"时间 = {d} ÷ {v} = {t}（小时）"

def xingcheng_3():
    """已知路程时间求速度"""
    v, t = ri(30, 60), ri(3, 6)
    d = v * t
    return f"一辆汽车{t}小时行了{d}千米，平均每小时行多少千米？", f"速度 = {d} ÷ {t} = {v}（千米/小时）"

def xingcheng_4():
    """往返行程"""
    v1, t1 = ri(40, 60), ri(2, 4)
    d = v1 * t1
    v2 = ri(30, 50)
    t2 = d // v2
    while t2 * v2 != d:
        v2 = ri(30, 50)
        t2 = d // v2
    return f"从A城到B城，去时每小时行{v1}千米，用了{t1}小时。返回时每小时行{v2}千米，返回用了几小时？", f"距离 = {v1}×{t1} = {d}千米，返回时间 = {d}÷{v2} = {t2}小时"

def xingcheng_5():
    """相遇问题"""
    v1, v2 = ri(40, 60), ri(30, 50)
    t = ri(2, 4)
    d = (v1 + v2) * t
    return f"甲乙两人同时从两地相向而行，甲每小时行{v1}千米，乙每小时行{v2}千米，{t}小时后相遇。两地相距多少千米？", f"距离 = ({v1}+{v2}) × {t} = {d}（千米）"

def xingcheng_6():
    """追及问题"""
    v1, v2 = ri(50, 70), ri(30, 45)
    head_start = ri(1, 3)
    catch_time = (v2 * head_start) // (v1 - v2)
    while catch_time <= 0 or catch_time > 10:
        v1, v2 = ri(50, 70), ri(30, 45)
        head_start = ri(1, 3)
        if v1 > v2:
            catch_time = (v2 * head_start) // (v1 - v2)
    return f"甲每小时行{v1}千米，乙每小时行{v2}千米。乙先走{head_start}小时，甲几小时后追上乙？", f"追及时间 = {v2}×{head_start}÷({v1}-{v2}) = {catch_time}小时"

def xingcheng_7():
    """流水行船"""
    v_boat, v_water = ri(20, 40), ri(4, 8)
    v_down = v_boat + v_water
    v_up = v_boat - v_water
    return f"船在静水中的速度是每小时{v_boat}千米，水流速度是每小时{v_water}千米。顺水速度是多少？逆水速度是多少？", f"顺水 = {v_boat}+{v_water} = {v_down}千米/时，逆水 = {v_boat}-{v_water} = {v_up}千米/时"

def xingcheng_8():
    """环形跑道"""
    v1, v2 = ri(200, 300), ri(150, 200)
    track = ri(400, 800)
    meet_time = track // (v1 - v2)
    while meet_time <= 0 or meet_time > 20:
        v1, v2 = ri(200, 300), ri(150, 200)
        track = ri(400, 800)
        meet_time = track // (v1 - v2)
    return f"环形跑道长{track}米，甲每分钟跑{v1}米，乙每分钟跑{v2}米，同时同地同向出发，几分钟后首次相遇？", f"相遇时间 = {track}÷({v1}-{v2}) = {meet_time}分钟"

def xingcheng_9():
    """过桥问题"""
    train_len, bridge_len = ri(100, 300), ri(200, 500)
    speed = ri(20, 40)
    total_dist = train_len + bridge_len
    time = total_dist // speed
    while time * speed != total_dist:
        speed = ri(20, 40)
        time = total_dist // speed
    return f"一列火车长{train_len}米，通过一座长{bridge_len}米的桥，速度是每秒{speed}米，需要多少秒？", f"总距离 = {train_len}+{bridge_len} = {total_dist}米，时间 = {total_dist}÷{speed} = {time}秒"

def xingcheng_10():
    """平均速度"""
    d1, v1 = ri(60, 120), ri(30, 50)
    d2, v2 = ri(60, 120), ri(40, 60)
    t1 = d1 // v1
    t2 = d2 // v2
    while t1 * v1 != d1 or t2 * v2 != d2:
        d1, v1 = ri(60, 120), ri(30, 50)
        d2, v2 = ri(60, 120), ri(40, 60)
        t1 = d1 // v1
        t2 = d2 // v2
    avg_v = (d1 + d2) // (t1 + t2)
    return f"上山{d1}千米，速度{v1}千米/时；下山{d2}千米，速度{v2}千米/时。上下山平均速度是多少？", f"总路程 = {d1+d2}，总时间 = {t1}+{t2} = {t1+t2}，平均 = {d1+d2}÷{t1+t2} = {avg_v}千米/时"

# ==================== 2. 工程问题（10个子类型） ====================

def gongcheng_1():
    """合作完成"""
    d1, d2 = ri(4, 10), ri(5, 12)
    lcm = d1 * d2 // gcd(d1, d2)
    w1, w2 = lcm // d1, lcm // d2
    together = lcm // (w1 + w2)
    return f"一项工程，甲独做{d1}天完成，乙独做{d2}天完成。合做几天完成？", f"甲效1/{d1}，乙效1/{d2}，合效={w1}/{lcm}+{w2}/{lcm}={w1+w2}/{lcm}，天数={lcm}/{w1+w2}={together}天"

def gongcheng_2():
    """先做后合"""
    d1, d2 = ri(6, 12), ri(8, 15)
    alone_days = ri(2, 4)
    lcm = d1 * d2 // gcd(d1, d2)
    w1, w2 = lcm // d1, lcm // d2
    remain = lcm - w1 * alone_days
    together = remain // (w1 + w2)
    return f"一项工程，甲独做{d1}天完成，乙独做{d2}天完成。甲先做{alone_days}天，剩下的合做几天完成？", f"甲完成{w1*alone_days}/{lcm}，剩{lcm-w1*alone_days}/{lcm}，合做={remain}÷({w1}+{w2})={together}天"

def gongcheng_3():
    """三人合作"""
    d1, d2, d3 = ri(6, 10), ri(8, 12), ri(10, 15)
    lcm = d1 * d2 // gcd(d1, d2)
    lcm = lcm * d3 // gcd(lcm, d3)
    w1, w2, w3 = lcm // d1, lcm // d2, lcm // d3
    together = lcm // (w1 + w2 + w3)
    return f"一项工程，甲独做{d1}天完成，乙独做{d2}天完成，丙独做{d3}天完成。三人合做几天完成？", f"三人效率和={w1}+{w2}+{w3}={w1+w2+w3}/{lcm}，天数={lcm}/{w1+w2+w3}={together}天"

def gongcheng_4():
    """效率比较"""
    d1, d2 = ri(5, 10), ri(8, 15)
    while d1 == d2:
        d2 = ri(8, 15)
    lcm = d1 * d2 // gcd(d1, d2)
    w1, w2 = lcm // d1, lcm // d2
    diff = abs(w1 - w2)
    faster = "甲" if w1 > w2 else "乙"
    return f"甲独做{d1}天完成，乙独做{d2}天完成。谁效率高？每天多完成几分之几？", f"甲效1/{d1}={w1}/{lcm}，乙效1/{d2}={w2}/{lcm}，{faster}效率高，多{diff}/{lcm}"

def gongcheng_5():
    """中途休息"""
    d1, d2 = ri(6, 10), ri(8, 12)
    rest_days = ri(1, 3)
    lcm = d1 * d2 // gcd(d1, d2)
    w1, w2 = lcm // d1, lcm // d2
    total_work = lcm
    b_work = total_work - w1 * rest_days
    b_days = b_work // w2
    total_days = rest_days + b_days
    return f"一项工程，甲独做{d1}天完成，乙独做{d2}天完成。甲先做{rest_days}天后休息，乙接着做，共需几天完成？", f"甲做{w1*rest_days}/{lcm}，剩{lcm-w1*rest_days}/{lcm}，乙做{b_work}÷{w2}={b_days}天，共{rest_days}+{b_days}={total_days}天"

def gongcheng_6():
    """交替工作"""
    d1, d2 = ri(4, 8), ri(6, 10)
    lcm = d1 * d2 // gcd(d1, d2)
    w1, w2 = lcm // d1, lcm // d2
    cycle_work = w1 + w2
    cycles = lcm // cycle_work
    remain = lcm % cycle_work
    if remain <= w1:
        extra = 1
    else:
        extra = 2
    total_days = cycles * 2 + extra
    return f"一项工程，甲独做{d1}天完成，乙独做{d2}天完成。两人交替各做1天，几天完成？", f"每2天完成{w1}+{w2}={cycle_work}/{lcm}，{cycles}个周期后剩{remain}，还需{extra}天，共{total_days}天"

def gongcheng_7():
    """注水问题"""
    fill_time, drain_time = ri(4, 8), ri(6, 12)
    lcm = fill_time * drain_time // gcd(fill_time, drain_time)
    fill_rate = lcm // fill_time
    drain_rate = lcm // drain_time
    net_rate = fill_rate - drain_rate
    fill_days = lcm // net_rate
    return f"一个水池，注水管{fill_time}小时注满，排水管{drain_time}小时排完。同时开两管，几小时注满？", f"注速1/{fill_time}={fill_rate}/{lcm}，排速1/{drain_time}={drain_rate}/{lcm}，净速={fill_rate-drain_rate}/{lcm}，时间={lcm}/{net_rate}={fill_days}小时"

def gongcheng_8():
    """部分完成"""
    d1, d2 = ri(6, 10), ri(8, 12)
    work_pct = random.choice([25, 30, 40, 50, 60])
    lcm = d1 * d2 // gcd(d1, d2)
    w1, w2 = lcm // d1, lcm // d2
    target_work = lcm * work_pct // 100
    together_days = target_work // (w1 + w2)
    return f"一项工程，甲独做{d1}天完成，乙独做{d2}天完成。合做几天完成工程的{work_pct}%？", f"合效={w1+w2}/{lcm}，{work_pct}%={target_work}/{lcm}，天数={target_work}÷({w1}+{w2})={together_days}天"

def gongcheng_9():
    """中途加入"""
    d1, d2 = ri(8, 12), ri(10, 15)
    late_days = ri(2, 4)
    lcm = d1 * d2 // gcd(d1, d2)
    w1, w2 = lcm // d1, lcm // d2
    work_alone = w1 * late_days
    remain = lcm - work_alone
    together = remain // (w1 + w2)
    total = late_days + together
    return f"一项工程，甲独做{d1}天完成，乙独做{d2}天完成。甲先做{late_days}天后乙加入，合做几天完成？总共几天？", f"甲先做{work_alone}/{lcm}，剩{lcm-work_alone}/{lcm}，合做={remain}÷({w1}+{w2})={together}天，共{late_days}+{together}={total}天"

def gongcheng_10():
    """效率提升"""
    d1 = ri(10, 20)
    pct = random.choice([20, 25, 50])
    new_rate = d1 * 100 // (100 + pct)
    while new_rate * (100 + pct) != d1 * 100:
        d1 = ri(10, 20)
        new_rate = d1 * 100 // (100 + pct)
    return f"一项工程原计划{d1}天完成，实际效率提高了{pct}%，实际几天完成？", f"原效率1/{d1}，新效率=1/{d1}×(1+{pct}%)={100+pct}/({d1*100})，天数={d1*100}÷{100+pct}={new_rate}天"

# ==================== 3. 浓度问题（10个子类型） ====================

def nongdu_1():
    """求溶质"""
    solution = random.choice([100, 200, 300, 400, 500])
    pct = random.choice([10, 15, 20, 25, 30])
    solute = solution * pct // 100
    return f"{solution}克盐水含盐{pct}%，含盐多少克？", f"含盐 = {solution} × {pct}% = {solute}克"

def nongdu_2():
    """求浓度"""
    solute, solution = ri(10, 50), random.choice([100, 200, 250, 500])
    while solute > solution:
        solute = ri(10, 50)
    pct = solute * 100 // solution
    while pct * solution != solute * 100:
        solution = random.choice([100, 200, 250, 500])
        pct = solute * 100 // solution
    return f"{solution}克盐水中含盐{solute}克，浓度是多少？", f"浓度 = {solute}÷{solution}×100% = {pct}%"

def nongdu_3():
    """加水稀释"""
    solution1, pct1 = ri(100, 300), random.choice([20, 25, 30])
    add_water = ri(50, 200)
    solute = solution1 * pct1 // 100
    solution2 = solution1 + add_water
    pct2 = solute * 100 // solution2
    while pct2 * solution2 != solute * 100:
        add_water = ri(50, 200)
        solution2 = solution1 + add_water
        pct2 = solute * 100 // solution2
    return f"{solution1}克盐水含盐{pct1}%，加入{add_water}克水后，浓度变为多少？", f"盐={solution1}×{pct1}%={solute}克，新溶液={solution2}克，浓度={solute}÷{solution2}×100%={pct2}%"

def nongdu_4():
    """加盐增浓"""
    solution1, pct1 = ri(200, 400), random.choice([10, 15, 20])
    add_salt = ri(10, 50)
    solute1 = solution1 * pct1 // 100
    solute2 = solute1 + add_salt
    solution2 = solution1 + add_salt
    pct2 = solute2 * 100 // solution2
    while pct2 * solution2 != solute2 * 100:
        add_salt = ri(10, 50)
        solute2 = solute1 + add_salt
        solution2 = solution1 + add_salt
        pct2 = solute2 * 100 // solution2
    return f"{solution1}克盐水含盐{pct1}%，加入{add_salt}克盐后，浓度变为多少？", f"原盐={solute1}克，新盐={solute2}克，新溶液={solution2}克，浓度={pct2}%"

def nongdu_5():
    """混合问题"""
    s1, p1 = ri(100, 200), random.choice([10, 20])
    s2, p2 = ri(100, 200), random.choice([20, 30])
    solute = s1 * p1 // 100 + s2 * p2 // 100
    solution = s1 + s2
    pct = solute * 100 // solution
    while pct * solution != solute * 100:
        s1, s2 = ri(100, 200), ri(100, 200)
        solution = s1 + s2
        solute = s1 * p1 // 100 + s2 * p2 // 100
        pct = solute * 100 // solution
    return f"{p1}%盐水{s1}克和{p2}%盐水{s2}克混合，浓度是多少？", f"盐={s1*p1//100}+{s2*p2//100}={solute}克，溶液={solution}克，浓度={pct}%"

def nongdu_6():
    """蒸发浓缩"""
    solution1, pct1 = ri(200, 400), random.choice([10, 15])
    evap = ri(50, 150)
    solute = solution1 * pct1 // 100
    solution2 = solution1 - evap
    pct2 = solute * 100 // solution2
    while pct2 * solution2 != solute * 100 or solution2 <= 0:
        evap = ri(50, 150)
        solution2 = solution1 - evap
        pct2 = solute * 100 // solution2
    return f"{solution1}克盐水含盐{pct1}%，蒸发{evap}克水后，浓度变为多少？", f"盐={solute}克不变，新溶液={solution2}克，浓度={solute}÷{solution2}×100%={pct2}%"

def nongdu_7():
    """求溶液"""
    solute, pct = ri(20, 80), random.choice([10, 20, 25, 40])
    solution = solute * 100 // pct
    return f"含盐{solute}克的盐水，浓度是{pct}%，盐水共多少克？", f"溶液 = {solute}÷{pct}% = {solution}克"

def nongdu_8():
    """配制问题"""
    target_solution, target_pct = random.choice([200, 300, 500]), random.choice([15, 20, 25])
    high_pct, low_pct = 30, 10
    # 十字交叉法
    high_ratio = target_pct - low_pct
    low_ratio = high_pct - target_pct
    total_ratio = high_ratio + low_ratio
    high_solution = target_solution * high_ratio // total_ratio
    low_solution = target_solution - high_solution
    return f"配制{target_pct}%盐水{target_solution}克，需{high_pct}%和{low_pct}%盐水各多少克？", f"高浓度比={high_ratio}，低浓度比={low_ratio}，高={high_solution}克，低={low_solution}克"

def nongdu_9():
    """多次稀释"""
    solution, pct = ri(200, 400), random.choice([30, 40])
    remove = ri(50, 100)
    add_water = remove
    solute = solution * pct // 100
    # 第一次
    solute1 = solute * (solution - remove) // solution
    solution1 = solution
    # 加水后
    solution2 = solution1
    pct2 = solute1 * 100 // solution2
    while pct2 * solution2 != solute1 * 100:
        remove = ri(50, 100)
        solute1 = solute * (solution - remove) // solution
        pct2 = solute1 * 100 // solution
    return f"{solution}克{pct}%盐水，倒出{remove}克后加满水，浓度是多少？", f"剩余盐={solute1}克，溶液仍{solution}克，浓度={pct2}%"

def nongdu_10():
    """溶质守恒"""
    solution1, pct1 = ri(100, 300), random.choice([20, 25, 30])
    target_pct = random.choice([10, 15])
    solute = solution1 * pct1 // 100
    solution2 = solute * 100 // target_pct
    add_water = solution2 - solution1
    return f"{solution1}克{pct1}%盐水稀释成{target_pct}%，需加多少克水？", f"盐={solute}克不变，新溶液={solution2}克，加水={add_water}克"

# ==================== 继续其他大类... ====================
# 为节省篇幅，这里简写，实际代码包含所有12个大类

# 4. 利润折扣
def lirun_1():
    cost = ri(50, 200)
    pct = random.choice([20, 30, 50])
    sell = cost * (100 + pct) // 100
    return f"成本{cost}元，加价{pct}%出售，售价多少？", f"售价 = {cost}×(1+{pct}%) = {sell}元"

def lirun_2():
    sell, cost = ri(100, 300), ri(50, 200)
    while sell <= cost:
        sell, cost = ri(100, 300), ri(50, 200)
    profit = sell - cost
    pct = profit * 100 // cost
    while pct * cost != profit * 100:
        cost = ri(50, 200)
        profit = sell - cost
        pct = profit * 100 // cost
    return f"售价{sell}元，成本{cost}元，利润率是多少？", f"利润={profit}元，利润率={profit}÷{cost}×100%={pct}%"

def lirun_3():
    price = ri(100, 500)
    discount = random.choice([80, 85, 90])
    final = price * discount // 100
    return f"原价{price}元，打{discount//10}折，现价多少？", f"现价 = {price}×{discount//10}0% = {final}元"

def lirun_4():
    final, discount = ri(100, 400), random.choice([80, 85, 90])
    original = final * 100 // discount
    while original * discount // 100 != final:
        discount = random.choice([80, 85, 90])
        original = final * 100 // discount
    return f"打{discount//10}折后{final}元，原价多少？", f"原价 = {final}÷{discount//10}0% = {original}元"

def lirun_5():
    buy, sell = ri(50, 150), ri(80, 200)
    while sell <= buy:
        sell = ri(80, 200)
    qty = ri(10, 50)
    profit = (sell - buy) * qty
    return f"进价{buy}元，售价{sell}元，卖出{qty}件，总利润多少？", f"单利润={sell-buy}元，总利润={sell-buy}×{qty}={profit}元"

def lirun_6():
    cost, fix_pct = ri(100, 300), random.choice([10, 15, 20])
    discount = random.choice([85, 90])
    sell = cost * (100 + fix_pct) // 100
    actual = sell * discount // 100
    profit = actual - cost
    return f"成本{cost}元，标价获利{fix_pct}%，打{discount//10}折出售，实际获利多少？", f"标价={sell}元，实际={actual}元，获利={profit}元"

def lirun_7():
    total_cost = ri(1000, 5000)
    qty = ri(50, 200)
    unit_cost = total_cost // qty
    while unit_cost * qty != total_cost:
        qty = ri(50, 200)
        unit_cost = total_cost // qty
    sell_price = ri(unit_cost + 10, unit_cost + 50)
    profit = (sell_price - unit_cost) * qty
    return f"总成本{total_cost}元购入{qty}件，每件售价{sell_price}元，总利润多少？", f"单成本={unit_cost}元，总利润=({sell_price}-{unit_cost})×{qty}={profit}元"

def lirun_8():
    original = ri(200, 800)
    d1, d2 = random.choice([90, 85]), random.choice([95, 90])
    final = original * d1 // 100 * d2 // 100
    return f"原价{original}元，先打{d1//10}折再打{d2//10}折，最终价格多少？", f"最终 = {original}×{d1//10}0%×{d2//10}0% = {final}元"

def lirun_9():
    cost, target_pct = ri(100, 300), random.choice([20, 30, 50])
    discount = random.choice([85, 90])
    sell = cost * (100 + target_pct) // 100 // (discount // 100)
    while sell * discount // 100 < cost:
        sell += 1
    return f"成本{cost}元，想获利{target_pct}%，打{discount//10}折后仍获利，至少标价多少？", f"设标价x，x×{discount}%≥{cost}×(1+{target_pct}%)，x≥{sell}元"

def lirun_10():
    buy_qty, buy_price = ri(100, 500), ri(5, 20)
    sell_qty = int(buy_qty * 0.8)
    sell_price = buy_price * 2
    revenue = sell_price * sell_qty
    cost = buy_price * buy_qty
    profit = revenue - cost
    return f"进货{buy_qty}件，每件{buy_price}元。售出80%，售价翻倍，利润多少？", f"成本={cost}元，收入={revenue}元，利润={profit}元"

# 6. 比例问题
def bili_1():
    a, b = ri(2, 5), ri(3, 7)
    total = random.choice([100, 200, 300, 400, 500])
    g = gcd(a, b)
    sa, sb = a // g, b // g
    pa = total * sa // (sa + sb)
    pb = total - pa
    return f"甲乙之比{a}:{b}，和为{total}，各是多少？", f"总份{sa+sb}，甲={total}×{sa}/{sa+sb}={pa}，乙={pb}"

def bili_2():
    ratio = random.choice([(2,3), (3,4), (3,5), (4,5)])
    known = ri(20, 100)
    other = known * ratio[1] // ratio[0]
    while other * ratio[0] != known * ratio[1]:
        known = ri(20, 100)
        other = known * ratio[1] // ratio[0]
    return f"甲乙之比{ratio[0]}:{ratio[1]}，甲是{known}，乙是多少？", f"乙 = {known}×{ratio[1]}÷{ratio[0]} = {other}"

def bili_3():
    a, b = ri(2, 4), ri(4, 6)  # ensure b > a so that diff is positive
    if a >= b:
        a, b = b, a
    diff = ri(10, 50)
    g = gcd(a, b)
    sa, sb = a // g, b // g
    if sb - sa <= 0:
        sa, sb = 1, 2
    unit = diff // (sb - sa)
    while unit * (sb - sa) != diff:
        diff = ri(10, 50)
        unit = diff // (sb - sa)
    pa, pb = unit * sa, unit * sb
    return f"甲乙之比{sa}:{sb}，乙比甲多{diff}，各是多少？", f"每份={diff}÷{sb-sa}={unit}，甲={pa}，乙={pb}"

def bili_4():
    scale = random.choice([100, 200, 500, 1000, 5000])
    map_dist = ri(2, 20)
    real = map_dist * scale
    if real >= 100000:
        return f"比例尺1:{scale}，图上{map_dist}cm，实际多少千米？", f"实际 = {map_dist}×{scale} = {real}cm = {real/100000}千米"
    return f"比例尺1:{scale}，图上{map_dist}cm，实际多少米？", f"实际 = {map_dist}×{scale} = {real}cm = {real/100}米"

def bili_5():
    a, b = ri(2, 4), ri(3, 5)
    total = random.choice([60, 90, 120, 180])
    g = gcd(a, b)
    sa, sb = a // g, b // g
    pa = total * sa // (sa + sb)
    pb = total - pa
    return f"按{a}:{b}分配{total}，两部分各是多少？", f"总份{sa+sb}，一部分={pa}，另一部分={pb}"

def bili_6():
    speed1, speed2 = ri(40, 60), ri(30, 50)
    time1, time2 = ri(2, 4), ri(3, 5)
    d1, d2 = speed1 * time1, speed2 * time2
    g = gcd(d1, d2)
    return f"甲速{speed1}行{time1}时，乙速{speed2}行{time2}时，路程比是多少？", f"路程比 = {d1}:{d2} = {d1//g}:{d2//g}"

def bili_7():
    ratio = random.choice([(2,3), (3,4), (4,5)])
    old_a, old_b = ri(10, 30), ri(10, 30)
    new_a = old_a * ratio[0] // ratio[1]
    while new_a * ratio[1] != old_a * ratio[0]:
        old_a = ri(10, 30)
        new_a = old_a * ratio[0] // ratio[1]
    return f"甲乙之比原为{old_a}:{old_b}，现甲变为{new_a}，比例变为{ratio[0]}:{ratio[1]}，乙变了吗？", f"原比例{old_a}:{old_b}，新比例{new_a}:{old_b}，乙不变"

def bili_8():
    a, b, c = ri(2, 4), ri(3, 5), ri(4, 6)
    total = random.choice([180, 270, 360, 450])
    g = gcd(gcd(a, b), c)
    sa, sb, sc = a // g, b // g, c // g
    unit = total // (sa + sb + sc)
    while unit * (sa + sb + sc) != total:
        total = random.choice([180, 270, 360, 450])
        unit = total // (sa + sb + sc)
    pa, pb, pc = unit * sa, unit * sb, unit * sc
    return f"甲乙丙之比{a}:{b}:{c}，总和{total}，各是多少？", f"总份{sa+sb+sc}，甲={pa}，乙={pb}，丙={pc}"

def bili_9():
    pct = random.choice([25, 50, 75, 80])
    return f"甲是乙的{pct}%，甲乙之比是多少？", f"甲:乙 = {pct}:100 = {pct//gcd(pct,100)}:{100//gcd(pct,100)}"

def bili_10():
    a, b = ri(2, 5), ri(3, 6)
    g = gcd(a, b)
    sa, sb = a // g, b // g
    new_a = sa * 2
    new_b = sb * 2
    return f"甲乙之比{a}:{b}，都扩大2倍后，比值变吗？", f"原比{sa}:{sb}，扩大后{new_a}:{new_b}={sa}:{sb}，比值不变"

# 7. 年龄问题
def age_1():
    age_now = ri(8, 15)
    diff = ri(20, 35)
    years = ri(3, 10)
    parent_now = age_now + diff
    parent_later = parent_now + years
    return f"小明{age_now}岁，爸大{diff}岁，{years}年后爸几岁？", f"爸现在{parent_now}岁，{years}年后{parent_later}岁"

def age_2():
    age1, age2 = ri(8, 12), ri(35, 45)
    years = ri(5, 15)
    later1, later2 = age1 + years, age2 + years
    ratio = later2 // later1
    while ratio * later1 != later2:
        years = ri(5, 15)
        later1, later2 = age1 + years, age2 + years
        ratio = later2 // later1
    return f"小明{age1}岁，爸{age2}岁，几年后爸年龄是小明的{ratio}倍？", f"{years}年后小明{later1}岁，爸{later2}岁，{ratio}倍"

def age_3():
    sum_age = random.choice([60, 70, 80, 90])
    diff = ri(2, 8)
    older = (sum_age + diff) // 2
    younger = sum_age - older
    return f"兄弟年龄和{sum_age}，哥比弟大{diff}岁，各几岁？", f"哥={older}岁，弟={younger}岁"

def age_4():
    ages = [ri(8, 12) for _ in range(3)]
    total = sum(ages)
    avg = total // 3
    while avg * 3 != total:
        ages = [ri(8, 12) for _ in range(3)]
        total = sum(ages)
        avg = total // 3
    return f"三人年龄和{total}，平均年龄是多少？", f"平均 = {total}÷3 = {avg}岁"

def age_5():
    age_now = ri(10, 15)
    years_ago = ri(3, 8)
    age_then = age_now - years_ago
    ratio = ri(3, 6)
    parent_then = age_then * ratio
    parent_now = parent_then + years_ago
    return f"小明{age_now}岁，{years_ago}年前爸年龄是他的{ratio}倍，爸现在几岁？", f"{years_ago}年前小明{age_then}岁，爸{parent_then}岁，现在{parent_now}岁"

def age_6():
    age1, age2 = ri(30, 40), ri(2, 8)
    years = ri(10, 20)
    diff = age1 - age2
    return f"爸{age1}岁，儿子{age2}岁，{years}年后相差几岁？", f"年龄差永远不变，相差{diff}岁"

def age_7():
    ages = [ri(30, 45), ri(28, 42), ri(8, 12)]
    total = sum(ages)
    return f"爸{ages[0]}岁，妈{ages[1]}岁，孩子{ages[2]}岁，全家年龄和？", f"全家 = {ages[0]}+{ages[1]}+{ages[2]} = {total}岁"

def age_8():
    age1, age2 = ri(35, 45), ri(8, 12)
    years = ri(5, 15)
    sum_future = age1 + age2 + years * 2
    return f"爸{age1}岁，儿子{age2}岁，{years}年后父子年龄和？", f"{years}年后爸{age1+years}岁，儿子{age2+years}岁，和={sum_future}岁"

def age_9():
    age_now = ri(10, 15)
    years_later = ri(5, 10)
    age_later = age_now + years_later
    ratio = random.choice([2, 3])
    parent_later = age_later * ratio
    parent_now = parent_later - years_later
    return f"小明{age_now}岁，{years_later}年后妈年龄是他的{ratio}倍，妈现在几岁？", f"{years_later}年后小明{age_later}岁，妈{parent_later}岁，现在{parent_now}岁"

def age_10():
    age1, age2 = ri(30, 40), ri(5, 10)
    years_ago = ri(5, 10)
    age1_then = age1 - years_ago
    age2_then = age2 - years_ago
    ratio = age1_then // age2_then
    while ratio * age2_then != age1_then or years_ago >= age2:
        years_ago = ri(5, 10)
        age1_then = age1 - years_ago
        age2_then = age2 - years_ago
        if age2_then > 0:
            ratio = age1_then // age2_then
    return f"爸{age1}岁，儿子{age2}岁，{years_ago}年前爸年龄是儿子的几倍？", f"{years_ago}年前爸{age1_then}岁，儿子{age2_then}岁，{ratio}倍"

# 8. 鸡兔同笼
def jitu_1():
    head = ri(10, 30)
    rabbit = ri(3, head - 3)
    chicken = head - rabbit
    leg = rabbit * 4 + chicken * 2
    return f"鸡兔同笼，{head}个头，{leg}条腿，鸡兔各几只？", f"假设全鸡：{head*2}条，差{leg-head*2}条，兔={leg-head*2}÷2={rabbit}只，鸡={chicken}只"

def jitu_2():
    head = ri(15, 35)
    leg_diff = ri(10, 40)
    rabbit = (head * 2 + leg_diff) // 4
    while rabbit * 4 != head * 2 + leg_diff:
        leg_diff = ri(10, 40)
        rabbit = (head * 2 + leg_diff) // 4
    chicken = head - rabbit
    leg = rabbit * 4 + chicken * 2
    diff = leg - head * 2
    return f"鸡兔{head}只，兔腿比鸡腿多{leg_diff}条，各几只？", f"兔{rabbit}只，鸡{chicken}只"

def jitu_3():
    coin_total = random.choice([20, 30, 50])
    coin_value = random.choice([50, 80, 100])
    # 1元和5角
    one_yuan = (coin_value - coin_total * 5) // 5
    while one_yuan * 10 + (coin_total - one_yuan) * 5 != coin_value:
        coin_total = random.choice([20, 30, 50])
        one_yuan = (coin_value - coin_total * 5) // 5
    half_yuan = coin_total - one_yuan
    return f"{coin_total}枚硬币（1元和5角）共{coin_value}角，各几枚？", f"1元{one_yuan}枚，5角{half_yuan}枚"

def jitu_4():
    vehicle = ri(10, 25)
    wheel = ri(30, 80)
    car = (wheel - vehicle * 2) // 2
    while car * 4 + (vehicle - car) * 2 != wheel:
        wheel = ri(30, 80)
        car = (wheel - vehicle * 2) // 2
    bike = vehicle - car
    return f"{vehicle}辆车（汽车4轮，自行车2轮）共{wheel}个轮，各几辆？", f"汽车{car}辆，自行车{bike}辆"

def jitu_5():
    head = ri(20, 40)
    leg = ri(50, 120)
    rabbit = (leg - head * 2) // 2
    while rabbit * 4 + (head - rabbit) * 2 != leg:
        leg = ri(50, 120)
        rabbit = (leg - head * 2) // 2
    chicken = head - rabbit
    return f"鸡兔{head}只，共{leg}条腿，用方程解各几只？", f"设兔x只，4x+2({head}-x)={leg}，x={rabbit}，鸡={chicken}只"

def jitu_6():
    head = ri(10, 25)
    leg = ri(30, 80)
    rabbit = (leg - head * 2) // 2
    chicken = head - rabbit
    diff = rabbit - chicken
    return f"鸡兔{head}只，{leg}条腿，兔比鸡多几只？", f"兔{rabbit}只，鸡{chicken}只，多{diff}只"

def jitu_7():
    head = ri(15, 30)
    leg = ri(40, 100)
    rabbit = (leg - head * 2) // 2
    chicken = head - rabbit
    return f"鸡兔{head}只，{leg}条腿，如果全是鸡有几条腿？全是兔呢？", f"全鸡{head*2}条，全兔{head*4}条，实际{leg}条"

def jitu_8():
    head = ri(20, 40)
    leg = ri(60, 130)
    rabbit = (leg - head * 2) // 2
    chicken = head - rabbit
    leg_change = rabbit * 4 + chicken * 2 + 10
    return f"鸡兔{head}只，{leg}条腿，如果增加5只鸡，共几条腿？", f"原兔{rabbit}鸡{chicken}，加5鸡后{leg_change}条腿"

def jitu_9():
    head = ri(10, 25)
    leg = ri(30, 80)
    rabbit = (leg - head * 2) // 2
    chicken = head - rabbit
    return f"鸡兔{head}只，{leg}条腿，鸡的腿数是多少？", f"鸡{chicken}只，鸡腿={chicken*2}条"

def jitu_10():
    head = ri(15, 30)
    leg = ri(40, 100)
    rabbit = (leg - head * 2) // 2
    chicken = head - rabbit
    return f"鸡兔{head}只，{leg}条腿，兔占总数的几分之几？", f"兔{rabbit}只，占{rabbit}/{head}"

# 9. 植树问题
def zhishu_1():
    length = random.choice([100, 150, 200, 250])
    interval = random.choice([5, 10, 20, 25])
    while length % interval != 0:
        length = random.choice([100, 150, 200, 250])
    seg = length // interval
    trees = seg + 1
    return f"路长{length}米，每隔{interval}米植一棵，两端都植，需几棵？", f"间隔{seg}个，树={seg+1}={trees}棵"

def zhishu_2():
    length = random.choice([100, 150, 200])
    interval = random.choice([5, 10, 20])
    while length % interval != 0:
        length = random.choice([100, 150, 200])
    seg = length // interval
    trees = seg
    return f"路长{length}米，每隔{interval}米植一棵，只植一端，需几棵？", f"间隔{seg}个，树={seg}={trees}棵"

def zhishu_3():
    length = random.choice([100, 150, 200])
    interval = random.choice([5, 10, 20])
    while length % interval != 0:
        length = random.choice([100, 150, 200])
    seg = length // interval
    trees = seg - 1
    return f"路长{length}米，每隔{interval}米植一棵，两端不植，需几棵？", f"间隔{seg}个，树={seg-1}={trees}棵"

def zhishu_4():
    trees = random.choice([20, 25, 30, 40])
    interval = random.choice([5, 10, 15])
    length = (trees - 1) * interval
    return f"植{trees}棵树（两端都植），间隔{interval}米，路长多少？", f"间隔{trees-1}个，路长={length}米"

def zhishu_5():
    perimeter = random.choice([200, 300, 400])
    interval = random.choice([5, 10, 20])
    while perimeter % interval != 0:
        perimeter = random.choice([200, 300, 400])
    trees = perimeter // interval
    return f"圆形池塘周长{perimeter}米，每隔{interval}米植一棵，需几棵？", f"树={perimeter}÷{interval}={trees}棵"

def zhishu_6():
    length = random.choice([100, 200, 300])
    interval = random.choice([5, 10])
    side = 2
    seg = length // interval
    trees_one = seg + 1
    trees_total = trees_one * side
    return f"路长{length}米，每隔{interval}米植一棵，两侧都植，共几棵？", f"一侧{trees_one}棵，两侧{trees_total}棵"

def zhishu_7():
    trees = random.choice([30, 40, 50])
    interval = random.choice([4, 5, 8])
    length = (trees - 1) * interval
    new_interval = interval + 2
    new_trees = length // new_interval + 1
    while new_trees * new_interval - new_interval + new_interval != length + new_interval:
        new_interval = interval + 2
        new_trees = length // new_interval + 1
    return f"原植{trees}棵，间隔{interval}米。现改为间隔{new_interval}米，需几棵？", f"路长{length}米，新间隔{new_interval}米，需{new_trees}棵"

def zhishu_8():
    floor = random.choice([5, 6, 8])
    steps_per_floor = random.choice([15, 18, 20])
    total_steps = (floor - 1) * steps_per_floor
    return f"从1楼到{floor}楼，每层{steps_per_floor}级台阶，共几级？", f"共{floor-1}层，{total_steps}级台阶"

def zhishu_9():
    length = random.choice([100, 200])
    interval = random.choice([5, 10])
    seg = length // interval
    trees = seg + 1
    cost_per_tree = random.choice([20, 25, 30])
    total_cost = trees * cost_per_tree
    return f"路长{length}米，每隔{interval}米植树，每棵{cost_per_tree}元，总费用？", f"树{trees}棵，费用={total_cost}元"

def zhishu_10():
    trees = random.choice([20, 30, 40])
    interval = random.choice([5, 10])
    length = (trees - 1) * interval
    return f"植{trees}棵树，间隔{interval}米，路长多少？如果改植{trees+5}棵，间隔变为多少？", f"路长{length}米，新间隔={length}÷{trees+5-1}米"

# 10. 分数百分数
def fenshu_1():
    total = random.choice([100, 200, 300, 400, 500])
    frac = random.choice([(1,2), (1,3), (1,4), (1,5), (2,3), (3,4)])
    part = total * frac[0] // frac[1]
    return f"书共{total}页，看了{frac[0]}/{frac[1]}，看了几页？", f"已看={total}×{frac[0]}/{frac[1]}={part}页"

def fenshu_2():
    part = random.choice([50, 60, 80, 100])
    frac = random.choice([(1,2), (1,3), (1,4), (1,5)])
    total = part * frac[1] // frac[0]
    return f"看了{part}页，占全书{frac[0]}/{frac[1]}，全书几页？", f"全书={part}÷{frac[0]}/{frac[1]}={total}页"

def fenshu_3():
    total = random.choice([100, 200, 300])
    frac1 = random.choice([(1,3), (1,4), (1,5)])
    frac2 = random.choice([(1,4), (1,5), (1,6)])
    part1 = total * frac1[0] // frac1[1]
    part2 = total * frac2[0] // frac2[1]
    remain = total - part1 - part2
    return f"{total}吨货物，第一次运{frac1[0]}/{frac1[1]}，第二次运{frac2[0]}/{frac2[1]}，剩几吨？", f"第一次{part1}吨，第二次{part2}吨，剩{remain}吨"

def fenshu_4():
    original = random.choice([100, 200, 300])
    pct = random.choice([10, 20, 25])
    increase = original * pct // 100
    new_val = original + increase
    return f"原{original}，增加{pct}%后是多少？", f"增加{increase}，现{new_val}"

def fenshu_5():
    original = random.choice([200, 300, 400])
    pct = random.choice([10, 20, 25])
    decrease = original * pct // 100
    new_val = original - decrease
    return f"原{original}，减少{pct}%后是多少？", f"减少{decrease}，现{new_val}"

def fenshu_6():
    pct = random.choice([20, 25, 50, 80])
    return f"{pct}%化成分数和小数各是多少？", f"分数={pct}/100={pct//gcd(pct,100)}/{100//gcd(pct,100)}，小数={pct/100}"

def fenshu_7():
    frac = random.choice([(1,2), (1,4), (3,4), (1,5), (2,5)])
    pct = frac[0] * 100 // frac[1]
    return f"{frac[0]}/{frac[1]}化成百分数是多少？", f"{frac[0]}/{frac[1]}={frac[0]*100//frac[1]}%={pct}%"

def fenshu_8():
    total = random.choice([100, 200, 300])
    pct1 = random.choice([20, 30, 40])
    pct2 = random.choice([30, 40, 50])
    part1 = total * pct1 // 100
    part2 = total * pct2 // 100
    diff = abs(part1 - part2)
    return f"{total}的{pct1}%比{pct2}%少几？", f"{pct1}%={part1}，{pct2}%={part2}，差{diff}"

def fenshu_9():
    original = random.choice([100, 200, 300])
    new_val = random.choice([120, 150, 200, 250])
    increase_pct = (new_val - original) * 100 // original
    while increase_pct * original != (new_val - original) * 100:
        new_val = random.choice([120, 150, 200, 250])
        increase_pct = (new_val - original) * 100 // original
    return f"从{original}增加到{new_val}，增幅百分之几？", f"增加{new_val-original}，增幅={increase_pct}%"

def fenshu_10():
    a, b = random.choice([(20, 30), (30, 50), (40, 60)])
    pct = a * 100 // b
    while pct * b != a * 100:
        a, b = random.choice([(20, 30), (30, 50), (40, 60)])
        pct = a * 100 // b
    return f"{a}是{b}的百分之几？", f"{a}÷{b}×100%={pct}%"

# 11. 几何应用
def geo_1():
    l, w = ri(5, 30), ri(3, 20)
    area = l * w
    return f"长方形长{l}cm，宽{w}cm，面积？", f"面积={l}×{w}={area}cm²"

def geo_2():
    l, w = ri(5, 30), ri(3, 20)
    perimeter = 2 * (l + w)
    return f"长方形长{l}cm，宽{w}cm，周长？", f"周长=({l}+{w})×2={perimeter}cm"

def geo_3():
    s = ri(3, 20)
    area = s * s
    perimeter = 4 * s
    return f"正方形边长{s}cm，周长和面积？", f"周长={s}×4={perimeter}cm，面积={s}×{s}={area}cm²"

def geo_4():
    base, h = ri(5, 30), ri(3, 20)
    area = base * h / 2
    return f"三角形底{base}cm，高{h}cm，面积？", f"面积={base}×{h}÷2={area}cm²"

def geo_5():
    base, h = ri(5, 30), ri(3, 20)
    area = base * h
    return f"平行四边形底{base}cm，高{h}cm，面积？", f"面积={base}×{h}={area}cm²"

def geo_6():
    a, b, h = ri(3, 15), ri(5, 20), ri(3, 15)
    area = (a + b) * h / 2
    return f"梯形上底{a}cm，下底{b}cm，高{h}cm，面积？", f"面积=({a}+{b})×{h}÷2={area}cm²"

def geo_7():
    r = ri(2, 10)
    c = round(2 * 3.14 * r, 2)
    s = round(3.14 * r * r, 2)
    return f"圆半径{r}cm，周长和面积？（π取3.14）", f"周长=2×3.14×{r}={c}cm，面积=3.14×{r}²={s}cm²"

def geo_8():
    l, w, h = ri(3, 15), ri(2, 10), ri(2, 10)
    volume = l * w * h
    return f"长方体长{l}cm，宽{w}cm，高{h}cm，体积？", f"体积={l}×{w}×{h}={volume}cm³"

def geo_9():
    a = ri(3, 10)
    volume = a * a * a
    return f"正方体棱长{a}cm，体积？", f"体积={a}×{a}×{a}={volume}cm³"

def geo_10():
    l, w = ri(5, 20), ri(3, 15)
    h = ri(2, 10)
    area_wall = 2 * (l * h + w * h)
    area_floor = l * w
    return f"房间长{l}m，宽{w}m，高{h}m。四壁面积？地面面积？", f"四壁={area_wall}m²，地面={area_floor}m²"

# 12. 平均数
def avg_1():
    nums = [ri(70, 95) for _ in range(5)]
    total = sum(nums)
    avg = total / 5
    return f"5次成绩：{'、'.join(map(str, nums))}，平均？", f"总分={total}，平均={avg}"

def avg_2():
    nums = [ri(60, 90) for _ in range(4)]
    target_avg = ri(75, 85)
    need = target_avg * 5 - sum(nums)
    return f"4次成绩{'、'.join(map(str, nums))}，想平均{target_avg}，第5次需几分？", f"需{need}分"

def avg_3():
    nums = [ri(70, 90) for _ in range(6)]
    sorted_nums = sorted(nums)
    median = (sorted_nums[2] + sorted_nums[3]) / 2
    return f"6个数：{'、'.join(map(str, nums))}，中位数？", f"排序后中位数={median}"

def avg_4():
    avg1, n1 = ri(75, 85), ri(3, 5)
    avg2, n2 = ri(80, 90), ri(3, 5)
    total = avg1 * n1 + avg2 * n2
    total_n = n1 + n2
    combined_avg = total / total_n
    return f"甲{ n1}人平均{avg1}分，乙{n2}人平均{avg2}分，总平均？", f"总分={total}，总人数={total_n}，平均={combined_avg}"

def avg_5():
    avg_old, n = ri(70, 80), ri(4, 6)
    new_score = ri(85, 95)
    total_old = avg_old * n
    new_avg = (total_old + new_score) / (n + 1)
    return f"原{n}人平均{avg_old}分，新加入{new_score}分，新平均？", f"新平均={new_avg}"

def avg_6():
    avg_old, n = ri(75, 85), ri(5, 7)
    leave_score = ri(60, 70)
    total_old = avg_old * n
    new_avg = round((total_old - leave_score) / (n - 1), 1)
    return f"原{n}人平均{avg_old}分，离开一人{leave_score}分，新平均？", f"新平均={new_avg}分"

def avg_7():
    speeds = [ri(40, 70) for _ in range(3)]
    avg_speed = round(sum(speeds) / 3, 1)
    return f"三段速度：{'、'.join(map(str, speeds))}km/h，平均速度？", f"平均={avg_speed}km/h"

def avg_8():
    prices = [ri(5, 20) for _ in range(4)]
    avg_price = sum(prices) / 4
    return f"4次价格：{'、'.join(map(str, prices))}元，平均？", f"平均={avg_price}元"

def avg_9():
    # 确保 avg 是整数，避免浮点精度问题
    avg = ri(20, 45) * 2  # even number, divisible by 2 for clean /5
    nums_total = avg * 5  # guaranteed divisible by 5
    new_avg = avg + 3
    need = new_avg * 6 - nums_total
    return f"5个数平均{avg}，加一个数后平均变{new_avg}，加的数？", f"加的数={need}"

def avg_10():
    heights = [ri(130, 160) for _ in range(6)]
    max_h, min_h = max(heights), min(heights)
    diff = max_h - min_h
    return f"6人身高：{'、'.join(map(str, heights))}cm，最高与最矮差多少？", f"差{diff}cm"

# 13. 还原问题
def huanyuan_1():
    original = ri(50, 200)
    add, sub = ri(10, 50), ri(5, 30)
    result = original + add - sub
    return f"某数加{add}减{sub}得{result}，原数？", f"原数={original}"

def huanyuan_2():
    original = ri(30, 100)
    mul = ri(2, 5)
    result = original * mul
    return f"某数乘{mul}得{result}，原数？", f"原数={original}"

def huanyuan_3():
    original = ri(100, 300)
    div = ri(2, 5)
    while original % div != 0:
        original = ri(100, 300)
    result = original // div
    return f"某数除以{div}得{result}，原数？", f"原数={original}"

def huanyuan_4():
    original = ri(50, 150)
    add, mul = ri(10, 30), ri(2, 4)
    result = (original + add) * mul
    return f"某数加{add}再乘{mul}得{result}，原数？", f"原数={original}"

def huanyuan_5():
    original = ri(100, 300)
    sub, div = ri(20, 50), ri(2, 4)
    while (original - sub) % div != 0:
        original = ri(100, 300)
    result = (original - sub) // div
    return f"某数减{sub}再除以{div}得{result}，原数？", f"原数={original}"

def huanyuan_6():
    original = ri(60, 200)
    steps = [(ri(1,3), ri(10,30)) for _ in range(3)]
    result = original
    for op, val in steps:
        if op == 1:
            result += val
        else:
            result -= val
    return f"某数经过+{steps[0][1]}、-{steps[1][1]}、+{steps[2][1]}得{result}，原数？", f"原数={original}"

def huanyuan_7():
    original = ri(40, 120)
    half = original // 2
    while half * 2 != original:
        original = ri(40, 120)
        half = original // 2
    return f"某数一半是{half}，原数？", f"原数={original}"

def huanyuan_8():
    original = ri(50, 150)
    add1, add2 = ri(10, 30), ri(10, 30)
    result = original + add1 + add2
    return f"某数加{add1}再加{add2}得{result}，原数？", f"原数={original}"

def huanyuan_9():
    original = ri(80, 200)
    mul, add = ri(2, 4), ri(10, 30)
    result = original * mul + add
    return f"某数乘{mul}加{add}得{result}，原数？", f"原数={original}"

def huanyuan_10():
    original = ri(100, 300)
    sub1, sub2 = ri(20, 50), ri(20, 50)
    result = original - sub1 - sub2
    return f"某数减{sub1}再减{sub2}得{result}，原数？", f"原数={original}"

# 类别配置：(大类名, 子类型函数列表)
WORD_CATEGORIES = [
    ("行程问题", [xingcheng_1, xingcheng_2, xingcheng_3, xingcheng_4, xingcheng_5,
                 xingcheng_6, xingcheng_7, xingcheng_8, xingcheng_9, xingcheng_10]),
    ("工程问题", [gongcheng_1, gongcheng_2, gongcheng_3, gongcheng_4, gongcheng_5,
                 gongcheng_6, gongcheng_7, gongcheng_8, gongcheng_9, gongcheng_10]),
    ("浓度问题", [nongdu_1, nongdu_2, nongdu_3, nongdu_4, nongdu_5,
                 nongdu_6, nongdu_7, nongdu_8, nongdu_9, nongdu_10]),
    ("利润折扣", [lirun_1, lirun_2, lirun_3, lirun_4, lirun_5,
                 lirun_6, lirun_7, lirun_8, lirun_9, lirun_10]),
    ("比例问题", [bili_1, bili_2, bili_3, bili_4, bili_5,
                 bili_6, bili_7, bili_8, bili_9, bili_10]),
    ("年龄问题", [age_1, age_2, age_3, age_4, age_5,
                 age_6, age_7, age_8, age_9, age_10]),
    ("鸡兔同笼", [jitu_1, jitu_2, jitu_3, jitu_4, jitu_5,
                 jitu_6, jitu_7, jitu_8, jitu_9, jitu_10]),
    ("植树问题", [zhishu_1, zhishu_2, zhishu_3, zhishu_4, zhishu_5,
                 zhishu_6, zhishu_7, zhishu_8, zhishu_9, zhishu_10]),
    ("分数百分数", [fenshu_1, fenshu_2, fenshu_3, fenshu_4, fenshu_5,
                   fenshu_6, fenshu_7, fenshu_8, fenshu_9, fenshu_10]),
    ("几何应用", [geo_1, geo_2, geo_3, geo_4, geo_5,
                 geo_6, geo_7, geo_8, geo_9, geo_10]),
    ("平均数", [avg_1, avg_2, avg_3, avg_4, avg_5,
               avg_6, avg_7, avg_8, avg_9, avg_10]),
    ("还原问题", [huanyuan_1, huanyuan_2, huanyuan_3, huanyuan_4, huanyuan_5,
                 huanyuan_6, huanyuan_7, huanyuan_8, huanyuan_9, huanyuan_10]),
]

def generate_word_problems(total=500):
    """生成应用题"""
    problems, answers = [], []
    
    # 计算每个大类的题目数
    num_categories = len(WORD_CATEGORIES)
    per_category = total // num_categories
    
    for cat_name, sub_funcs in WORD_CATEGORIES:
        # 每个子类型生成的题目数
        per_subtype = per_category // len(sub_funcs)
        if per_subtype < 1:
            per_subtype = 1
        
        for func in sub_funcs:
            for _ in range(per_subtype):
                try:
                    p, a = func()
                    problems.append(p)
                    answers.append(a)
                except:
                    pass
    
    # 打乱
    combined = list(zip(problems, answers))
    random.shuffle(combined)
    problems = [p for p, a in combined]
    answers = [a for p, a in combined]
    
    return problems[:total], answers[:total]

def create_doc(problems, answers, title, filename):
    """生成Word文档"""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(title)
    run.font.size = Pt(18)
    run.bold = True
    
    for idx, (p, a) in enumerate(zip(problems, answers), 1):
        para = doc.add_paragraph()
        para.space_before = Pt(4)
        run = para.add_run(f"{idx}. {p}")
        run.font.size = Pt(11)
        if "答案" in title:
            run = para.add_run(f"\n{a}")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(100, 100, 100)
        else:
            para.add_run("\n\n")
    
    doc.save(filename)
    return filename

def main():
    total = 500
    if len(sys.argv) > 1:
        total = int(sys.argv[1])
    
    print(f"生成{total}道应用题...")
    problems, answers = generate_word_problems(total)
    
    ts = time.strftime("%Y%m%d_%H%M%S")
    
    q_file = create_doc(problems, answers, f"小学数学应用题{total}道（题目卷）", f"应用题{total}道_题目卷_{ts}.docx")
    a_file = create_doc(problems, answers, f"小学数学应用题{total}道（答案卷）", f"应用题{total}道_答案卷_{ts}.docx")
    
    print(f"题目卷: {q_file}")
    print(f"答案卷: {a_file}")
    print("完成！")

if __name__ == "__main__":
    main()
