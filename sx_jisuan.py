"""
小学数学计算题生成器（可配置总数）
覆盖25个分类，按权重自动分配题数。
每类最少1题，最多不超过最少分类的5倍。
用法：python gen_math_2000.py [总题数]
      默认2000题
输出：题目卷 + 答案卷 两个 Word 文档
依赖：pip install python-docx
"""

import sys, random
import time
from fractions import Fraction
from math import gcd

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("请先安装 python-docx：pip install python-docx")
    exit(1)

random.seed(42)

# ═══════════════════════════════════════
# 配置
# ═══════════════════════════════════════
TOTAL = 200  # 默认总题数，可通过命令行参数覆盖

# 25个分类：(名称, 生成函数, 权重)
# 权重决定分配比例，原始2000题时的数量即为权重

# ═══════════════════════════════════════
# 各分类生成器（每个 yield 单道题）
# ═══════════════════════════════════════

def gen_int_add():
    """1. 整数加法"""
    while True:
        i = random.randint(0, 2)
        if i == 0:
            a, b = random.randint(10, 999), random.randint(10, 999)
        elif i == 1:
            a, b = random.randint(100, 9999), random.randint(100, 9999)
        else:
            a, b = random.randint(1000, 99999), random.randint(1000, 99999)
        yield f"{a} + {b} =", f"{a} + {b} = {a + b}"

def gen_int_sub():
    """2. 整数减法"""
    while True:
        i = random.randint(0, 2)
        if i == 0:
            a = random.randint(100, 999); b = random.randint(10, a)
        elif i == 1:
            a = random.randint(1000, 9999); b = random.randint(100, a)
        else:
            a = random.randint(10000, 99999); b = random.randint(1000, a)
        yield f"{a} - {b} =", f"{a} - {b} = {a - b}"

def gen_int_mul():
    """3. 整数乘法"""
    while True:
        i = random.randint(0, 3)
        if i == 0:
            a, b = random.randint(2, 9), random.randint(10, 99)
        elif i == 1:
            a, b = random.randint(10, 99), random.randint(10, 99)
        elif i == 2:
            a, b = random.randint(100, 999), random.randint(2, 9)
        else:
            a, b = random.randint(100, 499), random.randint(10, 99)
        yield f"{a} × {b} =", f"{a} × {b} = {a * b}"

def gen_int_div():
    """4. 整数除法"""
    while True:
        i = random.randint(0, 2)
        if i == 0:
            b = random.randint(2, 9); q = random.randint(10, 99)
        elif i == 1:
            b = random.randint(2, 9); q = random.randint(100, 999)
        else:
            b = random.randint(10, 50); q = random.randint(10, 99)
        a = b * q
        yield f"{a} ÷ {b} =", f"{a} ÷ {b} = {q}"

def gen_remainder_div():
    """5. 带余除法"""
    while True:
        b = random.randint(2, 9)
        q = random.randint(10, 99)
        r = random.randint(1, b - 1)
        a = b * q + r
        yield f"{a} ÷ {b} =", f"{a} ÷ {b} = {q}……{r}"

def gen_dec_add():
    """6. 小数加法"""
    while True:
        dp = random.choice([1, 2])
        a = round(random.uniform(0.1, 99.9) if dp == 1 else random.uniform(0.01, 99.99), dp)
        b = round(random.uniform(0.1, 99.9) if dp == 1 else random.uniform(0.01, 99.99), dp)
        ans = round(a + b, 2)
        yield f"{a} + {b} =", f"{a} + {b} = {ans}"

def gen_dec_sub():
    """7. 小数减法"""
    while True:
        dp = random.choice([1, 2])
        a = round(random.uniform(10, 999.9) if dp == 1 else random.uniform(10, 999.99), dp)
        b = round(random.uniform(0.1, a) if dp == 1 else random.uniform(0.01, a), dp)
        ans = round(a - b, 2)
        yield f"{a} - {b} =", f"{a} - {b} = {ans}"

def gen_dec_mul():
    """8. 小数乘法"""
    while True:
        i = random.randint(0, 2)
        if i == 0:
            a = round(random.uniform(0.1, 9.9), 1); b = random.randint(2, 9)
        elif i == 1:
            a = round(random.uniform(0.1, 9.9), 1); b = round(random.uniform(0.1, 9.9), 1)
        else:
            a = round(random.uniform(0.01, 9.99), 2); b = round(random.uniform(0.1, 9.9), 1)
        ans = round(a * b, 4)
        yield f"{a} × {b} =", f"{a} × {b} = {ans:g}"

def gen_dec_div():
    """9. 小数除法"""
    while True:
        i = random.randint(0, 1)
        if i == 0:
            b = random.randint(2, 9)
            q = round(random.uniform(0.1, 99.9), 1)
            a = round(b * q, 1)
        else:
            b = round(random.uniform(0.1, 9.9), 1)
            q = round(random.uniform(1, 99), 1)
            a = round(b * q, 2)
        ans = round(a / b, 2)
        yield f"{a} ÷ {b} =", f"{a} ÷ {b} = {ans}"

def gen_frac_add():
    """10. 分数加法"""
    while True:
        d1 = random.choice([2, 3, 4, 5, 6, 7, 8, 9, 10, 12])
        d2 = random.choice([2, 3, 4, 5, 6, 7, 8, 9, 10, 12])
        n1 = random.randint(1, d1 - 1); n2 = random.randint(1, d2 - 1)
        result = Fraction(n1, d1) + Fraction(n2, d2)
        if result.denominator == 1:
            ans_str = str(result.numerator)
        elif result.numerator > result.denominator:
            w = result.numerator // result.denominator
            r = result.numerator % result.denominator
            ans_str = f"{result.numerator}/{result.denominator} = {w}又{r}/{result.denominator}"
        else:
            ans_str = f"{result.numerator}/{result.denominator}"
        yield f"{n1}/{d1} + {n2}/{d2} =", f"{n1}/{d1} + {n2}/{d2} = {ans_str}"

def gen_frac_sub():
    """11. 分数减法"""
    while True:
        d1 = random.choice([2, 3, 4, 5, 6, 8, 9, 10, 12])
        d2 = random.choice([2, 3, 4, 5, 6, 8, 9, 10, 12])
        n1 = random.randint(1, d1 - 1); n2 = random.randint(1, d2 - 1)
        f1, f2 = Fraction(n1, d1), Fraction(n2, d2)
        if f1 < f2:
            f1, f2 = f2, f1
            n1, d1, n2, d2 = f1.numerator, f1.denominator, f2.numerator, f2.denominator
        result = f1 - f2
        if result == 0:
            ans_str = "0"
        elif result.denominator == 1:
            ans_str = str(result.numerator)
        else:
            ans_str = f"{result.numerator}/{result.denominator}"
        yield f"{n1}/{d1} - {n2}/{d2} =", f"{n1}/{d1} - {n2}/{d2} = {ans_str}"

def gen_frac_mul():
    """12. 分数乘法"""
    while True:
        d1 = random.choice([2, 3, 4, 5, 6, 7, 8, 9, 10])
        d2 = random.choice([2, 3, 4, 5, 6, 7, 8, 9, 10])
        n1 = random.randint(1, d1); n2 = random.randint(1, d2)
        result = Fraction(n1, d1) * Fraction(n2, d2)
        if result.denominator == 1:
            ans_str = str(result.numerator)
        else:
            ans_str = f"{result.numerator}/{result.denominator}"
        yield f"{n1}/{d1} × {n2}/{d2} =", f"{n1}/{d1} × {n2}/{d2} = {ans_str}"

def gen_frac_div():
    """13. 分数除法"""
    while True:
        d1 = random.choice([2, 3, 4, 5, 6, 8, 10])
        d2 = random.choice([2, 3, 4, 5, 6, 8, 10])
        n1 = random.randint(1, d1 - 1); n2 = random.randint(1, d2 - 1)
        result = Fraction(n1, d1) / Fraction(n2, d2)
        if result.denominator == 1:
            ans_str = str(result.numerator)
        elif result.numerator > result.denominator:
            w = result.numerator // result.denominator
            r = result.numerator % result.denominator
            ans_str = f"{result.numerator}/{result.denominator} = {w}又{r}/{result.denominator}"
        else:
            ans_str = f"{result.numerator}/{result.denominator}"
        yield f"{n1}/{d1} ÷ {n2}/{d2} =", f"{n1}/{d1} ÷ {n2}/{d2} = {ans_str}"

def gen_percent():
    """14. 百分数"""
    while True:
        i = random.randint(0, 2)
        if i == 0:
            a = random.randint(10, 200); b = random.randint(20, 500)
            pct = round(a / b * 100, 1)
            yield f"{a} 是 {b} 的百分之几？", f"{a} ÷ {b} × 100% = {pct}%"
        elif i == 1:
            a = random.randint(50, 500)
            pct = random.choice([10, 15, 20, 25, 30, 40, 50, 60, 75, 80])
            result = a * pct / 100
            yield f"{a} 的 {pct}% 是多少？", f"{a} × {pct}% = {a} × {pct/100} = {result}"
        else:
            pct = random.randint(1, 20) * 5
            yield f"把 {pct}% 化成小数", f"{pct}% = {pct / 100}"

def gen_ratio():
    """15. 比"""
    while True:
        if random.random() < 0.5:
            a = random.randint(2, 50); b = random.randint(2, 50)
            g = gcd(a, b)
            yield f"化简比 {a}:{b}", f"{a}:{b} = {a//g}:{b//g}"
        else:
            a = random.randint(2, 20); b = random.randint(2, 20)
            total = random.randint(100, 500)
            g = gcd(a, b); sa, sb = a // g, b // g
            pa = total * sa // (sa + sb); pb = total - pa
            yield f"按 {a}:{b} 分配 {total}，两部分各是多少？", f"总份数 = {sa+sb}，一部分 = {pa}，另一部分 = {pb}"

def gen_mixed_ops():
    """16. 四则混合运算"""
    while True:
        i = random.randint(0, 4)
        if i == 0:
            a = random.randint(1, 100); b = random.randint(2, 9); c = random.randint(2, 9)
            yield f"{a} + {b} × {c} =", f"{a} + {b} × {c} = {a} + {b*c} = {a + b*c}"
        elif i == 1:
            a = random.randint(1, 50); b = random.randint(1, 50); c = random.randint(2, 9)
            yield f"({a} + {b}) × {c} =", f"({a} + {b}) × {c} = {a+b} × {c} = {(a+b)*c}"
        elif i == 2:
            d = random.randint(2, 9); cp = random.randint(2, 9); c = d * cp
            a = random.randint(2, 20); b = random.randint(2, 20)
            yield f"{a} × {b} - {c} ÷ {d} =", f"{a} × {b} - {c} ÷ {d} = {a*b} - {cp} = {a*b - cp}"
        elif i == 3:
            a = random.randint(100, 500); b = random.randint(10, 99)
            c = random.randint(2, 9); d = random.randint(2, 9)
            yield f"{a} - {b} + {c} × {d} =", f"{a} - {b} + {c} × {d} = {a} - {b} + {c*d} = {a - b + c*d}"
        else:
            c = random.randint(2, 9)
            ab = c * random.randint(2, 20)
            a = random.randint(1, ab - 1); b = ab - a
            d = random.randint(1, ab // c - 1)
            yield f"({a} + {b}) ÷ {c} - {d} =", f"({a} + {b}) ÷ {c} - {d} = {ab} ÷ {c} - {d} = {ab//c} - {d} = {ab//c - d}"

def gen_equation():
    """17. 解方程"""
    while True:
        i = random.randint(0, 3)
        if i == 0:
            x = random.randint(1, 50); a = random.randint(2, 9); b = random.randint(1, 50)
            c = a * x + b
            yield f"解方程：{a}x + {b} = {c}", f"解：{a}x = {c} - {b} = {c-b}，x = {c-b} ÷ {a} = {x}"
        elif i == 1:
            x = random.randint(2, 50); a = random.randint(2, 9)
            b = random.randint(1, a * x - 1); c = a * x - b
            yield f"解方程：{a}x - {b} = {c}", f"解：{a}x = {c} + {b} = {c+b}，x = {c+b} ÷ {a} = {x}"
        elif i == 2:
            a = random.randint(2, 9); x = a * random.randint(1, 20)
            b = random.randint(1, 30); c = x // a + b
            yield f"解方程：x ÷ {a} + {b} = {c}", f"解：x ÷ {a} = {c} - {b} = {c-b}，x = {c-b} × {a} = {(c-b)*a}"
        else:
            x = random.randint(1, 30); a = random.randint(2, 6)
            b = random.randint(1, 20); c = a * (x + b)
            yield f"解方程：{a}(x + {b}) = {c}", f"解：x + {b} = {c} ÷ {a} = {c//a}，x = {c//a} - {b} = {c//a - b}"

def gen_unit_convert():
    """18. 单位换算"""
    unit_data = [
        ("长度", [("千米", "米", 1000), ("米", "厘米", 100), ("厘米", "毫米", 10),
                  ("米", "分米", 10), ("分米", "厘米", 10)]),
        ("面积", [("平方千米", "公顷", 100), ("公顷", "平方米", 10000),
                  ("平方米", "平方分米", 100), ("平方分米", "平方厘米", 100)]),
        ("体积", [("立方米", "立方分米", 1000), ("立方分米", "立方厘米", 1000), ("升", "毫升", 1000)]),
        ("重量", [("吨", "千克", 1000), ("千克", "克", 1000)]),
        ("时间", [("时", "分", 60), ("分", "秒", 60)]),
    ]
    while True:
        cat_name, units = random.choice(unit_data)
        uf, ut, factor = random.choice(units)
        val = random.randint(1, 20) if factor >= 100 else random.randint(2, 50)
        result = val * factor
        yield f"{val}{uf} = ____{ut}", f"{val}{uf} = {result}{ut}"

def gen_geometry():
    """19. 几何计算"""
    while True:
        i = random.randint(0, 7)
        if i == 0:
            l = random.randint(3, 50); w = random.randint(2, 30)
            yield f"长方形长{l}cm，宽{w}cm，周长是多少？", f"周长 = ({l} + {w}) × 2 = {(l+w)*2}cm"
        elif i == 1:
            l = random.randint(3, 30); w = random.randint(2, 20)
            yield f"长方形长{l}cm，宽{w}cm，面积是多少？", f"面积 = {l} × {w} = {l*w}cm²"
        elif i == 2:
            s = random.randint(2, 30)
            yield f"正方形边长{s}cm，周长和面积各是多少？", f"周长 = {s} × 4 = {s*4}cm，面积 = {s} × {s} = {s*s}cm²"
        elif i == 3:
            base = random.randint(4, 30); h = random.randint(2, 20)
            area = base * h / 2
            yield f"三角形底{base}cm，高{h}cm，面积是多少？", f"面积 = {base} × {h} ÷ 2 = {area}cm²"
        elif i == 4:
            base = random.randint(4, 30); h = random.randint(2, 20)
            yield f"平行四边形底{base}cm，高{h}cm，面积是多少？", f"面积 = {base} × {h} = {base*h}cm²"
        elif i == 5:
            a_s = random.randint(3, 20); b_s = random.randint(3, 20); h = random.randint(2, 15)
            area = (a_s + b_s) * h / 2
            yield f"梯形上底{a_s}cm，下底{b_s}cm，高{h}cm，面积是多少？", f"面积 = ({a_s} + {b_s}) × {h} ÷ 2 = {area}cm²"
        elif i == 6:
            r = random.randint(1, 10)
            c = round(2 * 3.14 * r, 2); s = round(3.14 * r * r, 2)
            yield f"圆半径{r}cm，周长和面积各是多少？（π取3.14）", f"周长 = 2 × 3.14 × {r} = {c}cm，面积 = 3.14 × {r}² = {s}cm²"
        else:
            l = random.randint(2, 15); w = random.randint(2, 10); h = random.randint(2, 10)
            yield f"长方体长{l}cm，宽{w}cm，高{h}cm，体积是多少？", f"体积 = {l} × {w} × {h} = {l*w*h}cm³"

def gen_statistics():
    """20. 统计"""
    while True:
        n = random.choice([4, 5, 6, 7, 9])
        nums = [random.randint(50, 120) for _ in range(n)]
        nums_str = "、".join(str(x) for x in nums)
        if random.random() < 0.5:
            avg = sum(nums) / n
            yield f"求以下数据的平均数：{nums_str}", f"平均数 = ({' + '.join(str(x) for x in nums)}) ÷ {n} = {sum(nums)} ÷ {n} = {avg}"
        else:
            sorted_nums = sorted(nums)
            median = sorted_nums[n // 2]
            yield f"求以下数据的中位数：{nums_str}", f"排序后：{'、'.join(str(x) for x in sorted_nums)}，中位数 = {median}"

def gen_scale():
    """21. 比例尺"""
    while True:
        scale = random.choice([100, 200, 500, 1000, 5000, 10000, 50000])
        map_dist = random.randint(2, 20)
        real_dist = map_dist * scale
        if real_dist >= 100000:
            yield f"比例尺1:{scale}，图上距离{map_dist}cm，实际距离多少千米？", f"实际距离 = {map_dist} × {scale} = {real_dist}cm = {real_dist/100000}千米"
        else:
            yield f"比例尺1:{scale}，图上距离{map_dist}cm，实际距离多少米？", f"实际距离 = {map_dist} × {scale} = {real_dist}cm = {real_dist/100}米"

def gen_interest_discount():
    """22. 利息折扣"""
    while True:
        if random.random() < 0.5:
            price = random.randint(50, 500)
            discount = random.choice([80, 75, 70, 60, 85, 90, 50])
            final_price = price * discount / 100
            yield f"商品原价{price}元，打{discount//10}折出售，现价多少元？", f"现价 = {price} × {discount/100} = {final_price}元"
        else:
            principal = random.randint(1000, 10000)
            rate = random.choice([2, 3, 4, 5])
            years = random.randint(1, 3)
            interest = principal * rate / 100 * years
            yield f"本金{principal}元，年利率{rate}%，存{years}年，利息是多少？", f"利息 = {principal} × {rate/100} × {years} = {interest}元"

def gen_negative():
    """23. 负数/温度"""
    while True:
        if random.random() < 0.5:
            t1 = random.randint(-20, -1); t2 = random.randint(1, 15)
            diff = t2 - t1
            yield f"某地气温从零下{abs(t1)}°C上升到{t2}°C，温差是多少？", f"温差 = {t2} - ({t1}) = {t2} + {abs(t1)} = {diff}°C"
        else:
            a = random.randint(-30, -5); b = random.randint(-10, 10)
            s = a + b
            yield f"({a}) + ({b}) =", f"({a}) + ({b}) = {s}"

def gen_shortcut():
    """24. 简便运算"""
    while True:
        i = random.randint(0, 3)
        if i == 0:
            a = random.randint(2, 9); b = random.randint(10, 50); c = random.randint(10, 50)
            yield f"用简便方法计算：{a} × {b} + {a} × {c}", f"{a} × {b} + {a} × {c} = {a} × ({b} + {c}) = {a} × {b+c} = {a*(b+c)}"
        elif i == 1:
            a = random.randint(2, 8)
            b = random.choice([25, 125, 50])
            c = random.choice([4, 8, 2])
            yield f"用简便方法计算：{a} × {b} × {c}", f"{a} × {b} × {c} = {a} × {b*c} = {a*b*c}"
        elif i == 2:
            a = random.randint(500, 1000); b = random.randint(100, 300); c = random.randint(100, 300)
            yield f"用简便方法计算：{a} - {b} - {c}", f"{a} - {b} - {c} = {a} - ({b} + {c}) = {a} - {b+c} = {a-b-c}"
        else:
            a = random.choice([99, 101, 198, 202, 98, 102])
            b = random.randint(2, 20)
            if a < 100:
                ans = 100 * b - (100 - a) * b
                yield f"用简便方法计算：{a} × {b}", f"{a} × {b} = (100-{100-a}) × {b} = {100*b} - {(100-a)*b} = {ans}"
            else:
                ans = 100 * b + (a - 100) * b
                yield f"用简便方法计算：{a} × {b}", f"{a} × {b} = (100+{a-100}) × {b} = {100*b} + {(a-100)*b} = {ans}"

def gen_datetime():
    """25. 日期时间"""
    day_names = {1: "一", 2: "二", 3: "三", 4: "四", 5: "五", 6: "六", 7: "日"}
    while True:
        i = random.randint(0, 2)
        if i == 0:
            h1 = random.randint(6, 16); m1 = random.randint(0, 59)
            dur_h = random.randint(1, 5); dur_m = random.randint(10, 50)
            h2 = h1 + dur_h; m2 = m1 + dur_m
            if m2 >= 60: h2 += 1; m2 -= 60
            yield f"从 {h1}:{m1:02d} 经过 {dur_h}小时{dur_m}分钟，是几时几分？", f"{h1}:{m1:02d} + {dur_h}小时{dur_m}分钟 = {h2}:{m2:02d}"
        elif i == 1:
            year = random.randint(1900, 2100)
            is_leap = (year % 4 == 0 and year % 100 != 0) or (year % 400 == 0)
            yield f"{year}年是平年还是闰年？全年有多少天？", f"{year}年{'是闰年' if is_leap else '是平年'}，全年有{'366' if is_leap else '365'}天"
        else:
            start_day = random.randint(1, 7)
            days_later = random.randint(7, 100)
            end_day = (start_day + days_later - 1) % 7 + 1
            yield f"今天是星期{day_names[start_day]}，{days_later}天后是星期几？", f"{days_later} ÷ 7 = {days_later//7}……{days_later%7}，从星期{day_names[start_day]}往后推{days_later%7}天，是星期{day_names[end_day]}"


# ═══════════════════════════════════════
# 分类注册表
# ═══════════════════════════════════════
CATEGORIES = [
    ("一、整数加法",       gen_int_add,          150),
    ("二、整数减法",       gen_int_sub,          150),
    ("三、整数乘法",       gen_int_mul,          150),
    ("四、整数除法",       gen_int_div,          120),
    ("五、带余除法",       gen_remainder_div,     80),
    ("六、小数加法",       gen_dec_add,          100),
    ("七、小数减法",       gen_dec_sub,          100),
    ("八、小数乘法",       gen_dec_mul,          100),
    ("九、小数除法",       gen_dec_div,           80),
    ("十、分数加法",       gen_frac_add,          80),
    ("十一、分数减法",     gen_frac_sub,          60),
    ("十二、分数乘法",     gen_frac_mul,          50),
    ("十三、分数除法",     gen_frac_div,          50),
    ("十四、百分数计算",   gen_percent,           80),
    ("十五、比的计算",     gen_ratio,             60),
    ("十六、四则混合运算", gen_mixed_ops,        150),
    ("十七、解方程",       gen_equation,          80),
    ("十八、单位换算",     gen_unit_convert,     100),
    ("十九、几何计算",     gen_geometry,         100),
    ("二十、统计计算",     gen_statistics,        50),
    ("二十一、比例尺计算", gen_scale,             30),
    ("二十二、利息和折扣", gen_interest_discount,  40),
    ("二十三、负数与温度", gen_negative,          30),
    ("二十四、简便运算",   gen_shortcut,         100),
    ("二十五、日期与时间", gen_datetime,          50),
]

NUM_CATS = len(CATEGORIES)  # 25


# ═══════════════════════════════════════
# 动态分配算法
# ═══════════════════════════════════════
def allocate_counts(total, categories):
    """
    根据总题数和权重，自动分配每个分类的题数。
    约束：每类最少1题，最多不超过最少分类的5倍。
    """
    n = len(categories)
    weights = [c[2] for c in categories]
    total_weight = sum(weights)

    if total < n:
        total = n

    # 第一步：按比例分配，每类至少1题
    alloc = [max(1, round(w / total_weight * total)) for w in weights]

    # 第二步：调整总和到 total
    diff = total - sum(alloc)
    if diff > 0:
        # 需要增加：按权重从大到小依次加
        indices = sorted(range(n), key=lambda i: weights[i], reverse=True)
        idx = 0
        while diff > 0:
            alloc[indices[idx % n]] += 1
            diff -= 1
            idx += 1
    elif diff < 0:
        # 需要减少：从最大的开始减（但不能低于1）
        diff = -diff
        indices = sorted(range(n), key=lambda i: alloc[i], reverse=True)
        idx = 0
        while diff > 0:
            i = indices[idx % n]
            if alloc[i] > 1:
                alloc[i] -= 1
                diff -= 1
            idx += 1
            if idx > n * 1000:  # 安全退出
                break

    # 第三步：检查 max <= 5 * min，不满足则压缩高值、提升低值
    for _ in range(100):  # 最多迭代100次
        mn = min(alloc)
        mx = max(alloc)
        if mx <= 5 * mn:
            break
        limit = 5 * mn
        # 把超过 limit 的减到 limit
        excess = 0
        for i in range(n):
            if alloc[i] > limit:
                excess += alloc[i] - limit
                alloc[i] = limit
        # 把 excess 分给低于 limit 的
        if excess > 0:
            indices = sorted(range(n), key=lambda i: alloc[i])
            idx = 0
            while excess > 0:
                i = indices[idx % n]
                if alloc[i] < limit:
                    alloc[i] += 1
                    excess -= 1
                idx += 1
                if idx > n * 100:
                    break

    # 最终微调确保总和精确
    diff = total - sum(alloc)
    if diff > 0:
        indices = sorted(range(n), key=lambda i: weights[i], reverse=True)
        for j in range(diff):
            alloc[indices[j % n]] += 1
    elif diff < 0:
        indices = sorted(range(n), key=lambda i: alloc[i], reverse=True)
        j = 0
        while diff < 0 and j < n * 100:
            i = indices[j % n]
            if alloc[i] > 1:
                alloc[i] -= 1
                diff += 1
            j += 1

    return alloc


# ═══════════════════════════════════════
# 生成所有题目
# ═══════════════════════════════════════
def gen_all_problems(counts):
    """根据各分类题数生成所有题目"""
    all_problems = []
    all_answers = []

    for idx, (name, gen_func, weight) in enumerate(CATEGORIES):
        count = counts[idx]
        gen = gen_func()
        for _ in range(count):
            p, a = next(gen)
            all_problems.append(p)
            all_answers.append(a)

    return all_problems, all_answers


# ═══════════════════════════════════════
# Word 文档生成
# ═══════════════════════════════════════
def create_doc(items, title, filename, counts):
    """创建Word文档，按分类输出"""
    doc = Document()

    # 页面设置 A4
    section = doc.sections[0]
    section.page_width = 210000000
    section.page_height = 297000000
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # 标题
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.size = Pt(18)

    # 说明
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("姓名：__________    班级：__________    得分：__________")
    run.font.size = Pt(11)

    doc.add_paragraph()

    # 按分类输出
    idx = 0
    for i, (cat_name, _, _) in enumerate(CATEGORIES):
        count = counts[i]
        h = doc.add_heading(f"{cat_name}（共{count}题）", level=2)
        for run in h.runs:
            run.font.size = Pt(13)

        for j in range(count):
            if idx >= len(items):
                break
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            run = p.add_run(f"{idx+1}. {items[idx]}")
            run.font.size = Pt(10.5)
            idx += 1

        doc.add_paragraph()

    doc.save(filename)
    print(f"[OK] 已生成: {filename}  ({len(items)} 题)")


# ═══════════════════════════════════════
# 主程序
# ═══════════════════════════════════════
def main():
    total = TOTAL
    if len(sys.argv) > 1:
        try:
            total = int(sys.argv[1])
            if total < NUM_CATS:
                print(f"[!] 最少需要 {NUM_CATS} 题（每分类至少1题），已自动调整")
                total = NUM_CATS
        except ValueError:
            print(f"用法：python {sys.argv[0]} [总题数]")
            print(f"示例：python {sys.argv[0]} 1000")
            exit(1)
    # 时间戳后缀
    ts = time.strftime("%Y%m%d_%H%M%S")

    print("=" * 55)
    print(f"  小学数学计算题 {total} 道生成器")
    print("=" * 55)

    # 分配题数
    counts = allocate_counts(total, CATEGORIES)
    print(f"\n共 {NUM_CATS} 个分类，分配方案：")
    print(f"  最少: {min(counts)} 题，最多: {max(counts)} 题")
    print(f"  比例: 1 : {max(counts)/max(1,min(counts)):.1f}")
    print()
    for i, (name, _, _) in enumerate(CATEGORIES):
        print(f"  {name}: {counts[i]} 题")
    print(f"  {'─' * 30}")
    print(f"  合计: {sum(counts)} 题")

    # 生成
    print("\n正在生成题目...")
    problems, answers_list = gen_all_problems(counts)
    print(f"[OK] 生成完毕: {len(problems)} 道")

    # 输出文档
    print("\n正在生成题目卷...")
    create_doc(problems, f"小学数学计算题 {total} 道（题目卷）",
               f"小学数学计算题{total}道_题目卷_{ts}.docx", counts)

    print("正在生成答案卷...")
    create_doc(answers_list, f"小学数学计算题 {total} 道（答案卷）",
               f"小学数学计算题{total}道_答案卷_{ts}.docx", counts)

    print(f"\n{'=' * 55}")
    print(f"  全部完成！共 {total} 题")
    print(f"  题目卷: 小学数学计算题{total}道_题目卷_{ts}.docx")
    print(f"  答案卷: 小学数学计算题{total}道_答案卷_{ts}.docx")
    print(f"{'=' * 55}")


if __name__ == "__main__":
    main()
