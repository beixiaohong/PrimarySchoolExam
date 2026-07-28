"""
小学数学试卷生成器 (小升初综合版)
包含38个计算题分类：整数/小数/分数运算、运算律、混合运算、解方程、单位换算等
输出：题目卷 + 答案卷 两个Word文档
依赖：pip install python-docx
"""

import random
import time
import sys
import math
from fractions import Fraction
import threading
import time

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("请先安装 python-docx：pip install python-docx")
    exit(1)

def _safe_call(func, timeout_sec=2):
    """带超时保护的函数调用，使用线程实现（兼容Windows）"""
    result_container = [None]
    exception_container = [None]
    
    def worker():
        try:
            result_container[0] = func()
        except Exception as e:
            exception_container[0] = e
    
    thread = threading.Thread(target=worker)
    thread.daemon = True
    thread.start()
    thread.join(timeout_sec)
    
    if thread.is_alive():
        # 线程超时，返回None
        return None
    
    if exception_container[0] is not None:
        # 函数执行出错，返回None
        return None
    
    return result_container[0]

# 导入应用题生成模块
try:
    from gen_word_problems import WORD_CATEGORIES as _WP_CATS

    def generate_word_problems(total=20):
        """从 gen_word_problems 模块生成应用题，自动跳过会卡死的生成器"""
        problems, answers = [], []
        num_cats = len(_WP_CATS)
        per_cat = max(total // num_cats, 1)

        for cat_name, sub_funcs in _WP_CATS:
            per_sub = max(per_cat // len(sub_funcs), 1)
            for func in sub_funcs:
                for _ in range(per_sub):
                    result = _safe_call(func, timeout_sec=2)
                    if result and len(result) == 2:
                        problems.append(result[0])
                        answers.append(result[1])

        combined = list(zip(problems, answers))
        random.shuffle(combined)
        problems = [p for p, a in combined]
        answers = [a for p, a in combined]
        return problems[:total], answers[:total]

    print(f"已加载 gen_word_problems 模块（{len(_WP_CATS)} 个应用题大类）")

except ImportError:
    print("警告：找不到gen_word_problems模块，使用内置简化版应用题生成器")

    def generate_word_problems(total=20):
        problems, answers = [], []
        for _ in range(total):
            v, t = random.randint(30, 80), random.randint(2, 6)
            d = v * t
            problems.append(f"一辆汽车每小时行{v}千米，行了{t}小时，一共行了多少千米？")
            answers.append(f"路程 = {v} × {t} = {d}（千米）")
        return problems, answers


random.seed(int(time.time()))

# ==================== 辅助函数 ====================

DENOMINATORS = [2, 3, 4, 5, 6, 7, 8, 9, 10, 12]


def rand_fraction(max_den=12, allow_improper=False):
    """随机生成一个分数，返回 (numerator, denominator, Fraction对象)"""
    d = random.choice([x for x in DENOMINATORS if x <= max_den])
    if allow_improper:
        n = random.randint(1, d * 3)
    else:
        n = random.randint(1, d - 1)
    return n, d, Fraction(n, d)


def fmt_frac(f):
    """格式化分数为字符串（带分数形式）"""
    if f.denominator == 1:
        return str(f.numerator)
    if abs(f) > 1 and f.numerator > f.denominator:
        whole = f.numerator // f.denominator
        rem = f.numerator % f.denominator
        if rem == 0:
            return str(whole)
        return f"{whole}又{rem}/{f.denominator}"
    return f"{f.numerator}/{f.denominator}"


def fmt_frac_simple(f):
    """简单格式化分数（不带带分数转换）"""
    if f.denominator == 1:
        return str(f.numerator)
    return f"{f.numerator}/{f.denominator}"


def lcm(a, b):
    return a * b // math.gcd(a, b)


# ==================== 整数运算生成器 ====================

def gen_int_add():
    """整数加法：三位数加法"""
    while True:
        a, b = random.randint(100, 999), random.randint(100, 999)
        yield f"{a} + {b} =", f"{a} + {b} = {a + b}"


def gen_int_sub():
    """整数减法：三位数减法"""
    while True:
        a = random.randint(100, 999)
        b = random.randint(10, a)
        yield f"{a} - {b} =", f"{a} - {b} = {a - b}"


def gen_int_mul():
    """整数乘法：两位数×一位数/两位数"""
    while True:
        a = random.randint(11, 99)
        b = random.randint(2, 25)
        yield f"{a} × {b} =", f"{a} × {b} = {a * b}"


def gen_int_div():
    """整数除法：整除"""
    while True:
        b = random.randint(2, 12)
        q = random.randint(10, 99)
        a = b * q
        yield f"{a} ÷ {b} =", f"{a} ÷ {b} = {q}"


def gen_int_mix():
    """整数四则混合运算"""
    while True:
        pattern = random.randint(1, 6)
        if pattern == 1:
            a = random.randint(10, 100)
            b = random.randint(2, 9)
            c = random.randint(2, 9)
            result = a + b * c
            yield f"{a} + {b} × {c} =", f"{a} + {b} × {c} = {a} + {b*c} = {result}"
        elif pattern == 2:
            a = random.randint(50, 200)
            b = random.randint(2, 9)
            c = random.randint(2, 9)
            while b * c >= a:
                a = random.randint(50, 200)
            result = a - b * c
            yield f"{a} - {b} × {c} =", f"{a} - {b} × {c} = {a} - {b*c} = {result}"
        elif pattern == 3:
            a = random.randint(5, 30)
            b = random.randint(5, 30)
            c = random.randint(2, 9)
            result = (a + b) * c
            yield f"({a} + {b}) × {c} =", f"({a} + {b}) × {c} = {a+b} × {c} = {result}"
        elif pattern == 4:
            a, b = random.randint(2, 20), random.randint(2, 9)
            c, d = random.randint(2, 20), random.randint(2, 9)
            result = a * b + c * d
            yield f"{a} × {b} + {c} × {d} =", f"{a} × {b} + {c} × {d} = {a*b} + {c*d} = {result}"
        elif pattern == 5:
            b = random.randint(2, 9)
            q = random.randint(5, 30)
            a = b * q
            c = random.randint(10, 50)
            result = q + c
            yield f"{a} ÷ {b} + {c} =", f"{a} ÷ {b} + {c} = {q} + {c} = {result}"
        else:
            c = random.randint(2, 9)
            q = random.randint(5, 20)
            diff = c * q
            b = random.randint(10, 50)
            a = diff + b
            yield f"({a} - {b}) ÷ {c} =", f"({a} - {b}) ÷ {c} = {diff} ÷ {c} = {q}"


# ==================== 小数运算生成器 ====================

def gen_decimal_add():
    """小数加法"""
    while True:
        dp = random.choice([1, 2])
        if dp == 1:
            a = round(random.uniform(0.1, 99.9), 1)
            b = round(random.uniform(0.1, 99.9), 1)
        else:
            a = round(random.uniform(0.01, 50.99), 2)
            b = round(random.uniform(0.01, 50.99), 2)
        ans = round(a + b, dp)
        yield f"{a} + {b} =", f"{a} + {b} = {ans}"


def gen_decimal_sub():
    """小数减法"""
    while True:
        dp = random.choice([1, 2])
        if dp == 1:
            a = round(random.uniform(10, 99.9), 1)
            b = round(random.uniform(0.1, a), 1)
        else:
            a = round(random.uniform(10, 99.99), 2)
            b = round(random.uniform(0.01, a), 2)
        ans = round(a - b, dp)
        if ans < 0:
            continue
        yield f"{a} - {b} =", f"{a} - {b} = {ans}"


def gen_decimal_mul():
    """小数乘法"""
    while True:
        pattern = random.randint(1, 3)
        if pattern == 1:
            a = round(random.uniform(0.1, 20.0), 1)
            b = random.randint(2, 20)
            ans = round(a * b, 1)
            yield f"{a} × {b} =", f"{a} × {b} = {ans}"
        elif pattern == 2:
            a = round(random.uniform(0.1, 9.9), 1)
            b = round(random.uniform(0.1, 9.9), 1)
            ans = round(a * b, 2)
            yield f"{a} × {b} =", f"{a} × {b} = {ans}"
        else:
            a = round(random.uniform(0.1, 5.0), 2)
            b = round(random.uniform(0.1, 5.0), 1)
            ans = round(a * b, 3)
            yield f"{a} × {b} =", f"{a} × {b} = {ans}"


def gen_decimal_div():
    """小数除法"""
    while True:
        pattern = random.randint(1, 3)
        if pattern == 1:
            b = random.randint(2, 8)
            q = round(random.uniform(0.1, 15.0), 1)
            a = round(b * q, 1)
            yield f"{a} ÷ {b} =", f"{a} ÷ {b} = {q}"
        elif pattern == 2:
            b = round(random.choice([0.5, 0.25, 0.2, 0.1, 1.5, 2.5]), 2)
            q = random.randint(2, 30)
            a = round(b * q, 2)
            yield f"{a} ÷ {b} =", f"{a} ÷ {b} = {q}"
        else:
            b = round(random.uniform(0.1, 3.0), 1)
            q = round(random.uniform(0.5, 10.0), 1)
            a = round(b * q, 2)
            if a <= 0 or b <= 0:
                continue
            yield f"{a} ÷ {b} =", f"{a} ÷ {b} = {q}"


def gen_decimal_mix():
    """小数四则混合"""
    while True:
        pattern = random.randint(1, 4)
        if pattern == 1:
            a = round(random.uniform(1, 20), 1)
            b = round(random.uniform(0.1, 5), 1)
            c = random.randint(2, 9)
            result = round(a + b * c, 1)
            yield f"{a} + {b} × {c} =", f"{a} + {b} × {c} = {a} + {round(b*c,1)} = {result}"
        elif pattern == 2:
            a = round(random.uniform(10, 50), 1)
            b = round(random.uniform(1, 10), 1)
            c = random.randint(2, 5)
            result = round((a - b) * c, 1)
            yield f"({a} - {b}) × {c} =", f"({a} - {b}) × {c} = {round(a-b,1)} × {c} = {result}"
        elif pattern == 3:
            a = round(random.uniform(1, 10), 1)
            b = round(random.uniform(1, 10), 1)
            c = round(random.uniform(1, 10), 1)
            result = round(a * b + c, 2)
            yield f"{a} × {b} + {c} =", f"{a} × {b} + {c} = {round(a*b,2)} + {c} = {result}"
        else:
            b_val = round(random.choice([0.5, 2.0, 2.5, 0.25]), 2)
            q = random.randint(5, 30)
            a = round(b_val * q, 2)
            c = round(random.uniform(1, q - 1), 1)
            result = round(q - c, 1)
            yield f"{a} ÷ {b_val} - {c} =", f"{a} ÷ {b_val} - {c} = {q} - {c} = {result}"


# ==================== 分数运算生成器 ====================

def gen_fraction_add():
    """分数加法（同分母/异分母）"""
    while True:
        d1 = random.choice(DENOMINATORS)
        d2 = random.choice(DENOMINATORS)
        n1 = random.randint(1, d1 - 1)
        n2 = random.randint(1, d2 - 1)
        f1 = Fraction(n1, d1)
        f2 = Fraction(n2, d2)
        result = f1 + f2
        if result <= 0 or result > 10:
            continue
        yield (f"{n1}/{d1} + {n2}/{d2} =",
               f"{n1}/{d1} + {n2}/{d2} = {fmt_frac_simple(result)}")


def gen_fraction_sub():
    """分数减法"""
    while True:
        d1 = random.choice(DENOMINATORS)
        d2 = random.choice(DENOMINATORS)
        n1 = random.randint(1, d1 - 1)
        n2 = random.randint(1, d2 - 1)
        f1 = Fraction(n1, d1)
        f2 = Fraction(n2, d2)
        if f1 <= f2:
            f1, f2 = f2, f1
        result = f1 - f2
        if result <= 0:
            continue
        yield (f"{f1.numerator}/{f1.denominator} - {f2.numerator}/{f2.denominator} =",
               f"{f1.numerator}/{f1.denominator} - {f2.numerator}/{f2.denominator} = {fmt_frac_simple(result)}")


def gen_fraction_mul():
    """分数乘法"""
    while True:
        n1, d1, f1 = rand_fraction(10)
        n2, d2, f2 = rand_fraction(10)
        result = f1 * f2
        if result <= 0:
            continue
        yield (f"{n1}/{d1} × {n2}/{d2} =",
               f"{n1}/{d1} × {n2}/{d2} = {fmt_frac_simple(result)}")


def gen_fraction_div():
    """分数除法"""
    while True:
        n1, d1, f1 = rand_fraction(10)
        n2, d2, f2 = rand_fraction(10)
        result = f1 / f2
        if result <= 0 or result > 20:
            continue
        yield (f"{n1}/{d1} ÷ {n2}/{d2} =",
               f"{n1}/{d1} ÷ {n2}/{d2} = {n1}/{d1} × {d2}/{n2} = {fmt_frac_simple(result)}")


def gen_fraction_addsub():
    """分数加减混合"""
    while True:
        d1 = random.choice(DENOMINATORS)
        d2 = random.choice(DENOMINATORS)
        d3 = random.choice(DENOMINATORS)
        n1 = random.randint(1, d1 - 1)
        n2 = random.randint(1, d2 - 1)
        n3 = random.randint(1, d3 - 1)
        f1 = Fraction(n1, d1)
        f2 = Fraction(n2, d2)
        f3 = Fraction(n3, d3)
        op1 = random.choice(['+', '-'])
        op2 = random.choice(['+', '-'])

        if op1 == '+':
            temp = f1 + f2
        else:
            temp = f1 - f2
        if temp <= 0:
            continue
        if op2 == '+':
            result = temp + f3
        else:
            result = temp - f3
        if result <= 0 or result > 10:
            continue

        expr = f"{n1}/{d1} {op1} {n2}/{d2} {op2} {n3}/{d3}"
        step1 = f"{fmt_frac_simple(f1)} {op1} {fmt_frac_simple(f2)} = {fmt_frac_simple(temp)}"
        step2 = f"{fmt_frac_simple(temp)} {op2} {fmt_frac_simple(f3)} = {fmt_frac_simple(result)}"
        yield f"{expr} =", f"{expr}：{step1}，{step2}"


def gen_fraction_muldiv():
    """分数乘除混合"""
    while True:
        n1, d1, f1 = rand_fraction(8)
        n2, d2, f2 = rand_fraction(8)
        n3, d3, f3 = rand_fraction(8)
        op1 = random.choice(['×', '÷'])
        op2 = random.choice(['×', '÷'])

        if op1 == '×':
            temp = f1 * f2
        else:
            temp = f1 / f2
        if temp <= 0 or temp > 20:
            continue

        if op2 == '×':
            result = temp * f3
        else:
            result = temp / f3
        if result <= 0 or result > 20:
            continue

        expr = f"{n1}/{d1} {op1} {n2}/{d2} {op2} {n3}/{d3}"
        yield f"{expr} =", f"{expr} = {fmt_frac_simple(result)}"


def gen_fraction_mix():
    """分数四则混合"""
    while True:
        pattern = random.randint(1, 4)
        if pattern == 1:
            n1, d1, f1 = rand_fraction(8)
            n2, d2, f2 = rand_fraction(8)
            n3, d3, f3 = rand_fraction(8)
            product = f2 * f3
            result = f1 + product
            if result > 10 or result <= 0:
                continue
            yield (f"{n1}/{d1} + {n2}/{d2} × {n3}/{d3} =",
                   f"{n1}/{d1} + {n2}/{d2} × {n3}/{d3} = {fmt_frac_simple(f1)} + {fmt_frac_simple(product)} = {fmt_frac_simple(result)}")
        elif pattern == 2:
            n1, d1, f1 = rand_fraction(8)
            n2, d2, f2 = rand_fraction(8)
            n3, d3, f3 = rand_fraction(8)
            sum_f = f1 + f2
            result = sum_f * f3
            if result > 10 or result <= 0:
                continue
            yield (f"({n1}/{d1} + {n2}/{d2}) × {n3}/{d3} =",
                   f"({n1}/{d1} + {n2}/{d2}) × {n3}/{d3} = {fmt_frac_simple(sum_f)} × {n3}/{d3} = {fmt_frac_simple(result)}")
        elif pattern == 3:
            n1, d1, f1 = rand_fraction(8)
            n2, d2, f2 = rand_fraction(8)
            n3, d3, f3 = rand_fraction(8)
            product = f1 * f2
            result = product - f3
            if result <= 0:
                continue
            yield (f"{n1}/{d1} × {n2}/{d2} - {n3}/{d3} =",
                   f"{n1}/{d1} × {n2}/{d2} - {n3}/{d3} = {fmt_frac_simple(product)} - {fmt_frac_simple(f3)} = {fmt_frac_simple(result)}")
        else:
            n1, d1, f1 = rand_fraction(8)
            n2, d2, f2 = rand_fraction(8)
            n3, d3, f3 = rand_fraction(8)
            quotient = f1 / f2
            result = quotient + f3
            if result > 10 or result <= 0:
                continue
            yield (f"{n1}/{d1} ÷ {n2}/{d2} + {n3}/{d3} =",
                   f"{n1}/{d1} ÷ {n2}/{d2} + {n3}/{d3} = {fmt_frac_simple(quotient)} + {fmt_frac_simple(f3)} = {fmt_frac_simple(result)}")


# ==================== 分数特殊运算 ====================

def gen_mixed_fraction():
    """带分数运算"""
    while True:
        pattern = random.choice(['add', 'sub', 'mul'])
        if pattern == 'add':
            w1 = random.randint(1, 9)
            d1 = random.choice([2, 3, 4, 5, 6, 8])
            n1 = random.randint(1, d1 - 1)
            w2 = random.randint(1, 9)
            d2 = random.choice([2, 3, 4, 5, 6, 8])
            n2 = random.randint(1, d2 - 1)
            f1 = Fraction(w1 * d1 + n1, d1)
            f2 = Fraction(w2 * d2 + n2, d2)
            result = f1 + f2
            s1 = f"{w1}又{n1}/{d1}"
            s2 = f"{w2}又{n2}/{d2}"
            yield f"{s1} + {s2} =", f"{s1} + {s2} = {fmt_frac(result)}"
        elif pattern == 'sub':
            w1 = random.randint(3, 15)
            d1 = random.choice([2, 3, 4, 5, 6, 8])
            n1 = random.randint(1, d1 - 1)
            w2 = random.randint(1, w1 - 1)
            d2 = random.choice([2, 3, 4, 5, 6, 8])
            n2 = random.randint(1, d2 - 1)
            f1 = Fraction(w1 * d1 + n1, d1)
            f2 = Fraction(w2 * d2 + n2, d2)
            if f1 <= f2:
                continue
            result = f1 - f2
            s1 = f"{w1}又{n1}/{d1}"
            s2 = f"{w2}又{n2}/{d2}"
            yield f"{s1} - {s2} =", f"{s1} - {s2} = {fmt_frac(result)}"
        else:
            w1 = random.randint(1, 5)
            d1 = random.choice([2, 3, 4, 5])
            n1 = random.randint(1, d1 - 1)
            b = random.randint(2, 6)
            f1 = Fraction(w1 * d1 + n1, d1)
            result = f1 * b
            s1 = f"{w1}又{n1}/{d1}"
            yield f"{s1} × {b} =", f"{s1} × {b} = {fmt_frac(result)}"


def gen_improper_fraction():
    """假分数运算"""
    while True:
        d1 = random.choice([2, 3, 4, 5, 6, 8])
        n1 = random.randint(d1 + 1, d1 * 4)
        d2 = random.choice([2, 3, 4, 5, 6, 8])
        n2 = random.randint(d2 + 1, d2 * 4)
        f1 = Fraction(n1, d1)
        f2 = Fraction(n2, d2)
        op = random.choice(['+', '-', '×'])
        if op == '+':
            result = f1 + f2
            expr = f"{n1}/{d1} + {n2}/{d2}"
        elif op == '-':
            if f1 < f2:
                f1, f2 = f2, f1
                n1, d1, n2, d2 = f1.numerator, f1.denominator, f2.numerator, f2.denominator
            result = f1 - f2
            expr = f"{n1}/{d1} - {n2}/{d2}"
        else:
            result = f1 * f2
            expr = f"{n1}/{d1} × {n2}/{d2}"
        if result <= 0 or result > 50:
            continue
        yield f"{expr} =", f"{expr} = {fmt_frac_simple(result)}"


def gen_common_denominator():
    """通分"""
    while True:
        d1 = random.choice([2, 3, 4, 5, 6, 8, 10])
        d2 = random.choice([2, 3, 4, 5, 6, 8, 10])
        if d1 == d2:
            continue
        n1 = random.randint(1, d1 - 1)
        n2 = random.randint(1, d2 - 1)
        common_d = lcm(d1, d2)
        new_n1 = n1 * (common_d // d1)
        new_n2 = n2 * (common_d // d2)
        yield (f"把 {n1}/{d1} 和 {n2}/{d2} 通分：",
               f"{n1}/{d1} = {new_n1}/{common_d}，{n2}/{d2} = {new_n2}/{common_d}（公分母为{common_d}）")


def gen_reduce_fraction():
    """约分"""
    while True:
        d = random.choice([2, 3, 4, 5, 6, 8, 10, 12])
        n = random.randint(1, d - 1)
        factor = random.randint(2, 8)
        big_n = n * factor
        big_d = d * factor
        if big_d > 100:
            continue
        f = Fraction(big_n, big_d)
        if f.denominator == 1:
            yield f"约分：{big_n}/{big_d} =", f"{big_n}/{big_d} = {f.numerator}（分子分母同除以{factor}）"
        else:
            yield f"约分：{big_n}/{big_d} =", f"{big_n}/{big_d} = {f.numerator}/{f.denominator}（分子分母同除以{factor}）"


def gen_reciprocal():
    """倒数"""
    while True:
        pattern = random.randint(1, 3)
        if pattern == 1:
            d = random.choice([2, 3, 4, 5, 6, 7, 8, 9, 10])
            n = random.randint(1, d - 1)
            yield f"{n}/{d} 的倒数是____", f"{n}/{d} 的倒数是 {d}/{n}"
        elif pattern == 2:
            n = random.randint(2, 30)
            yield f"{n} 的倒数是____", f"{n} 的倒数是 1/{n}"
        else:
            w = random.randint(1, 8)
            d = random.choice([2, 3, 4, 5, 6])
            n = random.randint(1, d - 1)
            improper_n = w * d + n
            yield f"{w}又{n}/{d} 的倒数是____", f"{w}又{n}/{d} = {improper_n}/{d}，倒数是 {d}/{improper_n}"


# ==================== 混合类型运算 ====================

def gen_int_decimal_mix():
    """整数+小数混合"""
    while True:
        pattern = random.randint(1, 4)
        if pattern == 1:
            a = random.randint(5, 50)
            b = round(random.uniform(0.1, 20.0), 1)
            result = round(a + b, 1)
            yield f"{a} + {b} =", f"{a} + {b} = {result}"
        elif pattern == 2:
            a = random.randint(20, 100)
            b = round(random.uniform(0.1, a - 1), 1)
            result = round(a - b, 1)
            yield f"{a} - {b} =", f"{a} - {b} = {result}"
        elif pattern == 3:
            a = random.randint(3, 25)
            b = round(random.uniform(0.1, 10.0), 1)
            result = round(a * b, 1)
            yield f"{a} × {b} =", f"{a} × {b} = {result}"
        else:
            b = round(random.uniform(0.5, 5.0), 1)
            q = random.randint(3, 20)
            a = round(b * q, 1)
            yield f"{a} ÷ {b} =", f"{a} ÷ {b} = {q}"


def gen_int_fraction_mix():
    """整数+分数混合"""
    while True:
        pattern = random.randint(1, 4)
        if pattern == 1:
            a = random.randint(2, 20)
            n, d, f = rand_fraction(10)
            result = Fraction(a) + f
            yield f"{a} + {n}/{d} =", f"{a} + {n}/{d} = {fmt_frac(result)}"
        elif pattern == 2:
            a = random.randint(2, 20)
            n, d, f = rand_fraction(10)
            result = Fraction(a) - f
            if result <= 0:
                continue
            yield f"{a} - {n}/{d} =", f"{a} - {n}/{d} = {fmt_frac_simple(result)}"
        elif pattern == 3:
            a = random.randint(2, 15)
            n, d, f = rand_fraction(10)
            result = Fraction(a) * f
            yield f"{a} × {n}/{d} =", f"{a} × {n}/{d} = {fmt_frac_simple(result)}"
        else:
            a = random.randint(2, 20)
            n, d, f = rand_fraction(10)
            result = Fraction(a) / f
            if result > 50:
                continue
            yield f"{a} ÷ {n}/{d} =", f"{a} ÷ {n}/{d} = {a} × {d}/{n} = {fmt_frac_simple(result)}"


def gen_decimal_fraction_mix():
    """小数+分数混合"""
    while True:
        pattern = random.randint(1, 3)
        if pattern == 1:
            a = round(random.uniform(0.5, 10.0), 1)
            n, d, f = rand_fraction(8)
            a_frac = Fraction(a).limit_denominator(100)
            result = a_frac + f
            yield f"{a} + {n}/{d} =", f"{a} + {n}/{d} = {fmt_frac_simple(result)}"
        elif pattern == 2:
            a = round(random.uniform(0.5, 10.0), 1)
            n, d, f = rand_fraction(8)
            a_frac = Fraction(a).limit_denominator(100)
            result = a_frac * f
            if result <= 0:
                continue
            yield f"{a} × {n}/{d} =", f"{a} × {n}/{d} = {fmt_frac_simple(result)}"
        else:
            n, d, f = rand_fraction(8, allow_improper=True)
            a = round(random.uniform(0.1, 2.0), 1)
            a_frac = Fraction(a).limit_denominator(100)
            result = f - a_frac
            if result <= 0:
                continue
            yield f"{n}/{d} - {a} =", f"{n}/{d} - {a} = {fmt_frac_simple(result)}"


def gen_all_mix():
    """整数+小数+分数综合混合"""
    while True:
        pattern = random.randint(1, 6)
        if pattern == 1:
            a = random.randint(2, 30)
            b = round(random.uniform(0.1, 10), 1)
            n, d, f = rand_fraction(8)
            b_frac = Fraction(b).limit_denominator(100)
            product = b_frac * f
            result = Fraction(a) + product
            if result > 100 or result <= 0:
                continue
            yield (f"{a} + {b} × {n}/{d} =",
                   f"{a} + {b} × {n}/{d} = {a} + {fmt_frac_simple(product)} = {fmt_frac_simple(result)}")
        elif pattern == 2:
            a = random.randint(5, 50)
            b = round(random.uniform(0.5, 5), 1)
            n, d, f = rand_fraction(8)
            diff = Fraction(a) - Fraction(b).limit_denominator(100)
            if diff <= 0:
                continue
            result = diff / f
            if result > 100 or result <= 0:
                continue
            yield (f"({a} - {b}) ÷ {n}/{d} =",
                   f"({a} - {b}) ÷ {n}/{d} = {fmt_frac_simple(diff)} × {d}/{n} = {fmt_frac_simple(result)}")
        elif pattern == 3:
            a = random.randint(2, 20)
            b = round(random.uniform(0.5, 5), 1)
            n, d, f = rand_fraction(8)
            product = Fraction(a) * Fraction(b).limit_denominator(100)
            result = product + f
            yield (f"{a} × {b} + {n}/{d} =",
                   f"{a} × {b} + {n}/{d} = {fmt_frac_simple(product)} + {n}/{d} = {fmt_frac_simple(result)}")
        elif pattern == 4:
            n1, d1, f1 = rand_fraction(8)
            a = random.randint(2, 15)
            b = round(random.uniform(0.1, 5), 1)
            product = f1 * a
            result = Fraction(product) + Fraction(b).limit_denominator(100)
            if result > 50:
                continue
            yield (f"{n1}/{d1} × {a} + {b} =",
                   f"{n1}/{d1} × {a} + {b} = {fmt_frac_simple(Fraction(product))} + {b} = {fmt_frac_simple(result)}")
        elif pattern == 5:
            a = random.randint(10, 100)
            n, d, f = rand_fraction(8)
            b = round(random.uniform(0.1, 5), 1)
            product = Fraction(a) * f
            b_frac = Fraction(b).limit_denominator(100)
            result = product - b_frac
            if result <= 0:
                continue
            yield (f"{a} × {n}/{d} - {b} =",
                   f"{a} × {n}/{d} - {b} = {fmt_frac_simple(product)} - {b} = {fmt_frac_simple(result)}")
        else:
            a = random.randint(2, 20)
            n1, d1, f1 = rand_fraction(8)
            n2, d2, f2 = rand_fraction(8)
            sum_f = f1 + f2
            result = Fraction(a) * sum_f
            if result > 100:
                continue
            yield (f"{a} × ({n1}/{d1} + {n2}/{d2}) =",
                   f"{a} × ({n1}/{d1} + {n2}/{d2}) = {a} × {fmt_frac_simple(sum_f)} = {fmt_frac_simple(result)}")


# ==================== 运算律生成器 ====================

def gen_add_exchange():
    """加法交换律: a + b = b + a"""
    while True:
        pattern = random.randint(1, 3)
        if pattern == 1:
            a = random.randint(10, 500)
            b = random.randint(10, 500)
            yield f"{a} + {b} = {b} + ____", f"{a} + {b} = {b} + {a}（加法交换律）"
        elif pattern == 2:
            a = round(random.uniform(1, 50), 1)
            b = round(random.uniform(1, 50), 1)
            yield f"{a} + {b} = {b} + ____", f"{a} + {b} = {b} + {a}（加法交换律）"
        else:
            a = random.randint(10, 200)
            b = random.randint(10, 200)
            c = random.randint(10, 200)
            yield f"{a} + {b} + {c} = {c} + ____ + {a}", f"填 {b}（加法交换律）"


def gen_add_associative():
    """加法结合律: (a + b) + c = a + (b + c)"""
    while True:
        a = random.randint(100, 500)
        c = random.randint(100, 500)
        target = ((c // 100) + 1) * 100
        b = target - c
        if b <= 0 or b > 500:
            continue
        yield (f"{a} + {b} + {c}（用简便方法计算）=",
               f"{a} + {b} + {c} = {a} + ({b} + {c}) = {a} + {target} = {a + target}（加法结合律）")


def gen_mul_exchange():
    """乘法交换律: a × b = b × a"""
    while True:
        pattern = random.randint(1, 2)
        if pattern == 1:
            a = random.randint(2, 50)
            b = random.randint(2, 50)
            yield f"{a} × {b} = {b} × ____", f"{a} × {b} = {b} × {a}（乘法交换律）"
        else:
            a = random.randint(2, 25)
            b = random.randint(2, 25)
            c = random.randint(2, 25)
            yield f"{a} × {b} × {c} = {c} × ____ × {a}", f"填 {b}（乘法交换律）"


def gen_mul_associative():
    """乘法结合律: (a × b) × c = a × (b × c)"""
    while True:
        pairs = [(25, 4), (125, 8), (50, 2), (20, 5), (40, 25)]
        a, c = random.choice(pairs)
        b = random.randint(2, 50)
        yield (f"{a} × {b} × {c}（用简便方法计算）=",
               f"{a} × {b} × {c} = ({a} × {c}) × {b} = {a*c} × {b} = {a*c*b}（乘法结合律）")


def gen_distributive():
    """乘法分配律: a × (b + c) = a × b + a × c"""
    while True:
        pattern = random.randint(1, 4)
        if pattern == 1:
            a = random.randint(2, 20)
            b = random.randint(10, 50)
            c = random.randint(10, 50)
            result = a * (b + c)
            yield (f"{a} × ({b} + {c}) =",
                   f"{a} × ({b} + {c}) = {a}×{b} + {a}×{c} = {a*b} + {a*c} = {result}（分配律）")
        elif pattern == 2:
            a = random.randint(2, 15)
            b = random.randint(10, 50)
            c = random.randint(10, 50)
            result = a * b + a * c
            yield (f"{a} × {b} + {a} × {c} =",
                   f"{a} × {b} + {a} × {c} = {a} × ({b} + {c}) = {a} × {b+c} = {result}（分配律）")
        elif pattern == 3:
            a = random.randint(2, 15)
            b = random.randint(30, 80)
            c = random.randint(10, b - 1)
            result = a * b - a * c
            yield (f"{a} × {b} - {a} × {c} =",
                   f"{a} × {b} - {a} × {c} = {a} × ({b} - {c}) = {a} × {b-c} = {result}（分配律）")
        else:
            a = random.randint(11, 99)
            result = 101 * a
            yield (f"101 × {a}（用简便方法计算）=",
                   f"101 × {a} = (100 + 1) × {a} = 100×{a} + 1×{a} = {100*a} + {a} = {result}（分配律）")


def gen_fraction_distributive():
    """分数分配律"""
    while True:
        pattern = random.randint(1, 2)
        if pattern == 1:
            a = random.randint(2, 12)
            n1, d1, f1 = rand_fraction(8)
            n2, d2, f2 = rand_fraction(8)
            sum_f = f1 + f2
            result = Fraction(a) * sum_f
            r1 = Fraction(a) * f1
            r2 = Fraction(a) * f2
            if result > 50:
                continue
            yield (f"{a} × ({n1}/{d1} + {n2}/{d2}) =",
                   f"{a} × ({n1}/{d1} + {n2}/{d2}) = {a}×{n1}/{d1} + {a}×{n2}/{d2} = {fmt_frac_simple(r1)} + {fmt_frac_simple(r2)} = {fmt_frac_simple(result)}")
        else:
            n, d, f = rand_fraction(8)
            a = random.randint(3, 20)
            b = random.randint(3, 20)
            r1 = f * a
            r2 = f * b
            result = r1 + r2
            yield (f"{n}/{d} × {a} + {n}/{d} × {b} =",
                   f"{n}/{d} × {a} + {n}/{d} × {b} = {n}/{d} × ({a} + {b}) = {n}/{d} × {a+b} = {fmt_frac_simple(result)}")


def gen_fraction_associative():
    """分数结合律"""
    while True:
        n1, d1, f1 = rand_fraction(8)
        complement_d = d1
        complement_n = d1 - n1
        if complement_n <= 0:
            continue
        f2 = Fraction(complement_n, complement_d)
        n3, d3, f3 = rand_fraction(8)
        result = f1 + f2 + f3
        yield (f"{n1}/{d1} + {complement_n}/{complement_d} + {n3}/{d3}（用简便方法计算）=",
               f"({n1}/{d1} + {complement_n}/{complement_d}) + {n3}/{d3} = 1 + {n3}/{d3} = {fmt_frac(result)}（结合律）")


def gen_sub_property():
    """减法性质: a - b - c = a - (b + c)"""
    while True:
        a = random.randint(100, 999)
        b = random.randint(10, a // 3)
        c = random.randint(10, a // 3)
        if b + c >= a:
            continue
        result = a - b - c
        yield (f"{a} - {b} - {c}（用简便方法计算）=",
               f"{a} - {b} - {c} = {a} - ({b} + {c}) = {a} - {b+c} = {result}（减法性质）")


def gen_div_property():
    """除法性质: a ÷ b ÷ c = a ÷ (b × c)"""
    while True:
        b = random.randint(2, 8)
        c = random.randint(2, 8)
        q = random.randint(5, 30)
        a = b * c * q
        yield (f"{a} ÷ {b} ÷ {c}（用简便方法计算）=",
               f"{a} ÷ {b} ÷ {c} = {a} ÷ ({b} × {c}) = {a} ÷ {b*c} = {q}（除法性质）")


def gen_fast_calc():
    """巧算（综合技巧）"""
    while True:
        pattern = random.randint(1, 5)
        if pattern == 1:
            n = random.randint(2, 40)
            result = 25 * n * 4
            yield (f"25 × {n} × 4 =",
                   f"25 × {n} × 4 = (25 × 4) × {n} = 100 × {n} = {result}")
        elif pattern == 2:
            n = random.randint(2, 20)
            result = 125 * n * 8
            yield (f"125 × {n} × 8 =",
                   f"125 × {n} × 8 = (125 × 8) × {n} = 1000 × {n} = {result}")
        elif pattern == 3:
            a = random.randint(2, 50)
            result = 999 * a
            yield (f"999 × {a}（用简便方法计算）=",
                   f"999 × {a} = (1000 - 1) × {a} = 1000×{a} - 1×{a} = {1000*a} - {a} = {result}")
        elif pattern == 4:
            a = random.randint(11, 99)
            result = 99 * a
            yield (f"99 × {a}（用简便方法计算）=",
                   f"99 × {a} = (100 - 1) × {a} = 100×{a} - 1×{a} = {100*a} - {a} = {result}")
        else:
            a = random.randint(2, 50)
            result = a * 10
            yield (f"{a} × 9.9 + {a} × 0.1 =",
                   f"{a} × 9.9 + {a} × 0.1 = {a} × (9.9 + 0.1) = {a} × 10 = {result}")


# ==================== 解方程 ====================

def gen_equation_simple():
    """解方程（小升初难度）"""
    while True:
        pattern = random.randint(1, 8)
        if pattern == 1:
            x = random.randint(1, 50)
            a = random.randint(2, 9)
            b = random.randint(1, 50)
            c = a * x + b
            yield (f"{a}x + {b} = {c}",
                   f"解：{a}x = {c} - {b}\n    {a}x = {c - b}\n    x = {c - b} ÷ {a}\n    x = {x}")
        elif pattern == 2:
            x = random.randint(5, 50)
            a = random.randint(2, 9)
            b = random.randint(1, a * x - 1)
            c = a * x - b
            yield (f"{a}x - {b} = {c}",
                   f"解：{a}x = {c} + {b}\n    {a}x = {c + b}\n    x = {c + b} ÷ {a}\n    x = {x}")
        elif pattern == 3:
            x = random.randint(10, 200)
            a = random.randint(10, 100)
            b = x + a
            yield (f"x + {a} = {b}",
                   f"解：x = {b} - {a}\n    x = {x}")
        elif pattern == 4:
            x = random.randint(10, 100)
            a = random.randint(x + 1, x + 200)
            b = a - x
            yield (f"{a} - x = {b}",
                   f"解：x = {a} - {b}\n    x = {x}")
        elif pattern == 5:
            x = random.randint(2, 30)
            a = random.randint(2, 8)
            b = random.randint(1, 8)
            c = (a + b) * x
            yield (f"{a}x + {b}x = {c}",
                   f"解：({a} + {b})x = {c}\n    {a+b}x = {c}\n    x = {c} ÷ {a+b}\n    x = {x}")
        elif pattern == 6:
            x = random.randint(2, 30)
            a = random.randint(2, 8)
            b = random.randint(1, 20)
            c = a * (x + b)
            yield (f"{a}(x + {b}) = {c}",
                   f"解：x + {b} = {c} ÷ {a}\n    x + {b} = {c // a}\n    x = {c // a} - {b}\n    x = {x}")
        elif pattern == 7:
            a = random.randint(2, 8)
            x_base = random.randint(2, 20)
            x = a * x_base
            b = random.randint(1, 30)
            c = x_base + b
            yield (f"x ÷ {a} + {b} = {c}",
                   f"解：x ÷ {a} = {c} - {b}\n    x ÷ {a} = {c - b}\n    x = {c - b} × {a}\n    x = {x}")
        else:
            x = random.randint(2, 20)
            a = random.randint(5, 12)
            c_coeff = random.randint(2, a - 1)
            b = random.randint(1, 30)
            d = (a - c_coeff) * x + b
            if d <= 0:
                continue
            yield (f"{a}x + {b} = {c_coeff}x + {d}",
                   f"解：{a}x - {c_coeff}x = {d} - {b}\n    {a - c_coeff}x = {d - b}\n    x = {d - b} ÷ {a - c_coeff}\n    x = {x}")


# ==================== 单位换算 ====================

def gen_length_convert():
    """单位换算（长度/重量/时间/面积）"""
    while True:
        pattern = random.randint(1, 12)
        if pattern == 1:
            val = random.randint(1, 20)
            yield f"{val}千米 = ____米", f"{val}千米 = {val * 1000}米"
        elif pattern == 2:
            val = random.randint(1, 50) * 1000
            yield f"{val}米 = ____千米", f"{val}米 = {val // 1000}千米"
        elif pattern == 3:
            val = random.randint(1, 30)
            yield f"{val}米 = ____厘米", f"{val}米 = {val * 100}厘米"
        elif pattern == 4:
            val = random.randint(1, 20) * 100
            yield f"{val}厘米 = ____米", f"{val}厘米 = {val // 100}米"
        elif pattern == 5:
            val = random.randint(1, 20)
            yield f"{val}分米 = ____厘米", f"{val}分米 = {val * 10}厘米"
        elif pattern == 6:
            val = random.randint(1, 50)
            yield f"{val}千克 = ____克", f"{val}千克 = {val * 1000}克"
        elif pattern == 7:
            val = random.randint(1, 20) * 1000
            yield f"{val}克 = ____千克", f"{val}克 = {val // 1000}千克"
        elif pattern == 8:
            val = random.randint(1, 10)
            yield f"{val}吨 = ____千克", f"{val}吨 = {val * 1000}千克"
        elif pattern == 9:
            val = random.randint(1, 20) * 1000
            yield f"{val}千克 = ____吨", f"{val}千克 = {val // 1000}吨"
        elif pattern == 10:
            val = random.randint(1, 24)
            yield f"{val}小时 = ____分钟", f"{val}小时 = {val * 60}分钟"
        elif pattern == 11:
            val = random.randint(1, 10) * 60
            yield f"{val}分钟 = ____小时", f"{val}分钟 = {val // 60}小时"
        else:
            val = random.randint(1, 20)
            yield f"{val}平方米 = ____平方分米", f"{val}平方米 = {val * 100}平方分米"


# ==================== 计算题分类配置 ====================

CALC_CATEGORIES = [
    ("整数加法", gen_int_add, 6),
    ("整数减法", gen_int_sub, 6),
    ("整数乘法", gen_int_mul, 6),
    ("整数除法", gen_int_div, 5),
    ("整数混合", gen_int_mix, 12),
    ("小数加法", gen_decimal_add, 5),
    ("小数减法", gen_decimal_sub, 5),
    ("小数乘法", gen_decimal_mul, 6),
    ("小数除法", gen_decimal_div, 6),
    ("小数混合", gen_decimal_mix, 8),
    ("分数加法", gen_fraction_add, 5),
    ("分数减法", gen_fraction_sub, 5),
    ("分数乘法", gen_fraction_mul, 6),
    ("分数除法", gen_fraction_div, 6),
    ("分数加减混合", gen_fraction_addsub, 6),
    ("分数乘除混合", gen_fraction_muldiv, 6),
    ("分数四则混合", gen_fraction_mix, 6),
    ("带分数", gen_mixed_fraction, 4),
    ("假分数", gen_improper_fraction, 2),
    ("通分", gen_common_denominator, 2),
    ("约分", gen_reduce_fraction, 2),
    ("倒数", gen_reciprocal, 3),
    ("整数+小数", gen_int_decimal_mix, 5),
    ("整数+分数", gen_int_fraction_mix, 5),
    ("小数+分数", gen_decimal_fraction_mix, 5),
    ("整数+小数+分数", gen_all_mix, 15),
    ("加法交换律", gen_add_exchange, 2),
    ("加法结合律", gen_add_associative, 2),
    ("乘法交换律", gen_mul_exchange, 2),
    ("乘法结合律", gen_mul_associative, 3),
    ("乘法分配律", gen_distributive, 6),
    ("分数分配律", gen_fraction_distributive, 2),
    ("分数结合律", gen_fraction_associative, 2),
    ("减法性质", gen_sub_property, 2),
    ("除法性质", gen_div_property, 2),
    ("巧算", gen_fast_calc, 2),
    ("解方程", gen_equation_simple, 15),
    ("单位换算", gen_length_convert, 10),
]

# ==================== 题目生成主逻辑 ====================


def generate_calc_problems():
    """按照配置的权重生成所有计算题，返回 (题目列表, 答案列表, 分类统计)"""
    problems = []
    answers = []
    category_stats = {}

    for name, gen_func, count in CALC_CATEGORIES:
        gen = gen_func()
        cat_problems = []
        cat_answers = []
        attempts = 0
        max_attempts = count * 20
        while len(cat_problems) < count and attempts < max_attempts:
            try:
                p, a = next(gen)
                cat_problems.append(p)
                cat_answers.append(a)
            except Exception:
                pass
            attempts += 1

        problems.extend(cat_problems)
        answers.extend(cat_answers)
        category_stats[name] = len(cat_problems)

    # 打乱顺序
    combined = list(zip(problems, answers))
    random.shuffle(combined)
    problems = [p for p, a in combined]
    answers = [a for p, a in combined]

    return problems, answers, category_stats


# ==================== 文档生成 ====================

def create_exam_doc(calc_problems, calc_answers, word_problems, word_answers,
                    category_stats=None, is_answer=False):
    """生成试卷Word文档"""
    doc = Document()

    # 页面设置 A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    calc_count = len(calc_problems)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_text = "小学数学综合试卷（答案卷）" if is_answer else "小学数学综合试卷（题目卷）"
    run = title.add_run(title_text)
    run.font.size = Pt(18)
    run.bold = True

    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("（小升初专项训练 · 计算题 + 应用题）")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # 信息栏
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = info.add_run("姓名：__________    班级：__________    日期：__________    得分：__________")
    run.font.size = Pt(11)

    doc.add_paragraph()

    # ===== 第一部分：计算题 =====
    section1 = doc.add_paragraph()
    run = section1.add_run(f"一、计算题（共{calc_count}题，共80分）")
    run.font.size = Pt(14)
    run.bold = True

    if not is_answer and category_stats:
        stats_p = doc.add_paragraph()
        run = stats_p.add_run("题目包含以下题型：")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(120, 120, 120)
        stats_text = "、".join([f"{k}({v}题)" for k, v in category_stats.items()])
        run = stats_p.add_run(stats_text)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(150, 150, 150)

    if is_answer:
        # 答案卷：按分类分组显示
        idx = 1
        offset = 0
        for name, gen_func, count in CALC_CATEGORIES:
            actual_count = category_stats.get(name, 0) if category_stats else count
            if actual_count == 0:
                offset += actual_count
                continue

            cat_p = doc.add_paragraph()
            cat_p.space_before = Pt(6)
            cat_p.space_after = Pt(2)
            run = cat_p.add_run(f"【{name}】（{actual_count}题）")
            run.font.size = Pt(11)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 80, 160)

            for i in range(actual_count):
                if offset + i < len(calc_answers):
                    p = doc.add_paragraph()
                    p.space_before = Pt(1)
                    p.space_after = Pt(1)
                    run = p.add_run(f"{idx}. {calc_answers[offset + i]}")
                    run.font.size = Pt(9)
                    idx += 1
            offset += actual_count
    else:
        # 题目卷：4列表格
        cols = 4
        rows = (calc_count + cols - 1) // cols
        table = doc.add_table(rows=rows, cols=cols)

        for i, prob in enumerate(calc_problems):
            row_idx = i % rows
            col_idx = i // rows
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.space_before = Pt(1)
            p.space_after = Pt(1)
            run = p.add_run(f"{i+1}. {prob}")
            run.font.size = Pt(9)

    doc.add_paragraph()

    # ===== 第二部分：应用题 =====
    word_count = len(word_problems)
    section2 = doc.add_paragraph()
    run = section2.add_run(f"二、应用题（共{word_count}题，共20分）")
    run.font.size = Pt(14)
    run.bold = True

    if is_answer:
        for idx, ans in enumerate(word_answers, 1):
            p = doc.add_paragraph()
            p.space_before = Pt(3)
            p.space_after = Pt(3)
            run = p.add_run(f"{idx}. {ans}")
            run.font.size = Pt(10)
    else:
        for idx, prob in enumerate(word_problems, 1):
            p = doc.add_paragraph()
            p.space_before = Pt(4)
            p.space_after = Pt(4)
            run = p.add_run(f"{idx}. {prob}")
            run.font.size = Pt(11)
            blank = doc.add_paragraph()
            run = blank.add_run("\n\n")
            run.font.size = Pt(11)

    return doc


# ==================== 主程序 ====================

def main():
    word_total = 20

    if len(sys.argv) >= 2:
        try:
            word_total = int(sys.argv[1])
        except ValueError:
            pass

    print("=" * 55)
    print("  小学数学综合试卷生成器（小升初专项训练版）")
    print("=" * 55)
    print(f"\n共 {len(CALC_CATEGORIES)} 个题型分类\n")

    # 生成计算题
    print("[1/2] 生成计算题...")
    calc_problems, calc_answers, category_stats = generate_calc_problems()
    calc_total = len(calc_problems)
    print(f"  共 {calc_total} 道计算题，各题型数量：")
    for name, count in category_stats.items():
        print(f"    {name}: {count}题")

    # 生成应用题
    print(f"\n[2/2] 生成应用题（{word_total}道）...")
    word_problems, word_answers = generate_word_problems(word_total)
    print(f"  应用题：{len(word_problems)} 道")

    # 时间戳
    ts = time.strftime("%Y%m%d_%H%M%S")

    # 生成题目卷
    print("\n生成题目卷...")
    doc_q = create_exam_doc(calc_problems, calc_answers, word_problems, word_answers,
                            category_stats, is_answer=False)
    fname_q = f"小升初数学试卷_题目卷_{ts}.docx"
    doc_q.save(fname_q)
    print(f"  -> {fname_q}")

    # 生成答案卷
    print("生成答案卷...")
    doc_a = create_exam_doc(calc_problems, calc_answers, word_problems, word_answers,
                            category_stats, is_answer=True)
    fname_a = f"小升初数学试卷_答案卷_{ts}.docx"
    doc_a.save(fname_a)
    print(f"  -> {fname_a}")

    print(f"\n{'=' * 55}")
    print(f"  [OK] 试卷生成完成！")
    print(f"  计算题 {calc_total} 道 + 应用题 {len(word_problems)} 道")
    print(f"  总分：100分（计算80分 + 应用20分）")
    print(f"{'=' * 55}")


if __name__ == "__main__":
    main()
