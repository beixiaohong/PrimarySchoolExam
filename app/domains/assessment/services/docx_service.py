"""Word文档生成服务"""
import os
from datetime import datetime
from typing import List, Dict
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from app.config import OUTPUT_DIR, BASE_DIR
from app.schemas.problem import ProblemItem


def _add_image(doc, image_path: str, width_cm: float = 5.5):
    """将 web 路径图片嵌入 Word 文档（居中）"""
    if not image_path:
        return
    # /output/figures/xxx.png -> 项目根/output/figures/xxx.png
    abs_path = BASE_DIR / image_path.lstrip("/")
    if not abs_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(abs_path), width=Cm(width_cm))


def build_math_docx(
    problems: List[ProblemItem],
    grade: int = 6,
    difficulty: str = "综合",
    title: str = None,
) -> str:
    """生成数学试卷Word文档，返回文件路径"""
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 标题（学段按年级判断：1-6 小学 / 7-9 初中）
    from app.services.semester import stage_label
    title_text = title or f"{stage_label(grade)}{grade}年级数学练习（{difficulty}）"
    h = doc.add_heading(title_text, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 信息栏
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"姓名：________    班级：________    得分：________")
    run.font.size = Pt(11)

    doc.add_paragraph()  # 空行

    # 按大类分组
    categories = {}
    for p in problems:
        categories.setdefault(p.category, []).append(p)

    section_num = 1
    cn_nums = "一二三四五六七八九十"
    for cat_name, items in categories.items():
        # 大题标题
        sec_title = doc.add_paragraph()
        run = sec_title.add_run(f"{cn_nums[section_num-1] if section_num <= 10 else section_num}、{cat_name}（共{len(items)}题）")
        run.bold = True
        run.font.size = Pt(12)

        for item in items:
            p = doc.add_paragraph()
            # 题号 + 难度标记
            diff_mark = "★" * item.difficulty
            run = p.add_run(f"{item.id}. {item.question}")
            run.font.size = Pt(11)
            if item.difficulty >= 4:
                star = p.add_run(f"  {diff_mark}")
                star.font.size = Pt(9)
                star.font.color.rgb = RGBColor(200, 50, 50)

            # 嵌入配图
            if item.image_path:
                _add_image(doc, item.image_path)

            # 答题空间
            doc.add_paragraph("   ")

        section_num += 1

    # 答案部分（分页）
    doc.add_page_break()
    ans_title = doc.add_heading("参考答案", level=1)
    ans_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for p in problems:
        para = doc.add_paragraph()
        run = para.add_run(f"{p.id}. ")
        run.bold = True
        run.font.size = Pt(10)
        run2 = para.add_run(p.answer)
        run2.font.size = Pt(10)

    # 保存
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"数学_{grade}年级_{difficulty}_{timestamp}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    return filepath


def build_english_docx(
    exercises: Dict[str, list],
    grade: int = 6,
    title: str = None,
    type_names: Dict[str, str] = None,
    filename_prefix: str = "英语",
) -> str:
    """生成英语/语文试卷Word文档"""
    from ..services.english_generator import TYPE_NAMES as EN_TYPE_NAMES

    if type_names is None:
        type_names = EN_TYPE_NAMES

    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    from app.services.semester import stage_label
    title_text = title or f"{stage_label(grade)}{grade}年级英语练习"
    h = doc.add_heading(title_text, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("姓名：________    班级：________    得分：________")
    run.font.size = Pt(11)

    doc.add_paragraph()

    cn_nums = "一二三四五六七八九十"
    section_num = 0

    # 有选项的题型
    choice_types = {"choice", "grammar_choice", "situational", "cloze"}

    for key, items in exercises.items():
        if not items:
            continue
        section_num += 1
        type_name = type_names.get(key, key)
        num_label = cn_nums[section_num - 1] if section_num <= 10 else str(section_num)
        sec_title = doc.add_paragraph()
        run = sec_title.add_run(f"{num_label}、{type_name}（共{len(items)}题）")
        run.bold = True
        run.font.size = Pt(12)

        if key in choice_types:
            for item in items:
                p = doc.add_paragraph()
                run = p.add_run(f"{item['id']}. {item['question']}")
                run.font.size = Pt(11)
                if item.get("options"):
                    for opt in item["options"]:
                        op = doc.add_paragraph(f"    {opt}")
                        op.runs[0].font.size = Pt(11)
        else:
            for item in items:
                p = doc.add_paragraph()
                run = p.add_run(f"{item['id']}. {item['question']}")
                run.font.size = Pt(11)
                # 答题空间
                blank = doc.add_paragraph("   ______________________________")
                blank.runs[0].font.size = Pt(11)

    # 答案
    doc.add_page_break()
    ans_title = doc.add_heading("参考答案", level=1)
    ans_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    section_num = 0
    for key, items in exercises.items():
        if not items:
            continue
        section_num += 1
        type_name = type_names.get(key, key)
        num_label = cn_nums[section_num - 1] if section_num <= 10 else str(section_num)
        sec_p = doc.add_paragraph()
        run = sec_p.add_run(f"{num_label}、{type_name}")
        run.bold = True
        run.font.size = Pt(10)
        for item in items:
            p = doc.add_paragraph()
            run = p.add_run(f"{item['id']}. {item['answer']}")
            run.font.size = Pt(10)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{filename_prefix}_{grade}年级_{timestamp}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    return filepath


def build_wrong_practice_docx(
    questions: list,
    include_answer: bool = True,
) -> str:
    """生成错题专项练习Word文档"""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    h = doc.add_heading("错题专项练习", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("姓名：________    班级：________    得分：________")
    run.font.size = Pt(11)

    doc.add_paragraph()

    # 按题型分组
    from collections import OrderedDict
    groups = OrderedDict()
    for q in questions:
        key = f"{q.subject} - {q.type_name or q.type_code}"
        groups.setdefault(key, []).append(q)

    cn_nums = "一二三四五六七八九十"
    section_num = 0

    for group_name, items in groups.items():
        section_num += 1
        num_label = cn_nums[section_num - 1] if section_num <= 10 else str(section_num)
        sec_title = doc.add_paragraph()
        run = sec_title.add_run(f"{num_label}、{group_name}（共{len(items)}题）")
        run.bold = True
        run.font.size = Pt(12)

        for idx, q in enumerate(items, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{idx}. {q.question}")
            run.font.size = Pt(11)

            # 如果有选项
            if q.options_json:
                import json
                try:
                    opts = json.loads(q.options_json)
                    if opts:
                        for opt in opts:
                            op = doc.add_paragraph(f"    {opt}")
                            op.runs[0].font.size = Pt(11)
                except (json.JSONDecodeError, TypeError):
                    pass

            # 嵌入配图
            if getattr(q, 'image_path', ''):
                _add_image(doc, q.image_path)

            # 答题空间
            doc.add_paragraph("   ______________________________")

    # 答案部分
    if include_answer:
        doc.add_page_break()
        ans_title = doc.add_heading("参考答案", level=1)
        ans_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        section_num = 0
        for group_name, items in groups.items():
            section_num += 1
            num_label = cn_nums[section_num - 1] if section_num <= 10 else str(section_num)
            sec_p = doc.add_paragraph()
            run = sec_p.add_run(f"{num_label}、{group_name}")
            run.bold = True
            run.font.size = Pt(10)
            for idx, q in enumerate(items, 1):
                p = doc.add_paragraph()
                run = p.add_run(f"{idx}. {q.answer}")
                run.font.size = Pt(10)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"错题练习_{timestamp}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    return filepath
