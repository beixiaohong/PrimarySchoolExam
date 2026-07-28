"""Word文档生成服务"""
import os
from datetime import datetime
from typing import List, Dict

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from ..config import OUTPUT_DIR
from ..schemas.problem import ProblemItem


def build_math_docx(
    problems: List[ProblemItem],
    grade: int = 6,
    difficulty: str = "综合",
    title: str = None,
) -> str:
    """生成数学试卷Word文档，返回文件路径"""
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 标题
    title_text = title or f"小学{grade}年级数学练习（{difficulty}）"
    h = doc.add_heading(title_text, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 信息栏
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"姓名：________    班级：________    得分：________")
    run.font.size = Pt(11)

    doc.add_paragraph()  # 空行

    # 按大类分组
    categories = {}
    for p in problems:
        categories.setdefault(p.category, []).append(p)

    section_num = 1
    cn_nums = "一二三四五六七八九十"
    for cat_name, items in categories.items():
        # 大题标题
        sec_title = doc.add_paragraph()
        run = sec_title.add_run(f"{cn_nums[section_num-1] if section_num <= 10 else section_num}、{cat_name}（共{len(items)}题）")
        run.bold = True
        run.font.size = Pt(12)

        for item in items:
            p = doc.add_paragraph()
            # 题号 + 难度标记
            diff_mark = "★" * item.difficulty
            run = p.add_run(f"{item.id}. {item.question}")
            run.font.size = Pt(11)
            if item.difficulty >= 4:
                star = p.add_run(f"  {diff_mark}")
                star.font.size = Pt(9)
                star.font.color.rgb = RGBColor(200, 50, 50)

            # 答题空间
            doc.add_paragraph("   ")

        section_num += 1

    # 答案部分（分页）
    doc.add_page_break()
    ans_title = doc.add_heading("参考答案", level=1)
    ans_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for p in problems:
        para = doc.add_paragraph()
        run = para.add_run(f"{p.id}. ")
        run.bold = True
        run.font.size = Pt(10)
        run2 = para.add_run(p.answer)
        run2.font.size = Pt(10)

    # 保存
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"数学_{grade}年级_{difficulty}_{timestamp}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    return filepath


def build_english_docx(
    exercises: Dict[str, list],
    grade: int = 6,
    title: str = None,
) -> str:
    """生成英语试卷Word文档"""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    title_text = title or f"小学{grade}年级英语练习"
    h = doc.add_heading(title_text, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("姓名：________    班级：________    得分：________")
    run.font.size = Pt(11)

    doc.add_paragraph()

    section_titles = {
        "dictation": "一、单词听写（根据中文和音标写出英文单词）",
        "choice": "二、选择题",
        "translation": "三、翻译题",
        "unscramble": "四、字母重排（组成正确单词）",
    }

    for key, items in exercises.items():
        if not items:
            continue
        sec_title = doc.add_paragraph()
        run = sec_title.add_run(section_titles.get(key, key))
        run.bold = True
        run.font.size = Pt(12)

        if key == "choice":
            for item in items:
                p = doc.add_paragraph()
                run = p.add_run(f"{item['id']}. {item['question']}")
                run.font.size = Pt(11)
                for opt in item["options"]:
                    op = doc.add_paragraph(f"    {opt}")
                    op.runs[0].font.size = Pt(11)
        else:
            for item in items:
                p = doc.add_paragraph()
                run = p.add_run(f"{item['id']}. {item['question']}")
                run.font.size = Pt(11)
                if key == "dictation":
                    # 添加下划线答题区
                    blank = doc.add_paragraph("   ______________________________")
                    blank.runs[0].font.size = Pt(11)

    # 答案
    doc.add_page_break()
    ans_title = doc.add_heading("参考答案", level=1)
    ans_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for key, items in exercises.items():
        if not items:
            continue
        sec_p = doc.add_paragraph()
        run = sec_p.add_run(section_titles.get(key, key).split("（")[0])
        run.bold = True
        for item in items:
            p = doc.add_paragraph()
            run = p.add_run(f"{item['id']}. {item['answer']}")
            run.font.size = Pt(10)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"英语_{grade}年级_{timestamp}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    return filepath
