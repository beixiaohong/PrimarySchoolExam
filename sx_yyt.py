#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
生成小学应用题500道（题目卷 + 答案卷）Word文档
覆盖20种经典题型：行程、工程、浓度、利润折扣、比例、年龄、鸡兔同笼、
植树、盈亏、和差、和倍差倍、归一、分数百分数、几何、平均数、还原、
统计、时间日期、利息税率、综合。
依赖：pip install python-docx
"""

import random, math, fractions
import time
from fractions import Fraction
from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

random.seed(2026)

# ═══════════════════════════════════════════
# 工具函数
# ═══════════════════════════════════════════
def ri(a, b):
    """随机整数 [a, b]"""
    return random.randint(a, b)

def rc(lst):
    """随机选一个"""
    return random.choice(lst)

def fmt_num(n, decimals=2):
    """格式化数字，整数去小数点"""
    if isinstance(n, int):
        return str(n)
    if isinstance(n, float):
        if n == int(n):
            return str(int(n))
        return f"{n:.{decimals}f}".rstrip('0').rstrip('.')
    return str(n)

def fmt_frac(f):
    """格式化分数"""
    if isinstance(f, Fraction):
        if f.denominator == 1:
            return str(f.numerator)
        return f"{f}"
    return fmt_num(f)

def round2(n):
    return round(n, 2)

# ═══════════════════════════════════════════
# 各题型生成器
# 每个返回 (题目字符串, 答案字符串) 列表
# ═══════════════════════════════════════════

# ---------- 1. 行程问题 (40) ----------
def gen_xingcheng(n=40):
    res = []
    names = ["小明","小红","小刚","小华","小丽","小强","小芳","小军","小玲","小伟"]
    places = ["学校","家","公园","图书馆","超市","车站","体育馆","少年宫","电影院","商场"]
    for i in range(n):
        sub = rc(names)
        place = rc([p for p in places])
        t = ri(1, 4)  # 子类型
        if t == 1:  # 基本行程
            v = ri(3, 15) * 10  # 速度
            time = ri(2, 8)     # 时间
            d = v * time
            res.append((
                f"{sub}从{place}出发，每小时行{v}千米，{time}小时到达目的地。{sub}一共行了多少千米？",
                f"{v}×{time}={d}（千米）\n答：{sub}一共行了{d}千米。"
            ))
        elif t == 2:  # 求时间 - 构造法
            v = ri(3, 12) * 10
            time = ri(2, 8)
            d = v * time
            res.append((
                f"{sub}从{place}到{place}的距离是{d}千米，每小时行{v}千米，需要几小时到达？",
                f"{d}÷{v}={time}（小时）\n答：需要{time}小时到达。"
            ))
        elif t == 3:  # 求速度 - 构造法
            v = ri(3, 15) * 10
            time = ri(2, 6)
            d = v * time
            res.append((
                f"{sub}从{place}出发，{time}小时走了{d}千米，平均每小时行多少千米？",
                f"{d}÷{time}={v}（千米/小时）\n答：平均每小时行{v}千米。"
            ))
        elif t == 4:  # 相遇问题
            v1 = ri(3, 8) * 10
            v2 = ri(3, 8) * 10
            time = ri(2, 5)
            d = (v1 + v2) * time
            n1, n2 = random.sample(names, 2)
            p1, p2 = random.sample(places, 2)
            res.append((
                f"{n1}从{p1}、{n2}从{p2}同时相向而行，{n1}每小时行{v1}千米，{n2}每小时行{v2}千米，{time}小时后两人相遇。{p1}和{p2}之间相距多少千米？",
                f"（{v1}+{v2}）×{time}={v1+v2}×{time}={d}（千米）\n答：{p1}和{p2}之间相距{d}千米。"
            ))
    return res

# ---------- 2. 工程问题 (30) ----------
def gen_gongcheng(n=30):
    res = []
    names = ["甲队","乙队","丙队","A组","B组","第一车间","第二车间","张师傅","李师傅","王师傅"]
    tasks = ["修一条路","挖一条水渠","粉刷一面墙","打扫校园","完成一批零件","铺设管道","修建花坛","装订一批书","加工一批零件","搬运一批货物"]
    for i in range(n):
        t = ri(1, 3)
        n1, n2 = random.sample(names, 2)
        task = rc(tasks)
        if t == 1:  # 合作求时间
            d1 = ri(3, 10)
            d2 = ri(4, 12)
            # 1/d1 + 1/d2 = (d1+d2)/(d1*d2), time = d1*d2/(d1+d2)
            lcm = (d1 * d2) // math.gcd(d1, d2)
            t_time = Fraction(d1 * d2, d1 + d2)
            res.append((
                f"{n1}单独{task}需要{d1}天完成，{n2}单独{task}需要{d2}天完成。两队合作几天可以完成？",
                f"1÷（1/{d1}+1/{d2}）=1÷{Fraction(d1+d2, d1*d2)}={fmt_frac(t_time)}（天）\n答：两队合作{fmt_frac(t_time)}天可以完成。"
            ))
        elif t == 2:  # 合作后剩余
            d1 = ri(5, 15)
            d2 = ri(5, 15)
            work_days = ri(2, min(d1, d2) - 1)
            remaining = Fraction(1) - Fraction(work_days) * Fraction(1, d1) - Fraction(work_days) * Fraction(1, d2)
            res.append((
                f"{n1}单独{task}需要{d1}天，{n2}单独{task}需要{d2}天。两队合作{work_days}天后，还剩全部工程的几分之几？",
                f"合作效率=1/{d1}+1/{d2}={Fraction(d1+d2, d1*d2)}\n"
                f"已完成={Fraction(work_days) * Fraction(d1+d2, d1*d2)}\n"
                f"剩余=1-{Fraction(work_days) * Fraction(d1+d2, d1*d2)}={fmt_frac(remaining)}\n"
                f"答：还剩全部工程的{fmt_frac(remaining)}。"
            ))
        else:  # 已知效率比
            r1 = ri(2, 5)
            r2 = ri(2, 5)
            while r1 == r2:
                r2 = ri(2, 5)
            total = ri(100, 500)
            # 按比分配
            v1 = total * r1 // (r1 + r2)
            v2 = total * r2 // (r1 + r2)
            while v1 + v2 != total:
                total = ri(100, 500)
                v1 = total * r1 // (r1 + r2)
                v2 = total * r2 // (r1 + r2)
            res.append((
                f"{task}共{total}个，{n1}和{n2}按{r1}:{r2}的比例分配任务，各应做多少个？",
                f"{n1}：{total}×{r1}/({r1}+{r2})={v1}（个）\n"
                f"{n2}：{total}×{r2}/({r1}+{r2})={v2}（个）\n"
                f"答：{n1}应做{v1}个，{n2}应做{v2}个。"
            ))
    return res

# ---------- 3. 浓度问题 (20) ----------
def gen_nongdu(n=20):
    res = []
    for i in range(n):
        t = ri(1, 3)
        if t == 1:  # 求浓度
            salt = ri(5, 50)
            water = ri(100, 400)
            total = salt + water
            conc = round2(salt / total * 100)
            res.append((
                f"把{salt}克盐溶解在{water}克水中，盐水的浓度是多少？（百分号前保留一位小数）",
                f"{salt}÷（{salt}+{water}）={salt}÷{total}={round2(salt/total*100)}%\n答：盐水的浓度是{conc}%。"
            ))
        elif t == 2:  # 求溶质
            # 确保 solution 是100的倍数，这样任何pct都能整除
            solution = ri(1, 5) * 100
            pct = ri(5, 30)
            salt = solution * pct // 100
            res.append((
                f"有{solution}克浓度为{pct}%的盐水，其中含盐多少克？",
                f"{solution}×{pct}%={salt}（克）\n答：其中含盐{salt}克。"
            ))
        else:  # 加水稀释
            # 构造法：先选初始盐水，再选加水量使结果为整数百分比
            solution = ri(1, 3) * 100  # 100/200/300
            c1 = ri(10, 40)
            salt = solution * c1 // 100  # 盐量（整数）
            # 找加水量使 salt*100/(solution+add_water) 为整数
            # 即 (solution+add_water) 必须是 salt*100 的因子
            target_total = salt * 100
            # 找在 [solution+50, solution+200] 范围内的因子
            add_water = None
            for divisor in range(solution + 50, solution + 201):
                if target_total % divisor == 0:
                    add_water = divisor - solution
                    break
            if add_water is None:
                # 退路：选加水量使总量为 salt*10 的倍数
                add_water = salt * 10 - solution
                if add_water < 50:
                    add_water = salt * 5 - solution
                if add_water < 50:
                    # 直接不要求整除，保留一位小数
                    add_water = ri(50, 200)
                    c2 = round2(salt * 100 / (solution + add_water))
                    res.append((
                        f"有{solution}克浓度为{c1}%的盐水，加入{add_water}克水后，盐水的浓度变为百分之几？（保留一位小数）",
                        f"含盐量={solution}×{c1}%={salt}（克）\n"
                        f"新浓度={salt}÷（{solution}+{add_water}）×100%≈{c2}%\n"
                        f"答：浓度约为{c2}%。"
                    ))
                    continue
            c2 = salt * 100 // (solution + add_water)
            res.append((
                f"有{solution}克浓度为{c1}%的盐水，加入{add_water}克水后，盐水的浓度变为多少？",
                f"含盐量={solution}×{c1}%={salt}（克）\n"
                f"新浓度={salt}÷（{solution}+{add_water}）×100%={c2}%\n"
                f"答：浓度变为{c2}%。"
            ))
    return res

# ---------- 4. 利润/折扣/经济问题 (30) ----------
def gen_lirun(n=30):
    res = []
    goods = ["书包","文具盒","篮球","足球","运动鞋","外套","衬衫","裙子","帽子","手表",
             "台灯","闹钟","水杯","雨伞","自行车"]
    for i in range(n):
        t = ri(1, 4)
        good = rc(goods)
        if t == 1:  # 求利润
            cost = ri(20, 200)
            sell = cost + ri(5, 80)
            qty = ri(5, 50)
            profit = (sell - cost) * qty
            res.append((
                f"商店购进{qty}个{good}，每个进价{cost}元，售价{sell}元，全部卖出后赚了多少元？",
                f"（{sell}-{cost}）×{qty}={sell-cost}×{qty}={profit}（元）\n答：全部卖出后赚了{profit}元。"
            ))
        elif t == 2:  # 折扣
            price = ri(50, 500)
            while price % 10 != 0:
                price = ri(50, 500)
            discount = ri(6, 9)  # 打几折
            final = price * discount // 10
            saved = price - final
            res.append((
                f"一件{good}原价{price}元，现在打{discount}折出售，现价多少元？比原价便宜了多少元？",
                f"现价={price}×{discount}0%={final}（元）\n"
                f"便宜了{price}-{final}={saved}（元）\n"
                f"答：现价{final}元，便宜了{saved}元。"
            ))
        elif t == 3:  # 利润率
            cost = ri(30, 200)
            sell = ri(cost + 10, cost + 150)
            rate = round2((sell - cost) / cost * 100)
            res.append((
                f"一个{good}进价{cost}元，售价{sell}元，利润率是百分之几？（保留一位小数）",
                f"利润率=（{sell}-{cost}）÷{cost}×100%={sell-cost}÷{cost}×100%={rate}%\n答：利润率是{rate}%。"
            ))
        else:  # 满减
            price = ri(100, 500)
            threshold = ri(200, 400)
            reduce = ri(30, 80)
            qty = ri(1, 4)
            total = price * qty
            if total >= threshold:
                final = total - reduce
                res.append((
                    f"{good}每个{price}元，买{qty}个，商场活动满{threshold}元减{reduce}元，实际应付多少元？",
                    f"总价={price}×{qty}={total}（元）\n"
                    f"{total}≥{threshold}，可减{reduce}元\n"
                    f"实付={total}-{reduce}={final}（元）\n"
                    f"答：实际应付{final}元。"
                ))
            else:
                res.append((
                    f"{good}每个{price}元，买{qty}个，商场活动满{threshold}元减{reduce}元，实际应付多少元？",
                    f"总价={price}×{qty}={total}（元）\n"
                    f"{total}<{threshold}，不满足优惠条件\n"
                    f"实付={total}元\n"
                    f"答：实际应付{total}元。"
                ))
    return res

# ---------- 5. 比例问题 (30) ----------
def gen_bili(n=30):
    res = []
    for i in range(n):
        t = ri(1, 3)
        if t == 1:  # 正比例 - 构造法
            unit_price = ri(2, 15)
            a = ri(2, 8)
            c = ri(2, 10)
            b = unit_price * a
            d = unit_price * c
            items = ["苹果","橘子","笔记本","铅笔","面包","牛奶"]
            item = rc(items)
            res.append((
                f"买{a}千克{item}需要{b}元，买{c}千克{item}需要多少元？",
                f"单价={b}÷{a}={unit_price}（元/千克）\n"
                f"总价={unit_price}×{c}={d}（元）\n"
                f"答：买{c}千克{item}需要{d}元。"
            ))
        elif t == 2:  # 反比例 - 构造法
            v1 = ri(3, 10)
            t1 = ri(4, 20)
            total = v1 * t1
            # 找total的因子作为v2
            v2 = None
            for candidate in range(2, 9):
                if total % candidate == 0:
                    v2 = candidate
                    break
            if v2 is None:
                v2 = 2  # 退路
                total = v2 * ri(5, 20)
                v1 = total // ri(4, 20)
                t1 = total // v1
            t2 = total // v2
            res.append((
                f"一辆汽车每小时行{v1}千米，{t1}小时到达。如果每小时行{v2}千米，几小时到达？",
                f"路程={v1}×{t1}={total}（千米）\n"
                f"时间={total}÷{v2}={t2}（小时）\n"
                f"答：每小时行{v2}千米，{t2}小时到达。"
            ))
        else:  # 比例尺
            scale = rc([1000, 2000, 5000, 10000, 50000, 100000])
            map_dist = ri(2, 20)
            real_dist = map_dist * scale  # 厘米
            if real_dist >= 100000:
                real_km = real_dist / 100000
                res.append((
                    f"在比例尺1:{scale}的地图上，量得两地距离为{map_dist}厘米，两地实际距离是多少千米？",
                    f"实际距离={map_dist}×{scale}={real_dist}（厘米）={real_km}（千米）\n答：两地实际距离是{real_km}千米。"
                ))
            else:
                real_m = real_dist / 100
                res.append((
                    f"在比例尺1:{scale}的地图上，量得两地距离为{map_dist}厘米，两地实际距离是多少米？",
                    f"实际距离={map_dist}×{scale}={real_dist}（厘米）={real_m}（米）\n答：两地实际距离是{real_m}米。"
                ))
    return res

# ---------- 6. 年龄问题 (20) ----------
def gen_nianling(n=20):
    res = []
    names = [("小明","爸爸"),("小红","妈妈"),("小刚","爷爷"),("小华","奶奶"),
            ("小丽","爸爸"),("小强","妈妈"),("小芳","哥哥"),("小军","姐姐"),
            ("小伟","爸爸"),("小玲","妈妈")]
    for i in range(n):
        t = ri(1, 3)
        child, elder = rc(names)
        if t == 1:  # 几年后几倍
            age_c = ri(5, 12)
            age_e = ri(age_c + 18, age_c + 35)
            # n年后 elder+n = k*(child+n)
            # 尝试找到合理的k
            found = False
            for years in range(1, 20):
                nc = age_c + years
                ne = age_e + years
                if ne % nc == 0:
                    k = ne // nc
                    if 2 <= k <= 5:
                        res.append((
                            f"{child}今年{age_c}岁，{elder}今年{age_e}岁。几年后{elder}的年龄是{child}的{k}倍？",
                            f"设x年后\n{age_e}+x={k}×（{age_c}+x）\n{age_e}+x={k*age_c}+{k}x\n"
                            f"{age_e-k*age_c}={k-1}x\nx={years}\n"
                            f"答：{years}年后{elder}的年龄是{child}的{k}倍。"
                        ))
                        found = True
                        break
            if not found:
                # 简单版本：求年龄差
                diff = age_e - age_c
                res.append((
                    f"{child}今年{age_c}岁，{elder}今年{age_e}岁，{elder}比{child}大多少岁？10年后{elder}比{child}大多少岁？",
                    f"年龄差={age_e}-{age_c}={diff}（岁）\n"
                    f"年龄差不变，10后仍大{diff}岁\n"
                    f"答：{elder}比{child}大{diff}岁，10年后仍大{diff}岁。"
                ))
        elif t == 2:  # 年龄和
            age_c = ri(6, 14)
            age_e = ri(age_c + 20, age_c + 32)
            total = age_c + age_e
            res.append((
                f"{child}和{elder}的年龄之和是{total}岁，{elder}比{child}大{age_e-age_c}岁，两人各多少岁？",
                f"{elder}=（{total}+{age_e-age_c}）÷2={age_e}（岁）\n"
                f"{child}=（{total}-{age_e-age_c}）÷2={age_c}（岁）\n"
                f"答：{child}{age_c}岁，{elder}{age_e}岁。"
            ))
        else:  # 倍数关系
            age_c = ri(5, 10)
            k = ri(3, 6)
            age_e = age_c * k
            res.append((
                f"{child}今年{age_c}岁，{elder}的年龄是{child}的{k}倍，{elder}今年多少岁？{elder}比{child}大多少岁？",
                f"{elder}年龄={age_c}×{k}={age_e}（岁）\n"
                f"年龄差={age_e}-{age_c}={age_e-age_c}（岁）\n"
                f"答：{elder}今年{age_e}岁，比{child}大{age_e-age_c}岁。"
            ))
    return res

# ---------- 7. 鸡兔同笼 (20) ----------
def gen_jitu(n=20):
    res = []
    scenarios = [
        ("鸡","兔","只","头","条腿", 2, 4),
        ("小猫","鹦鹉","只","只","条腿", 4, 2),
        ("三轮车","小汽车","辆","辆","个轮子", 3, 4),
        ("大船","小船","条","条","人", 6, 4),
        ("5元票","2元票","张","张","元", 5, 2),
        ("男生","女生","人","人","人", 0, 0),  # 特殊处理
    ]
    for i in range(n):
        sc = rc(scenarios)
        a_name, b_name, unit1, unit2, unit2_name, default_a, default_b = sc
        
        if a_name == "男生":
            # 特殊：男女生人数问题（不是鸡兔同笼型）
            total = ri(30, 60)
            diff = ri(2, 10)
            while (total + diff) % 2 != 0:
                diff = ri(2, 10)
            boys = (total + diff) // 2
            girls = (total - diff) // 2
            res.append((
                f"全班{total}人，男生比女生多{diff}人，男生和女生各多少人？",
                f"男生=（{total}+{diff}）÷2={boys}（人）\n"
                f"女生=（{total}-{diff}）÷2={girls}（人）\n"
                f"答：男生{boys}人，女生{girls}人。"
            ))
            continue
        
        a_legs = default_a
        b_legs = default_b
        # 构造法：先定答案再算总数
        a_count = ri(3, 12)
        b_count = ri(3, 12)
        total_head = a_count + b_count
        total_legs = a_legs * a_count + b_legs * b_count
        
        if a_name in ("三轮车", "大船"):
            res.append((
                f"停车场有{a_name}和{b_name}共{total_head}{unit1}，数{unit2_name}共{total_legs}，{a_name}和{b_name}各有多少？",
                f"假设全是{b_name}：{b_legs}×{total_head}={b_legs*total_head}\n"
                f"差额：{total_legs}-{b_legs*total_head}={total_legs-b_legs*total_head}\n"
                f"每{unit1}差额：{a_legs}-{b_legs}={a_legs-b_legs}\n"
                f"{a_name}：{total_legs-b_legs*total_head}÷{a_legs-b_legs}={a_count}（{unit1}）\n"
                f"{b_name}：{total_head}-{a_count}={b_count}（{unit1}）\n"
                f"答：{a_name}{a_count}{unit1}，{b_name}{b_count}{unit1}。"
            ))
        elif a_name == "5元票":
            res.append((
                f"买了{a_name}和{b_name}共{total_head}{unit1}，一共花了{total_legs}元，两种票各买了多少张？",
                f"假设全是{b_name}：{b_legs}×{total_head}={b_legs*total_head}元\n"
                f"多花：{total_legs}-{b_legs*total_head}={total_legs-b_legs*total_head}元\n"
                f"每张差额：{a_legs}-{b_legs}={a_legs-b_legs}元\n"
                f"{a_name}：{total_legs-b_legs*total_head}÷{a_legs-b_legs}={a_count}（张）\n"
                f"{b_name}：{total_head}-{a_count}={b_count}（张）\n"
                f"答：{a_name}{a_count}张，{b_name}{b_count}张。"
            ))
        else:
            res.append((
                f"笼子里有{a_name}和{b_name}共{total_head}{unit1}，数{unit2_name}共{total_legs}，{a_name}和{b_name}各有多少{unit1}？",
                f"假设全是{b_name}：{b_legs}×{total_head}={b_legs*total_head}\n"
                f"差额：{total_legs}-{b_legs*total_head}={total_legs-b_legs*total_head}\n"
                f"每{unit1}差额：{a_legs}-{b_legs}={a_legs-b_legs}\n"
                f"{a_name}：{total_legs-b_legs*total_head}÷{a_legs-b_legs}={a_count}（{unit1}）\n"
                f"{b_name}：{total_head}-{a_count}={b_count}（{unit1}）\n"
                f"答：{a_name}{a_count}{unit1}，{b_name}{b_count}{unit1}。"
            ))
    return res

# ---------- 8. 植树问题 (15) ----------
def gen_zhishu(n=15):
    res = []
    for i in range(n):
        t = ri(1, 3)
        interval = ri(4, 20)
        seg = ri(3, 20)
        length = interval * seg
        if t == 1:  # 两端都栽
            trees = seg + 1
            res.append((
                f"在一条长{length}米的小路一侧植树，每隔{interval}米栽一棵，两端都栽，共需多少棵树苗？",
                f"间隔数={length}÷{interval}={seg}\n"
                f"棵数={seg}+1={trees}（棵）\n"
                f"答：共需{trees}棵树苗。"
            ))
        elif t == 2:  # 只栽一端
            trees = seg
            res.append((
                f"在一条长{length}米的小路一侧植树，每隔{interval}米栽一棵，只栽一端，共需多少棵树苗？",
                f"间隔数={length}÷{interval}={seg}\n"
                f"棵数={seg}（棵）\n"
                f"答：共需{trees}棵树苗。"
            ))
        else:  # 两端都不栽
            trees = seg - 1
            res.append((
                f"在一条长{length}米的小路一侧植树，每隔{interval}米栽一棵，两端都不栽，共需多少棵树苗？",
                f"间隔数={length}÷{interval}={seg}\n"
                f"棵数={seg}-1={trees}（棵）\n"
                f"答：共需{trees}棵树苗。"
            ))
    return res

# ---------- 9. 盈亏问题 (15) ----------
def gen_yingkui(n=15):
    res = []
    for i in range(n):
        t = ri(1, 2)
        if t == 1:  # 一盈一亏
            n1 = ri(3, 8)  # 每人分n1个多a个
            n2 = ri(2, n1-1)  # 每人分n2个少b个
            a = ri(2, 10)
            b = ri(2, 10)
            people = (a + b) // (n1 - n2)
            while people * (n1 - n2) != a + b:
                a = ri(2, 10)
                b = ri(2, 10)
                people = (a + b) // (n1 - n2)
            total = people * n1 - a
            items = ["苹果","桃子","饼干","糖果","铅笔"]
            item = rc(items)
            res.append((
                f"把一些{item}分给小朋友，每人分{n1}个，多出{a}个；每人分{n2}个，少了{b}个。有几个小朋友？一共有多少个{item}？",
                f"人数=（{a}+{b}）÷（{n1}-{n2}）={a+b}÷{n1-n2}={people}（人）\n"
                f"总数={people}×{n1}-{a}={total}（个）\n"
                f"答：有{people}个小朋友，一共有{total}个{item}。"
            ))
        else:  # 两盈
            n1 = ri(5, 10)
            n2 = ri(2, n1 - 1)
            a = ri(1, 8)
            b = ri(a + 1, a + 10)
            diff = b - a
            people = diff // (n1 - n2)
            while people * (n1 - n2) != diff:
                b = ri(a + 1, a + 10)
                diff = b - a
                people = diff // (n1 - n2)
            total = people * n1 - a
            res.append((
                f"把一些糖果分给同学，每人分{n1}颗多{a}颗，每人分{n2}颗多{b}颗，有几个同学？一共有多少颗糖果？",
                f"人数=（{b}-{a}）÷（{n1}-{n2}）={diff}÷{n1-n2}={people}（人）\n"
                f"总数={people}×{n1}-{a}={total}（颗）\n"
                f"答：有{people}个同学，一共有{total}颗糖果。"
            ))
    return res

# ---------- 10. 和差问题 (20) ----------
def gen_hecha(n=20):
    res = []
    for i in range(n):
        a_plus_b = ri(50, 300)
        diff = ri(5, 50)
        while (a_plus_b + diff) % 2 != 0:
            diff = ri(5, 50)
        big = (a_plus_b + diff) // 2
        small = (a_plus_b - diff) // 2
        scenarios = [
            ("甲班","乙班","人数"),
            ("果园里苹果树","梨树","棵数"),
            ("两个数","它们","值"),
            ("长方形相邻两边","它们","长度"),
            ("哥哥","弟弟","零花钱"),
        ]
        sc = rc(scenarios)
        if sc[2] == "人数":
            res.append((
                f"{sc[0]}和{sc[1]}共{a_plus_b}人，{sc[0]}比{sc[1]}多{diff}人，两{sc[2]}各是多少？",
                f"大数=（{a_plus_b}+{diff}）÷2={big}\n"
                f"小数=（{a_plus_b}-{diff}）÷2={small}\n"
                f"答：{sc[0]}{big}人，{sc[1]}{small}人。"
            ))
        elif sc[2] == "棵数":
            res.append((
                f"{sc[0]}和{sc[1]}共{a_plus_b}{sc[2][:1]}，{sc[0]}比{sc[1]}多{diff}{sc[2][:1]}，两种树各多少{sc[2][:1]}？",
                f"多的一种=（{a_plus_b}+{diff}）÷2={big}（{sc[2][:1]}）\n"
                f"少的一种=（{a_plus_b}-{diff}）÷2={small}（{sc[2][:1]}）\n"
                f"答：{sc[0]}{big}{sc[2][:1]}，{sc[1]}{small}{sc[2][:1]}。"
            ))
        elif sc[2] == "值":
            res.append((
                f"两个数的和是{a_plus_b}，大数比小数多{diff}，这两个数分别是多少？",
                f"大数=（{a_plus_b}+{diff}）÷2={big}\n"
                f"小数=（{a_plus_b}-{diff}）÷2={small}\n"
                f"答：大数是{big}，小数是{small}。"
            ))
        elif sc[2] == "长度":
            res.append((
                f"长方形的周长是{a_plus_b*2}厘米，长比宽多{diff}厘米，长和宽各是多少？",
                f"长+宽={a_plus_b*2}÷2={a_plus_b}（厘米）\n"
                f"长=（{a_plus_b}+{diff}）÷2={big}（厘米）\n"
                f"宽=（{a_plus_b}-{diff}）÷2={small}（厘米）\n"
                f"答：长{big}厘米，宽{small}厘米。"
            ))
        else:
            res.append((
                f"{sc[0]}和{sc[1]}的{sc[2]}一共{a_plus_b}元，{sc[0]}比{sc[1]}多{diff}元，两人各有多少元？",
                f"{sc[0]}=（{a_plus_b}+{diff}）÷2={big}（元）\n"
                f"{sc[1]}=（{a_plus_b}-{diff}）÷2={small}（元）\n"
                f"答：{sc[0]}{big}元，{sc[1]}{small}元。"
            ))
    return res

# ---------- 11. 和倍/差倍问题 (20) ----------
def gen_hebeicha(n=20):
    res = []
    for i in range(n):
        t = ri(1, 2)
        if t == 1:  # 和倍
            small = ri(5, 30)
            k = ri(2, 6)
            big = small * k
            total = small + big
            items = [("白兔","黑兔","只"),("苹果","梨","个"),("男生","女生","人"),("语文书","数学书","本")]
            a_n, b_n, unit = rc(items)
            res.append((
                f"{a_n}和{b_n}共{total}{unit}，{a_n}的数量是{b_n}的{k}倍，两种各多少？",
                f"{b_n}=（小数）{total}÷（{k}+1）={total}÷{k+1}={small}（{unit}）\n"
                f"{a_n}=（大数）{small}×{k}={big}（{unit}）\n"
                f"答：{a_n}{big}{unit}，{b_n}{small}{unit}。"
            ))
        else:  # 差倍
            small = ri(5, 25)
            k = ri(2, 5)
            big = small * k
            diff = big - small
            items = [("足球","篮球","个"),("红花","黄花","朵"),("大米","面粉","千克")]
            a_n, b_n, unit = rc(items)
            res.append((
                f"{a_n}比{b_n}多{diff}{unit}，{a_n}是{b_n}的{k}倍，两种各多少？",
                f"{b_n}=（小数）{diff}÷（{k}-1）={diff}÷{k-1}={small}（{unit}）\n"
                f"{a_n}={small}×{k}={big}（{unit}）\n"
                f"答：{a_n}{big}{unit}，{b_n}{small}{unit}。"
            ))
    return res

# ---------- 12. 归一问题 (20) ----------
def gen_guiyi(n=20):
    res = []
    for i in range(n):
        t = ri(1, 2)
        items = [("钢笔","支","元"),("大米","千克","元"),("布料","米","元"),
                 ("牛奶","瓶","元"),("水泥","吨","元"),("煤","吨","元")]
        item, unit1, unit2 = rc(items)
        qty1 = ri(2, 8)
        price1 = qty1 * ri(3, 20)
        unit_price = price1 // qty1
        if t == 1:  # 正归一
            qty2 = ri(3, 15)
            price2 = unit_price * qty2
            res.append((
                f"买{qty1}{unit1}{item}需要{price1}{unit2}，买{qty2}{unit1}{item}需要多少{unit2}？",
                f"单价={price1}÷{qty1}={unit_price}（{unit2}/{unit1}）\n"
                f"总价={unit_price}×{qty2}={price2}（{unit2}）\n"
                f"答：买{qty2}{unit1}{item}需要{price2}{unit2}。"
            ))
        else:  # 反归一
            total = ri(50, 300)
            while total % unit_price != 0:
                total = ri(50, 300)
            qty2 = total // unit_price
            res.append((
                f"买{qty1}{unit1}{item}需要{price1}{unit2}，{total}{unit2}可以买多少{unit1}{item}？",
                f"单价={price1}÷{qty1}={unit_price}（{unit2}/{unit1}）\n"
                f"数量={total}÷{unit_price}={qty2}（{unit1}）\n"
                f"答：{total}{unit2}可以买{qty2}{unit1}{item}。"
            ))
    return res

# ---------- 13. 分数/百分数应用 (35) ----------
def gen_fenshu(n=35):
    res = []
    for i in range(n):
        t = ri(1, 5)
        if t == 1:  # 求一个数的几分之几 - 构造法
            denom = ri(2, 8)
            numer = ri(1, denom - 1)
            frac = Fraction(numer, denom)
            # total必须是denom的倍数
            k = ri(5, 25)
            total = k * denom
            result = total * numer // denom
            res.append((
                f"学校图书馆有{total}本故事书，其中{fmt_frac(frac)}是科技书，科技书有多少本？",
                f"{total}×{fmt_frac(frac)}={result}（本）\n答：科技书有{result}本。"
            ))
        elif t == 2:  # 已知几分之几求总数 - 构造法
            denom = ri(2, 6)
            numer = ri(1, denom - 1)
            frac = Fraction(numer, denom)
            # part必须是numer的倍数
            k = ri(2, 15)
            part = k * numer
            total = part * denom // numer  # = k * denom
            res.append((
                f"小明看一本书，已经看了{part}页，正好占全书的{fmt_frac(frac)}，这本书共多少页？",
                f"{part}÷{fmt_frac(frac)}={total}（页）\n答：这本书共{total}页。"
            ))
        elif t == 3:  # 百分数应用 - 出勤率
            total = ri(40, 60)
            absent = ri(0, 5)
            present = total - absent
            rate = round2(present / total * 100)
            res.append((
                f"六年级有{total}名学生，今天缺席{absent}人，今天的出勤率是多少？",
                f"出勤人数={total}-{absent}={present}（人）\n"
                f"出勤率={present}÷{total}×100%={rate}%\n"
                f"答：今天的出勤率是{rate}%。"
            ))
        elif t == 4:  # 增加/减少百分之几
            old = ri(50, 300)
            change = ri(5, 50)
            t_dir = ri(1, 2)
            if t_dir == 1:
                new = old + change
                pct = round2(change / old * 100)
                res.append((
                    f"某商品原来售价{old}元，现在涨价到{new}元，涨价了百分之几？",
                    f"涨价={new}-{old}={change}（元）\n"
                    f"涨价百分比={change}÷{old}×100%={pct}%\n"
                    f"答：涨价了{pct}%。"
                ))
            else:
                new = old - change
                pct = round2(change / old * 100)
                res.append((
                    f"某商品原来售价{old}元，现在降价到{new}元，降价了百分之几？",
                    f"降价={old}-{new}={change}（元）\n"
                    f"降价百分比={change}÷{old}×100%={pct}%\n"
                    f"答：降价了{pct}%。"
                ))
        else:  # 分数比较
            b = ri(3, 8)
            d = ri(3, 8)
            a = ri(1, b - 1)
            c = ri(1, d - 1)
            f1 = Fraction(a, b)
            f2 = Fraction(c, d)
            diff = abs(f1 - f2)
            if f1 > f2:
                res.append((
                    f"甲绳长{fmt_frac(f1)}米，乙绳长{fmt_frac(f2)}米，甲绳比乙绳长多少米？",
                    f"{fmt_frac(f1)}-{fmt_frac(f2)}={fmt_frac(diff)}（米）\n答：甲绳比乙绳长{fmt_frac(diff)}米。"
                ))
            else:
                res.append((
                    f"甲绳长{fmt_frac(f1)}米，乙绳长{fmt_frac(f2)}米，乙绳比甲绳长多少米？",
                    f"{fmt_frac(f2)}-{fmt_frac(f1)}={fmt_frac(diff)}（米）\n答：乙绳比甲绳长{fmt_frac(diff)}米。"
                ))
    return res

# ---------- 14. 几何应用 (30) ----------
def gen_jihe(n=30):
    res = []
    for i in range(n):
        t = ri(1, 6)
        if t == 1:  # 长方形面积/周长
            l = ri(5, 30)
            w = ri(3, 20)
            area = l * w
            peri = 2 * (l + w)
            res.append((
                f"一块长方形菜地，长{l}米，宽{w}米。这块菜地的面积是多少平方米？围栏一周需要多少米？",
                f"面积={l}×{w}={area}（平方米）\n"
                f"周长=2×（{l}+{w}）={peri}（米）\n"
                f"答：面积{area}平方米，围栏{peri}米。"
            ))
        elif t == 2:  # 正方形
            s = ri(4, 25)
            area = s * s
            peri = 4 * s
            res.append((
                f"一个正方形花坛，边长{s}米。花坛面积是多少？护栏长多少米？",
                f"面积={s}×{s}={area}（平方米）\n"
                f"周长=4×{s}={peri}（米）\n"
                f"答：面积{area}平方米，护栏{peri}米。"
            ))
        elif t == 3:  # 三角形
            base = ri(4, 20) * 2  # 确保面积整数
            height = ri(3, 15)
            area = base * height // 2
            res.append((
                f"一块三角形铁皮，底{base}厘米，高{height}厘米，面积是多少平方厘米？",
                f"面积={base}×{height}÷2={area}（平方厘米）\n答：面积{area}平方厘米。"
            ))
        elif t == 4:  # 平行四边形
            base = ri(5, 20)
            height = ri(3, 15)
            area = base * height
            res.append((
                f"一块平行四边形玻璃，底{base}分米，高{height}分米，面积是多少平方分米？",
                f"面积={base}×{height}={area}（平方分米）\n答：面积{area}平方分米。"
            ))
        elif t == 5:  # 梯形
            a = ri(3, 10)
            b = ri(a + 2, a + 15)
            h = ri(3, 12) * 2  # 确保面积整数
            area = (a + b) * h // 2
            res.append((
                f"一块梯形田地，上底{a}米，下底{b}米，高{h}米，面积是多少平方米？",
                f"面积=（{a}+{b}）×{h}÷2={a+b}×{h}÷2={area}（平方米）\n答：面积{area}平方米。"
            ))
        elif t == 6:  # 圆
            r = ri(2, 15)
            area = round2(3.14 * r * r)
            circ = round2(2 * 3.14 * r)
            res.append((
                f"一个圆形花坛，半径{r}米，面积是多少平方米？周长是多少米？（π取3.14）",
                f"面积=3.14×{r}²=3.14×{r*r}={area}（平方米）\n"
                f"周长=2×3.14×{r}={circ}（米）\n"
                f"答：面积{area}平方米，周长{circ}米。"
            ))
    return res

# ---------- 15. 平均数问题 (20) ----------
def gen_pingjun(n=20):
    res = []
    for i in range(n):
        t = ri(1, 3)
        if t == 1:  # 基本平均数
            n_count = ri(4, 8)
            scores = [ri(60, 100) for _ in range(n_count)]
            total = sum(scores)
            avg = round2(total / n_count)
            scores_str = "、".join(str(s) for s in scores)
            res.append((
                f"小明{n_count}次测验的成绩分别是{scores_str}分，平均每次测验多少分？",
                f"总分={'+' .join(str(s) for s in scores)}={total}\n"
                f"平均={total}÷{n_count}={avg}（分）\n"
                f"答：平均每次测验{avg}分。"
            ))
        elif t == 2:  # 已知平均数求某值
            n_count = ri(4, 6)
            avg = ri(70, 95)
            total = avg * n_count
            known = [ri(60, 100) for _ in range(n_count - 1)]
            known_sum = sum(known)
            last = total - known_sum
            while last < 0 or last > 100:
                known = [ri(60, 100) for _ in range(n_count - 1)]
                known_sum = sum(known)
                last = total - known_sum
            known_str = "、".join(str(s) for s in known)
            res.append((
                f"{n_count}个同学的平均分是{avg}分，其中前{n_count-1}个同学的成绩分别是{known_str}分，最后一个同学得了多少分？",
                f"总分={avg}×{n_count}={total}\n"
                f"前{n_count-1}人总分={'+' .join(str(s) for s in known)}={known_sum}\n"
                f"最后一人={total}-{known_sum}={last}（分）\n"
                f"答：最后一个同学得了{last}分。"
            ))
        else:  # 加权平均
            g1 = ri(80, 95)
            g2 = ri(70, 85)
            n1 = ri(20, 40)
            n2 = ri(20, 40)
            total = g1 * n1 + g2 * n2
            avg = round2(total / (n1 + n2))
            res.append((
                f"甲班{n1}人平均分{g1}分，乙班{n2}人平均分{g2}分，两班合在一起平均分是多少？（保留一位小数）",
                f"甲班总分={g1}×{n1}={g1*n1}\n"
                f"乙班总分={g2}×{n2}={g2*n2}\n"
                f"总分={g1*n1}+{g2*n2}={total}\n"
                f"平均={total}÷（{n1}+{n2}）={total}÷{n1+n2}={avg}（分）\n"
                f"答：两班合在一起平均分是{avg}分。"
            ))
    return res

# ---------- 16. 还原问题 (15) ----------
def gen_huanyuan(n=15):
    res = []
    for i in range(n):
        t = ri(1, 3)
        if t == 1:  # 倒推
            final = ri(20, 80)
            step1 = ri(3, 15)
            step2 = ri(2, 8)
            # 正向：x - step1 - step2 = final => x = final + step1 + step2
            original = final + step1 + step2
            res.append((
                f"小明有一些糖果，先给了小红{step1}颗，又给了小华{step2}颗，最后还剩{final}颗。小明原来有多少颗糖果？",
                f"原来={final}+{step1}+{step2}={original}（颗）\n答：小明原来有{original}颗糖果。"
            ))
        elif t == 2:  # 倒推（乘除）
            original = ri(5, 20)
            add = ri(5, 20)
            multiply = ri(2, 5)
            result = (original + add) * multiply
            res.append((
                f"一个数加上{add}，再乘以{multiply}，结果是{result}。这个数是多少？",
                f"逆运算：{result}÷{multiply}={result//multiply if result%multiply==0 else result/multiply}\n"
                f"{result//multiply if result%multiply==0 else result/multiply}-{add}={original}\n"
                f"答：这个数是{original}。"
            ))
        else:  # 连续操作
            original = ri(30, 100)
            if original % 2 != 0:
                original = original + 1
            half = original // 2
            sold2 = ri(3, min(10, half - 1))
            remain = half - sold2
            res.append((
                f"一筐苹果，第一次卖出一半，第二次又卖出{sold2}千克，还剩{remain}千克。原来有多少千克？（提示：倒推）",
                f"第二次卖出后剩{remain}千克\n"
                f"卖出一半后剩{remain}+{sold2}={half}千克\n"
                f"原来={half}×2={original}（千克）\n"
                f"答：原来有{original}千克。"
            ))
    return res

# ---------- 17. 统计问题 (15) ----------
def gen_tongji(n=15):
    res = []
    for i in range(n):
        t = ri(1, 3)
        if t == 1:  # 求中位数
            n_count = ri(5, 9)
            if n_count % 2 == 0:
                n_count += 1  # 确保奇数
            data = sorted([ri(50, 100) for _ in range(n_count)])
            median = data[n_count // 2]
            data_str = "、".join(str(d) for d in data)
            res.append((
                f"以下是{n_count}名同学的跳绳成绩（次/分钟）：{data_str}。这组数据的中位数是多少？",
                f"数据已排序，共{n_count}个，中位数是第{n_count//2+1}个\n"
                f"中位数={median}\n"
                f"答：这组数据的中位数是{median}。"
            ))
        elif t == 2:  # 求众数
            mode = ri(70, 95)
            n_count = ri(7, 10)
            data = [mode] * ri(2, 3)
            while len(data) < n_count:
                v = ri(60, 100)
                if v != mode:
                    data.append(v)
            random.shuffle(data)
            data_str = "、".join(str(d) for d in data)
            res.append((
                f"以下是{n_count}名同学的体重（千克）：{data_str}。这组数据的众数是多少？",
                f"出现次数最多的是{mode}，出现了{data.count(mode)}次\n"
                f"众数={mode}\n"
                f"答：这组数据的众数是{mode}。"
            ))
        else:  # 读表题
            categories = ["语文","数学","英语","科学","体育"]
            cats = random.sample(categories, 3)
            scores = [ri(70, 98) for _ in range(3)]
            avg = round2(sum(scores) / 3)
            res.append((
                f"小明期末考试成绩：{'、'.join(f'{c}{s}分' for c,s in zip(cats,scores))}。三科平均分是多少？哪科最高？哪科最低？",
                f"总分={'+' .join(str(s) for s in scores)}={sum(scores)}\n"
                f"平均={sum(scores)}÷3={avg}（分）\n"
                f"最高={max(scores)}分（{cats[scores.index(max(scores))]}）\n"
                f"最低={min(scores)}分（{cats[scores.index(min(scores))]}）\n"
                f"答：平均分{avg}分，{cats[scores.index(max(scores))]}最高，{cats[scores.index(min(scores))]}最低。"
            ))
    return res

# ---------- 18. 时间/日期问题 (15) ----------
def gen_shijian(n=15):
    res = []
    for i in range(n):
        t = ri(1, 3)
        if t == 1:  # 经过时间
            h1 = ri(6, 14)
            m1 = ri(0, 59)
            h2 = ri(h1 + 1, min(h1 + 6, 22))
            m2 = ri(0, 59)
            elapsed_h = h2 - h1
            elapsed_m = m2 - m1
            if elapsed_m < 0:
                elapsed_h -= 1
                elapsed_m += 60
            res.append((
                f"小明{h1}时{m1}分出发上学，{h2}时{m2}分到校，他在路上花了多长时间？",
                f"从{h1}:{m1:02d}到{h2}:{m2:02d}\n"
                f"={elapsed_h}小时{elapsed_m}分钟\n"
                f"答：在路上花了{elapsed_h}小时{elapsed_m}分钟。"
            ))
        elif t == 2:  # 日期推算
            month = ri(1, 11)
            day = ri(10, 25)
            days_later = ri(7, 30)
            # 简化计算
            days_in_month = [0,31,28,31,30,31,30,31,31,30,31,30,31]
            new_day = day + days_later
            new_month = month
            while new_day > days_in_month[new_month]:
                new_day -= days_in_month[new_month]
                new_month += 1
                if new_month > 12:
                    new_month = 1
            res.append((
                f"今天是{month}月{day}日，{days_later}天后是几月几日？",
                f"{month}月{day}日+{days_later}天\n"
                f"={new_month}月{new_day}日\n"
                f"答：{days_later}天后是{new_month}月{new_day}日。"
            ))
        else:  # 周期问题
            period = ri(3, 7)
            weeks = ri(2, 5)
            remainder = (weeks * 7) % period
            res.append((
                f"一串彩灯按{period}个一组循环排列（红、黄、蓝、绿、紫……），第{weeks*7}个彩灯是什么颜色？",
                f"{weeks*7}÷{period}={weeks*7//period}……{remainder}\n"
                f"答：第{weeks*7}个彩灯是第{remainder if remainder else period}种颜色。"
            ))
    return res

# ---------- 19. 利息/税率问题 (15) ----------
def gen_lixishuilv(n=15):
    res = []
    for i in range(n):
        t = ri(1, 3)
        if t == 1:  # 利息
            principal = ri(1000, 10000)
            while principal % 100 != 0:
                principal = ri(1000, 10000)
            rate = rc([1.5, 2.0, 2.5, 3.0, 3.5])
            years = ri(1, 3)
            interest = principal * rate / 100 * years
            total = principal + interest
            res.append((
                f"爸爸存入银行{principal}元，年利率{rate}%，存{years}年，到期后利息是多少元？本息一共多少元？",
                f"利息={principal}×{rate}%×{years}={interest}（元）\n"
                f"本息={principal}+{interest}={total}（元）\n"
                f"答：到期利息{interest}元，本息一共{total}元。"
            ))
        elif t == 2:  # 税率
            income = ri(5000, 10000)
            while income % 100 != 0:
                income = ri(5000, 10000)
            tax_rate = ri(3, 10)
            tax = income * tax_rate // 100
            after = income - tax
            res.append((
                f"某人的劳务报酬为{income}元，按{tax_rate}%的税率缴税，税后实际收入多少元？",
                f"税额={income}×{tax_rate}%={tax}（元）\n"
                f"税后收入={income}-{tax}={after}（元）\n"
                f"答：税后实际收入{after}元。"
            ))
        else:  # 增值税
            cost = ri(100, 1000)
            while cost % 10 != 0:
                cost = ri(100, 1000)
            markup_rate = ri(10, 50)
            sell = cost * (100 + markup_rate) // 100
            vat_rate = ri(3, 13)
            vat = sell * vat_rate // 100
            res.append((
                f"一件商品成本{cost}元，加价{markup_rate}%出售，再按售价的{vat_rate}%缴纳增值税。售价多少元？增值税多少元？",
                f"售价={cost}×（1+{markup_rate}%）={cost}×{100+markup_rate}/100={sell}（元）\n"
                f"增值税={sell}×{vat_rate}%={vat}（元）\n"
                f"答：售价{sell}元，增值税{vat}元。"
            ))
    return res

# ---------- 20. 综合应用题 (30) ----------
def gen_zonghe(n=30):
    res = []
    names = ["小明","小红","小刚","小华","小丽","小强"]
    for i in range(n):
        t = ri(1, 6)
        name = rc(names)
        if t == 1:  # 购物找零
            items_prices = [("笔记本", ri(3,8)), ("铅笔", ri(1,3)), ("橡皮", ri(1,3)), ("尺子", ri(2,5))]
            chosen = random.sample(items_prices, ri(2,3))
            total = sum(p for _, p in chosen)
            qty = ri(2, 5)
            grand_total = total * qty
            paid = ri(grand_total + 10, grand_total + 50)
            while paid % 10 != 0:
                paid = ri(grand_total + 10, grand_total + 50)
            change = paid - grand_total
            items_str = "、".join(f"{n}{p}元" for n,p in chosen)
            res.append((
                f"{name}买了{'各' if qty==1 else ''}{qty}份文具，每份包含{items_str}。付了{paid}元，应找回多少元？",
                f"每份={'+' .join(str(p) for _,p in chosen)}={total}（元）\n"
                f"总价={total}×{qty}={grand_total}（元）\n"
                f"找零={paid}-{grand_total}={change}（元）\n"
                f"答：应找回{change}元。"
            ))
        elif t == 2:  # 铺砖问题
            room_l = ri(4, 10)
            room_w = ri(3, 8)
            tile_side = ri(2, 5)
            room_area_dm = room_l * 10 * room_w * 10  # 平方分米
            tile_area = tile_side * tile_side
            while room_area_dm % tile_area != 0:
                tile_side = ri(2, 5)
                tile_area = tile_side * tile_side
            tiles = room_area_dm // tile_area
            res.append((
                f"一间房间长{room_l}米、宽{room_w}米，用边长{tile_side}分米的正方形地砖铺满地面，需要多少块地砖？",
                f"房间面积={room_l}×{room_w}={room_l*room_w}（平方米）={room_area_dm}（平方分米）\n"
                f"地砖面积={tile_side}×{tile_side}={tile_area}（平方分米）\n"
                f"需要{room_area_dm}÷{tile_area}={tiles}（块）\n"
                f"答：需要{tiles}块地砖。"
            ))
        elif t == 3:  # 注水/排水 - 构造法
            fill_rate = ri(3, 10)
            drain_rate = ri(1, fill_rate - 1)
            net = fill_rate - drain_rate
            hours = ri(5, 20)
            pool_size = net * hours
            res.append((
                f"一个水池容量{pool_size}吨，进水管每小时注{fill_rate}吨，出水管每小时排水{drain_rate}吨。两管同时开，几小时注满？",
                f"每小时净注水={fill_rate}-{drain_rate}={net}（吨）\n"
                f"时间={pool_size}÷{net}={hours}（小时）\n"
                f"答：{hours}小时注满。"
            ))
        elif t == 4:  # 鸡兔同笼变式 - 租车
            big_cap = ri(40, 55)
            small_cap = ri(20, 35)
            total_people = ri(100, 300)
            # 尝试找整数解
            found = False
            for big_n in range(1, total_people // big_cap + 1):
                remaining = total_people - big_n * big_cap
                if remaining > 0 and remaining % small_cap == 0:
                    small_n = remaining // small_cap
                    res.append((
                        f"学校组织{total_people}人出游，大车每辆坐{big_cap}人，小车每辆坐{small_cap}人。用大车{big_n}辆、小车{small_n}辆刚好坐满，大车和小车各几辆？（提示：列方程）",
                        f"大车{big_n}辆×{big_cap}人={big_n*big_cap}人\n"
                        f"剩余{total_people}-{big_n*big_cap}={remaining}人\n"
                        f"小车={remaining}÷{small_cap}={small_n}辆\n"
                        f"答：大车{big_n}辆，小车{small_n}辆。"
                    ))
                    found = True
                    break
            if not found:
                # 简单版
                total_people = big_cap * 2 + small_cap * 3
                res.append((
                    f"学校组织出游，大车每辆坐{big_cap}人，小车每辆坐{small_cap}人。用了大车2辆、小车3辆，一共去了多少人？",
                    f"{big_cap}×2+{small_cap}×3={big_cap*2}+{small_cap*3}={total_people}（人）\n"
                    f"答：一共去了{total_people}人。"
                ))
        elif t == 5:  # 比例分配 - 构造法
            r1, r2, r3 = ri(2,5), ri(2,5), ri(2,5)
            total_parts = r1 + r2 + r3
            k = ri(20, 100)
            total = total_parts * k
            v1 = r1 * k
            v2 = r2 * k
            v3 = r3 * k
            res.append((
                f"把{total}千克化肥按{r1}:{r2}:{r3}分给三个村，各村分得多少千克？",
                f"总份数={r1}+{r2}+{r3}={r1+r2+r3}\n"
                f"第一个村：{total}×{r1}/{r1+r2+r3}={v1}（千克）\n"
                f"第二个村：{total}×{r2}/{r1+r2+r3}={v2}（千克）\n"
                f"第三个村：{total}×{r3}/{r1+r2+r3}={v3}（千克）\n"
                f"答：三个村分别分得{v1}、{v2}、{v3}千克。"
            ))
        else:  # 行程+时间综合 - 构造法
            v = ri(30, 80)
            t = ri(2, 6)
            d = v * t
            # 找d的因子作为v2
            v2 = None
            for candidate in range(20, 61):
                if d % candidate == 0 and d // candidate >= 2:
                    v2 = candidate
                    break
            if v2 is None:
                v2 = v  # 返回同速度
            t2 = d // v2
            res.append((
                f"{name}从A城到B城，去时每小时行{v}千米，用了{t}小时。返回时每小时行{v2}千米，返回用了几小时？",
                f"AB距离={v}×{t}={d}（千米）\n"
                f"返回时间={d}÷{v2}={t2}（小时）\n"
                f"答：返回用了{t2}小时。"
            ))
    return res


# ═══════════════════════════════════════════
# 手动精选灵活题（约100道，数值随机化）
# ═══════════════════════════════════════════
def gen_manual(n=100):
    """手动设计的灵活应用题，参数随机化"""
    res = []
    names = ["小明","小红","小刚","小华","小丽","小强","小芳","小军"]
    
    # 经典故事型
    a, b = random.sample(names, 2)
    x = ri(5, 25) * 2  # 确保偶数
    half_x = x // 2
    b_count = half_x + 3
    res.append((
        f"{a}有{x}颗糖，{b}的糖果数是{a}的一半还多3颗。{b}有多少颗糖？",
        f"{x}÷2+3={half_x}+3={b_count}（颗）\n答：{b}有{b_count}颗糖。"
    ))
    
    # 购物问题
    price_a = ri(5, 15)
    price_b = ri(3, 10)
    qty = ri(3, 8)
    money = ri(100, 200)
    total = price_a * qty + price_b * qty
    change = money - total
    res.append((
        f"妈妈带了{money}元去买水果，苹果每千克{price_a}元，香蕉每千克{price_b}元，各买了{qty}千克，应找回多少元？",
        f"苹果：{price_a}×{qty}={price_a*qty}（元）\n"
        f"香蕉：{price_b}×{qty}={price_b*qty}（元）\n"
        f"共花：{price_a*qty}+{price_b*qty}={total}（元）\n"
        f"找回：{money}-{total}={change}（元）\n"
        f"答：应找回{change}元。"
    ))
    
    # 工程合作变式
    a_days = ri(6, 12)
    b_days = ri(8, 16)
    work_days = ri(2, 5)
    a_rate = Fraction(1, a_days)
    b_rate = Fraction(1, b_days)
    done = (a_rate + b_rate) * work_days
    remain = Fraction(1) - done
    res.append((
        f"修一条路，甲队单独修{a_days}天完成，乙队单独修{b_days}天完成。两队合修{work_days}天后，剩下的由甲队单独修，还需几天？",
        f"合修效率=1/{a_days}+1/{b_days}={Fraction(a_days+b_days, a_days*b_days)}\n"
        f"合修{work_days}天完成={Fraction(work_days)}×{Fraction(a_days+b_days, a_days*b_days)}={done}\n"
        f"剩余=1-{done}={remain}\n"
        f"甲单独修需{remain}/{a_rate}={fmt_frac(remain/a_rate)}天\n"
        f"答：还需{fmt_frac(remain/a_rate)}天。"
    ))
    
    # 经典"倒推法"
    a_extra = ri(1, 3)
    b_extra = ri(1, 2)
    final = ri(5, 15)
    # 倒推：第二次拿前 = (final + b_extra) * 2
    before_2nd = (final + b_extra) * 2
    # 第一次拿前 = (before_2nd + a_extra) * 2
    original = (before_2nd + a_extra) * 2
    half1 = original // 2
    after_1st = half1 - a_extra
    half2 = after_1st // 2
    res.append((
        f"一筐鸡蛋，第一次拿出全部的一半多{a_extra}个，第二次拿出剩下的一半多{b_extra}个，还剩{final}个。原来有多少个？",
        f"倒推：第二次拿之前剩=（{final}+{b_extra}）×2={before_2nd}（个）\n"
        f"第一次拿之前（即原来）=（{before_2nd}+{a_extra}）×2={original}（个）\n"
        f"验证：全部{original}个，一半多{a_extra}={half1+a_extra}个，剩{original-half1-a_extra}={after_1st}个\n"
        f"再拿一半多{b_extra}={half2+b_extra}个，剩{after_1st-half2-b_extra}={final}个✓\n"
        f"答：原来有{original}个鸡蛋。"
    ))
    
    # 经典"替换法"
    big_box = ri(3, 6)
    small_box = ri(2, 5)
    while big_box == small_box:
        small_box = ri(2, 5)
    bx = ri(2, 6)
    sy = ri(2, 6)
    total_items = big_box * bx + small_box * sy
    total_boxes = bx + sy
    res.append((
        f"{total_items}个球装入{total_boxes}个盒子，大盒装{big_box}个，小盒装{small_box}个。大盒和小盒各几个？",
        f"设大盒x个，小盒({total_boxes}-x)个\n"
        f"{big_box}x+{small_box}×({total_boxes}-x)={total_items}\n"
        f"{big_box-small_box}x={total_items-small_box*total_boxes}\n"
        f"x={bx}\n"
        f"答：大盒{bx}个，小盒{sy}个。"
    ))
    
    # 流水行船 - 构造法
    boat_speed = ri(15, 30)
    water_speed = ri(2, 6)
    downstream = boat_speed + water_speed
    upstream = boat_speed - water_speed
    time_down = ri(2, 5)
    time_up = ri(3, 8)
    # distance必须是downstream和upstream的公倍数
    lcm = (downstream * upstream) // math.gcd(downstream, upstream)
    distance = lcm * ri(1, 3)
    time_down_actual = distance // downstream
    time_up_actual = distance // upstream
    res.append((
        f"一艘船在静水中的速度是每小时{boat_speed}千米，水流速度每小时{water_speed}千米。这艘船顺流航行{distance}千米需要几小时？逆流返回需要几小时？",
        f"顺流速度={boat_speed}+{water_speed}={downstream}（千米/时）\n"
        f"顺流时间={distance}÷{downstream}={time_down_actual}（小时）\n"
        f"逆流速度={boat_speed}-{water_speed}={upstream}（千米/时）\n"
        f"逆流时间={distance}÷{upstream}={time_up_actual}（小时）\n"
        f"答：顺流{time_down_actual}小时，逆流{time_up_actual}小时。"
    ))
    
    # 经典"盈亏"变式 - 住宿
    room3 = ri(3, 8)
    room2 = ri(2, 6)
    total_people = room3 * 3 + room2 * 2
    total_rooms = room3 + room2
    res.append((
        f"旅客{total_people}人住宿，3人间和2人间共{total_rooms}间，刚好住满。3人间和2人间各几间？",
        f"设3人间x间，2人间({total_rooms}-x)间\n"
        f"3x+2×({total_rooms}-x)={total_people}\n"
        f"3x+{total_rooms*2}-2x={total_people}\n"
        f"x={total_people-total_rooms*2}\n"
        f"3人间{total_people-total_rooms*2}间，2人间{total_rooms-(total_people-total_rooms*2)}间\n"
        f"答：3人间{room3}间，2人间{room2}间。"
    ))
    
    # 牛吃草问题（简化版）- 构造法
    grass_rate = ri(1, 3)
    cow_eat = ri(1, 3)
    n_cows = ri(5, 15)
    net_rate = n_cows * cow_eat - grass_rate  # 保证 > 0
    days = ri(5, 20)
    initial_grass = net_rate * days
    res.append((
        f"牧场原有{initial_grass}份草，每天新长{grass_rate}份。每头牛每天吃{cow_eat}份草。{n_cows}头牛几天能把草吃完？",
        f"每天净减少={n_cows}×{cow_eat}-{grass_rate}={net_rate}（份）\n"
        f"天数={initial_grass}÷{net_rate}={days}（天）\n"
        f"答：{days}天能把草吃完。"
    ))
    
    # 填充到n道
    while len(res) < n:
        t = ri(1, 8)
        name = rc(names)
        if t == 1:
            # 倍数+差
            a_val = ri(10, 50)
            k = ri(2, 5)
            b_val = a_val * k
            diff = b_val - a_val
            res.append((
                f"{name}看一本书，已看页数是未看页数的{k}倍，如果再看{diff//2}页，已看和未看就一样多。这本书共多少页？",
                f"已看={b_val}页，未看={a_val}页（差{diff}页）\n"
                f"共{a_val}+{b_val}={a_val+b_val}页\n"
                f"答：这本书共{a_val+b_val}页。"
            ))
        elif t == 2:
            # 方阵
            side = ri(5, 12)
            total = side * side
            res.append((
                f"同学们排成一个实心方阵，每边{side}人。一共有多少人？最外层有多少人？",
                f"总人数={side}×{side}={total}（人）\n"
                f"最外层=4×{side}-4={4*side-4}（人）\n"
                f"答：一共{total}人，最外层{4*side-4}人。"
            ))
        elif t == 3:
            # 锯木头
            cuts = ri(3, 8)
            time_per_cut = ri(2, 6)
            total_time = cuts * time_per_cut
            res.append((
                f"把一根木头锯成{cuts+1}段，每锯一次需要{time_per_cut}分钟，一共需要多少分钟？",
                f"锯{cuts+1}段需锯{cuts}次\n"
                f"总时间={cuts}×{time_per_cut}={total_time}（分钟）\n"
                f"答：一共需要{total_time}分钟。"
            ))
        elif t == 4:
            # 爬井 - 构造法
            up = ri(3, 5)
            down = ri(1, up - 1)
            net = up - down
            days = ri(3, 8)
            # depth = net * (days-1) + up
            depth = net * (days - 1) + up
            res.append((
                f"一只蜗牛从{depth}米深的井底往上爬，白天向上爬{up}米，晚上滑下{down}米。蜗牛几天能爬出井口？",
                f"每天净爬={up}-{down}={net}（米）\n"
                f"前{days-1}天爬了{net}×{days-1}={net*(days-1)}米\n"
                f"第{days}天白天爬{up}米到达{net*(days-1)+up}米={depth}米，到达井口\n"
                f"答：{days}天能爬出井口。"
            ))
        elif t == 5:
            # 时钟问题
            h = ri(1, 11)
            res.append((
                f"{h}点整时，时针和分针的夹角是多少度？",
                f"每小时对应30°\n"
                f"{h}点整时，分针指向12，时针指向{h}\n"
                f"夹角={h}×30°={h*30}°\n"
                f"答：夹角是{h*30}°。"
            ))
        elif t == 6:
            # 容斥原理
            a_count = ri(15, 30)
            b_count = ri(12, 25)
            both = ri(3, 10)
            neither = ri(2, 8)
            total = a_count + b_count - both + neither
            res.append((
                f"班上有{total}人，喜欢数学的有{a_count}人，喜欢语文的有{b_count}人，两科都喜欢的有{both}人。两科都不喜欢的有多少人？",
                f"至少喜欢一科={a_count}+{b_count}-{both}={a_count+b_count-both}（人）\n"
                f"两科都不喜欢={total}-（{a_count+b_count-both}）={neither}（人）\n"
                f"答：两科都不喜欢的有{neither}人。"
            ))
        elif t == 7:
            # 浓度混合
            s1 = ri(100, 300)
            c1 = ri(10, 30)
            s2 = ri(100, 300)
            c2 = ri(5, 15)
            salt = s1 * c1 // 100 + s2 * c2 // 100
            total_s = s1 + s2
            conc = round2(salt / total_s * 100)
            res.append((
                f"把{s1}克浓度为{c1}%的盐水和{s2}克浓度为{c2}%的盐水混合，混合后的浓度是多少？",
                f"盐1={s1}×{c1}%={s1*c1//100}克\n"
                f"盐2={s2}×{c2}%={s2*c2//100}克\n"
                f"总盐={s1*c1//100}+{s2*c2//100}={salt}克\n"
                f"总溶液={s1}+{s2}={total_s}克\n"
                f"浓度={salt}÷{total_s}×100%={conc}%\n"
                f"答：混合后浓度是{conc}%。"
            ))
        elif t == 8:
            # 等差数列
            first = ri(2, 10)
            diff = ri(2, 5)
            n_terms = ri(8, 15)
            last = first + (n_terms - 1) * diff
            total = (first + last) * n_terms // 2
            res.append((
                f"一列数：{first}、{first+diff}、{first+2*diff}、{first+3*diff}……第{n_terms}个数是多少？前{n_terms}个数的和是多少？",
                f"第{n_terms}个={first}+（{n_terms}-1）×{diff}={last}\n"
                f"和=（{first}+{last}）×{n_terms}÷2={total}\n"
                f"答：第{n_terms}个数是{last}，前{n_terms}个数的和是{total}。"
            ))
    
    return res[:n]


# ═══════════════════════════════════════════
# 汇总 & 生成文档
# ═══════════════════════════════════════════

CATEGORY_CONFIG = [
    ("行程问题", gen_xingcheng, 40),
    ("工程问题", gen_gongcheng, 30),
    ("浓度问题", gen_nongdu, 20),
    ("利润折扣问题", gen_lirun, 30),
    ("比例问题", gen_bili, 30),
    ("年龄问题", gen_nianling, 20),
    ("鸡兔同笼", gen_jitu, 20),
    ("植树问题", gen_zhishu, 15),
    ("盈亏问题", gen_yingkui, 15),
    ("和差问题", gen_hecha, 20),
    ("和倍差倍问题", gen_hebeicha, 20),
    ("归一问题", gen_guiyi, 20),
    ("分数百分数应用", gen_fenshu, 35),
    ("几何应用", gen_jihe, 30),
    ("平均数问题", gen_pingjun, 20),
    ("还原问题", gen_huanyuan, 15),
    ("统计问题", gen_tongji, 15),
    ("时间日期问题", gen_shijian, 15),
    ("利息税率问题", gen_lixishuilv, 15),
    ("综合应用题", gen_zonghe, 30),
    ("精选灵活题", gen_manual, 100),
]

def gen_all():
    all_problems = []
    for cat_name, gen_func, count in CATEGORY_CONFIG:
        try:
            problems = gen_func(count)
            for q, a in problems:
                all_problems.append((cat_name, q, a))
        except Exception as e:
            print(f"[WARN] {cat_name} 生成异常: {e}")
    return all_problems

def pad_to_500(problems):
    """确保恰好500道"""
    if len(problems) > 500:
        return problems[:500]
    while len(problems) < 500:
        idx = len(problems)
        cat = f"补充题{idx-499}"
        x = ri(10, 50)
        y = ri(10, 50)
        problems.append((cat,
            f"{x}+{y}=？",
            f"{x}+{y}={x+y}\n答：结果是{x+y}。"))
    return problems

def create_doc(title, problems, show_answer):
    doc = Document()
    # A4页面
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    # 标题
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(title)
    run.font.size = Pt(18)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.bold = True
    
    # 按分类输出
    current_cat = None
    idx = 0
    for cat, q, a in problems:
        if cat != current_cat:
            current_cat = cat
            p_cat = doc.add_paragraph()
            run_cat = p_cat.add_run(f"【{cat}】")
            run_cat.font.size = Pt(13)
            run_cat.font.name = "宋体"
            run_cat._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run_cat.bold = True
            idx = 0
        
        idx += 1
        p = doc.add_paragraph()
        if show_answer:
            content = f"{idx}. {q}\n【答案】{a}\n"
        else:
            content = f"{idx}. {q}\n\n"
        
        run = p.add_run(content)
        run.font.size = Pt(10.5)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    
    return doc


if __name__ == "__main__":
    print("正在生成500道应用题……")
    problems = gen_all()
    print(f"生成了 {len(problems)} 道，正在调整至500道……")
    problems = pad_to_500(problems)
    
    # 时间戳后缀
    ts = time.strftime("%Y%m%d_%H%M%S")
    # 统计各分类

    from collections import Counter
    cat_count = Counter(cat for cat, _, _ in problems)
    print("各分类题目数：")
    for cat, cnt in cat_count.most_common():
        print(f"  {cat}: {cnt}")
    print(f"总计: {len(problems)}")
    
    print("\n正在生成题目卷……")
    doc_q = create_doc("小学数学应用题500道（题目卷）", problems, show_answer=False)
    fname_q = f"小学数学应用题500道_题目卷_{ts}.docx"
    doc_q.save(fname_q)
    print(f"已保存: {fname_q}")
    
    print("正在生成答案卷……")
    doc_a = create_doc("小学数学应用题500道（答案卷）", problems, show_answer=True)
    fname_a = f"小学数学应用题500道_答案卷_{ts}.docx"
    doc_a.save(fname_a)
    print(f"已保存: {fname_a}")
    
    print("\n完成！")
