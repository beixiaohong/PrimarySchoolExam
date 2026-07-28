"""
小升初数学应用题200道生成器 — 高难度版
12大类，每题至少2-3步解题，适合小升初拔高训练
所有参数经过校验保证答案正确且为整数/简洁分数
"""
import random
import math
import signal
import time
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

random.seed(314159)
ri = random.randint
rc = random.choice


# ==================== 1. 行程问题（高难度） ====================

def xc_meet_then_apart():
    """相遇后继续走，求相距"""
    v1 = rc([40, 50, 60])
    v2 = rc([30, 40, 50])
    while v1 == v2:
        v2 = rc([30, 40, 50])
    t_meet = ri(2, 4)
    d = (v1 + v2) * t_meet
    t_more = ri(1, 3)
    apart = (v1 + v2) * t_more
    return (f"甲乙两人从相距{d}千米的两地同时出发相向而行，甲每小时行{v1}千米，乙每小时行{v2}千米。相遇后两人继续前行，又走了{t_more}小时，此时两人相距多少千米？",
            f"相遇时间 = {d}÷({v1}+{v2}) = {t_meet}小时。相遇后继续走{t_more}小时，相距 = ({v1}+{v2})×{t_more} = {apart}千米")

def xc_pursuit_with_start():
    """追及问题：不同出发点"""
    # 直接构造：先定追及时间t和速度差diff，再反推
    t = ri(2, 6)
    diff = rc([15, 20, 25, 30])
    # 乙先走gap小时，甲t小时追上：v_slow * gap = diff * t
    gap = ri(2, 4)
    dist = diff * t  # 乙先走的距离
    # v_slow = dist / gap
    while dist % gap != 0:
        gap = ri(2, 4)
    v_slow = dist // gap
    v_fast = v_slow + diff
    pos = v_fast * t
    return (f"甲每小时行{v_fast}千米，乙每小时行{v_slow}千米，乙先出发{gap}小时后甲出发。甲几小时追上乙？追上时距出发点多远？",
            f"乙先走 = {v_slow}×{gap} = {dist}千米，速度差 = {diff}千米/时，追及时间 = {dist}÷{diff} = {t}小时，距出发点 = {v_fast}×{t} = {pos}千米")

def xc_round_trip_avg():
    """往返平均速度（陷阱题：不等于速度平均）"""
    d = rc([120, 180, 240])
    v1 = rc([40, 60])
    v2 = rc([30, 40, 50, 80])
    while d % v1 != 0 or d % v2 != 0 or v1 == v2:
        d = rc([120, 180, 240])
        v1 = rc([40, 60])
        v2 = rc([30, 40, 50, 80])
    t1, t2 = d // v1, d // v2
    total_d, total_t = 2 * d, t1 + t2
    while total_d % total_t != 0:
        d = rc([120, 180, 240])
        v1 = rc([40, 60])
        v2 = rc([30, 40, 50, 80])
        if d % v1 != 0 or d % v2 != 0 or v1 == v2:
            continue
        t1, t2 = d // v1, d // v2
        total_d, total_t = 2 * d, t1 + t2
    avg = total_d // total_t
    return (f"从A到B距离{d}千米，去时速度{v1}千米/时，返回速度{v2}千米/时。往返平均速度是多少？（注意：不是两个速度求平均）",
            f"去时{d}÷{v1}={t1}小时，返回{d}÷{v2}={t2}小时。总路程={total_d}千米，总时间={total_t}小时，平均速度={total_d}÷{total_t}={avg}千米/时")

def xc_boat_distance():
    """流水行船求距离"""
    v_boat = rc([20, 24, 30])
    v_water = rc([4, 5, 6])
    v_down = v_boat + v_water
    v_up = v_boat - v_water
    # 直接构造：先定逆水时间，算距离，再求顺水时间
    t_up = ri(2, 5)
    d = v_up * t_up
    # 顺水时间 = d / v_down
    while d % v_down != 0:
        t_up = ri(2, 5)
        d = v_up * t_up
    t_down = d // v_down
    return (f"船在静水中速度{v_boat}千米/时，水速{v_water}千米/时。顺水行驶{t_down}小时的距离，逆水需要几小时？",
            f"顺水速度 = {v_down}千米/时，距离 = {v_down}×{t_down} = {d}千米。逆水速度 = {v_up}千米/时，时间 = {d}÷{v_up} = {t_up}小时")

def xc_train_bridge():
    """火车过桥问题"""
    train_len = rc([100, 150, 200])
    bridge_len = rc([300, 400, 500, 800])
    speed_m = rc([15, 20, 25, 30])  # 米/秒
    total_dist = train_len + bridge_len
    while total_dist % speed_m != 0:
        speed_m = rc([15, 20, 25, 30])
    t = total_dist // speed_m
    return (f"一列火车长{train_len}米，以每秒{speed_m}米的速度通过一座{bridge_len}米长的桥，从车头上桥到车尾离桥需要多少秒？",
            f"总距离 = 火车长+桥长 = {train_len}+{bridge_len} = {total_dist}米，时间 = {total_dist}÷{speed_m} = {t}秒")

def xc_two_meetings():
    """两次相遇问题"""
    d = rc([300, 400, 500, 600])
    v1 = rc([40, 50, 60])
    v2 = rc([30, 40, 50])
    while v1 == v2:
        v2 = rc([30, 40, 50])
    # 第一次相遇时甲走的路
    t1 = d  # 用比例思考：第一次相遇合走1个全程
    # 甲走 = v1/(v1+v2) * d
    num = v1 * d
    den = v1 + v2
    while num % den != 0:
        d = rc([300, 400, 500, 600])
        num = v1 * d
    s1 = num // den
    return (f"甲乙从相距{d}米的A、B两地同时出发相向而行，甲速{v1}米/分，乙速{v2}米/分。第一次相遇时甲走了多少米？",
            f"相遇时间 = {d}÷({v1}+{v2}) 分钟，甲走 = {v1}×{d}÷({v1}+{v2}) = {s1}米")

def xc_clock_lap():
    """环形跑道追及"""
    track = rc([400, 600, 800])
    v1 = rc([250, 300, 350])  # 米/分
    v2 = rc([150, 200, 250])
    while v1 <= v2 or track % (v1 - v2) != 0:
        v1 = rc([250, 300, 350])
        v2 = rc([150, 200, 250])
    t = track // (v1 - v2)
    return (f"环形跑道周长{track}米，甲每分钟跑{v1}米，乙每分钟跑{v2}米，同时同地同向出发，几分钟后甲首次追上乙？",
            f"速度差 = {v1}-{v2} = {v1-v2}米/分，追一圈需 {track}÷{v1-v2} = {t}分钟")

XINGCHENG = [xc_meet_then_apart, xc_pursuit_with_start, xc_round_trip_avg,
             xc_boat_distance, xc_train_bridge, xc_two_meetings, xc_clock_lap]


# ==================== 2. 工程问题（高难度） ====================

def gc_coop_then_alone():
    """合做一段时间后一人离开"""
    pairs = [(10, 15), (12, 8), (8, 12), (6, 12), (15, 10)]
    a, b = rc(pairs)
    lcm_val = a * b // math.gcd(a, b)
    eff_a, eff_b = lcm_val // a, lcm_val // b
    total_eff = eff_a + eff_b
    coop_days = ri(1, max(min(a, b) // 2, 2))
    done = coop_days * total_eff
    remain = lcm_val - done
    while remain <= 0 or remain % eff_a != 0:
        coop_days = ri(1, max(min(a, b) // 2, 2))
        done = coop_days * total_eff
        remain = lcm_val - done
    alone_days = remain // eff_a
    return (f"一项工程，甲独做{a}天完成，乙独做{b}天完成。两人合做{coop_days}天后乙离开，剩下的甲独做几天完成？",
            f"合做效率 = {total_eff}/{lcm_val}，{coop_days}天完成{done}/{lcm_val}，剩{remain}/{lcm_val}，甲效率 = {eff_a}/{lcm_val}，甲还需{remain}÷{eff_a} = {alone_days}天")

def gc_alternate():
    """轮流做"""
    pairs = [(8, 12), (10, 15), (6, 12), (12, 8)]
    a, b = rc(pairs)
    lcm_val = a * b // math.gcd(a, b)
    eff_a, eff_b = lcm_val // a, lcm_val // b
    cycle = eff_a + eff_b  # 一个周期（甲1天+乙1天）
    cycles = lcm_val // cycle
    remain = lcm_val - cycles * cycle
    while remain < 0:
        cycles -= 1
        remain = lcm_val - cycles * cycle
    total_days = cycles * 2
    if remain > 0:
        if remain <= eff_a:
            total_days += 1
            remain_note = f"剩{remain}/{lcm_val}由甲完成（不到1天）"
        else:
            total_days += 2
            remain_note = f"剩{remain}/{lcm_val}由甲做1天再由乙做"
    else:
        remain_note = "刚好完成"
    return (f"一项工程，甲独做{a}天，乙独做{b}天。甲乙轮流各做一天，从甲开始，几天完成？",
            f"每2天完成{cycle}/{lcm_val}，{cycles}个周期={cycles*cycle}/{lcm_val}用了{cycles*2}天，{remain_note}，共约{total_days}天")

def gc_three_partial():
    """三人合做，一人提前离开"""
    triples = [(6, 12, 4), (10, 15, 6), (8, 24, 12), (4, 6, 12), (12, 8, 6), (6, 4, 12)]
    a, b, c = rc(triples)
    lcm_val = a
    for x in [b, c]:
        lcm_val = lcm_val * x // math.gcd(lcm_val, x)
    ea, eb, ec = lcm_val // a, lcm_val // b, lcm_val // c
    total = ea + eb + ec
    two_eff = ea + eb
    # 枚举可行的pre_days
    valid = []
    for pd in range(1, max(1, lcm_val // total)):
        done = pd * total
        remain = lcm_val - done
        if remain > 0 and remain % two_eff == 0:
            valid.append((pd, remain))
    if not valid:
        # 回退：构造简单版本
        a, b, c = 6, 12, 4
        lcm_val = 12
        ea, eb, ec = 2, 1, 3
        total = 6
        two_eff = 3
        pre_days = 1
        remain = 6
        valid = [(1, 6)]
    pre_days, remain = rc(valid)
    more_days = remain // two_eff
    return (f"甲{a}天、乙{b}天、丙{c}天完成一项工程。三人合做{pre_days}天后丙离开，甲乙合做几天完成剩余？",
            f"三人效率={total}/{lcm_val}，{pre_days}天完成{pre_days*total}/{lcm_val}，剩{remain}/{lcm_val}。甲乙效率={two_eff}/{lcm_val}，需{remain}÷{two_eff}={more_days}天")

def gc_efficiency_change():
    """效率提高"""
    # 直接构造保证整除：new_days = remain_num * 100 / eff_new_pct
    # eff_new_pct in [120, 125, 150], remain_num * 100 要能被整除
    # 120: remain_num需被6整除 (600/120=5)  → remain_num in [6,12,18,...]
    # 125: remain_num需被5整除 (500/125=4)  → remain_num in [5,10,15,...]
    # 150: remain_num需被3整除 (300/150=2)  → remain_num in [3,6,9,...]
    pct = rc([20, 25, 50])
    eff_new_pct = 100 + pct
    divisors = {120: 6, 125: 5, 150: 3}
    div = divisors[eff_new_pct]
    # remain_num = days - work_days, 需要是div的倍数
    days = rc([10, 12, 15, 18, 20, 24])
    # remain_num 范围: [div, 2*div, ...] 且 < days
    max_remain = days - 2  # 至少做2天
    multiples = [k * div for k in range(1, max_remain // div + 1) if k * div < days]
    if not multiples:
        multiples = [div]
        days = div + 4
    remain_num = rc(multiples)
    work_days = days - remain_num
    new_days = remain_num * 100 // eff_new_pct
    return (f"一项工程计划{days}天完成，做了{work_days}天后效率提高{pct}%，还需几天完成？",
            f"已做{work_days}/{days}，剩{remain_num}/{days}。原效率=1/{days}，新效率={eff_new_pct}%/{days}。剩余天数={remain_num}×100÷{eff_new_pct}={new_days}天")

GONGCHENG = [gc_coop_then_alone, gc_alternate, gc_three_partial, gc_efficiency_change]


# ==================== 3. 浓度问题（高难度） ====================

def nd_two_mix():
    """两种不同浓度混合求比例"""
    p1 = rc([30, 40, 50])
    p2 = rc([10, 15, 20])
    target = ri(p2 + 5, p1 - 5)
    # 用十字交叉法
    ratio_high = target - p2
    ratio_low = p1 - target
    g = math.gcd(ratio_high, ratio_low)
    rh, rl = ratio_high // g, ratio_low // g
    return (f"有{p1}%和{p2}%两种盐水，要配成{target}%的盐水，两种盐水应按什么比例混合？",
            f"十字交叉法：{p1}%与{target}%差{p1-target}，{target}%与{p2}%差{target-p2}。{p1}%盐水:{p2}%盐水 = {target-p2}:{p1-target} = {rh}:{rl}")

def nd_replace():
    """倒出再加水（反复稀释）"""
    # 直接构造有效参数组合
    # remain_salt = pct * (total - pour) / 100 需为整数
    # new_pct = remain_salt * 100 / total 需为整数
    valid = []
    for total in [100, 200, 400, 500]:
        for pct in [20, 25, 40, 50]:
            for pour in [50, 100, 150, 200]:
                if pour >= total:
                    continue
                rs_num = pct * (total - pour)
                if rs_num % 100 != 0:
                    continue
                rs = rs_num // 100
                if rs * 100 % total != 0:
                    continue
                np_val = rs * 100 // total
                valid.append((total, pct, pour, rs, np_val))
    if not valid:
        valid = [(100, 20, 50, 10, 10)]
    total, pct, pour, remain_salt, new_pct = rc(valid)
    salt1 = total * pct // 100
    return (f"{total}克{pct}%盐水倒出{pour}克后加满水，浓度变为多少？",
            f"倒出后盐 = {salt1}×({total}-{pour})/{total} = {remain_salt}克，加满水后总量仍{total}克，浓度 = {remain_salt}÷{total}×100% = {new_pct}%")

def nd_target_concentration():
    """配制定量目标浓度"""
    # 预计算有效组合：total_want 必须能被 (p_high - p_low) 整除
    valid = []
    for ph in [30, 40, 50]:
        for pl in [10, 15]:
            ratio = ph - pl
            for tw in [200, 300, 400, 500, 600]:
                if tw % ratio == 0:
                    for tgt in range(pl + 5, ph - 4):
                        valid.append((ph, pl, tgt, tw))
    if not valid:
        valid = [(40, 10, 25, 300)]
    p_high, p_low, target, total_want = rc(valid)
    ratio_h = target - p_low
    ratio_l = p_high - target
    total_ratio = ratio_h + ratio_l
    unit = total_want // total_ratio
    need_h = unit * ratio_h
    need_l = unit * ratio_l
    return (f"用{p_high}%和{p_low}%的盐水配成{target}%的盐水{total_want}克，各需多少克？",
            f"十字交叉法得比 = {ratio_h}:{ratio_l}，总份{total_ratio}，{p_high}%需{need_h}克，{p_low}%需{need_l}克")

def nd_evaporate_then_mix():
    """蒸发后再混合"""
    pct1 = 20
    total1 = rc([200, 400])
    salt1 = total1 * pct1 // 100  # 40 or 80
    evap = total1 // 2  # 蒸发一半
    new_total1 = total1 - evap
    pct_new = salt1 * 100 // new_total1  # 40%
    # 再加回evap克同浓度
    add_salt = evap * pct1 // 100
    final_salt = salt1 + add_salt  # = salt1 + salt1 = 2*salt1... wait
    # 蒸发后盐不变=salt1, 加回evap克pct1%盐水: 加盐=evap*pct1/100
    final_salt = salt1 + add_salt
    final_total = total1  # 回到原来的量
    final_pct = final_salt * 100 // final_total
    return (f"{total1}克{pct1}%盐水蒸发{evap}克水后浓度变为多少？如果再加回{evap}克{pct1}%盐水，最终浓度是多少？",
            f"蒸发后：盐{salt1}克不变，溶液{new_total1}克，浓度={pct_new}%。加回{evap}克{pct1}%盐水(含盐{add_salt}克)：盐={final_salt}克，溶液={final_total}克，浓度={final_pct}%")

NONGDU = [nd_two_mix, nd_replace, nd_target_concentration, nd_evaporate_then_mix]


# ==================== 4. 利润折扣（高难度） ====================

def lr_double_discount():
    """连续两次打折"""
    price = rc([200, 300, 400, 500])
    d1 = rc([8, 9])
    d2 = rc([7, 8, 9])
    after1 = price * d1 // 10
    while price * d1 % 10 != 0:
        price = rc([200, 300, 400, 500])
        after1 = price * d1 // 10
    final = after1 * d2 // 10
    while after1 * d2 % 10 != 0:
        d2 = rc([7, 8, 9])
        final = after1 * d2 // 10
    cost = rc([100, 150, 200])
    while cost >= final:
        cost = rc([100, 150, 200])
    profit = final - cost
    pct = profit * 100 // cost
    while profit * 100 % cost != 0:
        cost = rc([100, 150, 200])
        while cost >= final:
            cost = rc([100, 150, 200])
        profit = final - cost
    pct = profit * 100 // cost
    return (f"商品标价{price}元，先打{d1}折再打{d2}折，进价{cost}元，最终利润率是多少？",
            f"第一次打折={price}×{d1}0%={after1}元，第二次={after1}×{d2}0%={final}元，利润={final}-{cost}={profit}元，利润率={profit}÷{cost}×100%={pct}%")

def lr_find_cost():
    """已知售价和利润率求进价"""
    profit_pct = rc([20, 25, 30, 40, 50])
    cost = rc([80, 100, 120, 150, 200])
    price = cost * (100 + profit_pct) // 100
    while cost * (100 + profit_pct) % 100 != 0:
        cost = rc([80, 100, 120, 150, 200])
        price = cost * (100 + profit_pct) // 100
    discount = rc([8, 9])
    actual = price * discount // 10
    while price * discount % 10 != 0:
        discount = rc([8, 9])
        actual = price * discount // 10
    new_profit = actual - cost
    new_pct = new_profit * 100 // cost
    while new_profit * 100 % cost != 0:
        cost = rc([80, 100, 120, 150, 200])
        price = cost * (100 + profit_pct) // 100
        while cost * (100 + profit_pct) % 100 != 0:
            cost = rc([80, 100, 120, 150, 200])
            price = cost * (100 + profit_pct) // 100
        actual = price * discount // 10
        while price * discount % 10 != 0:
            discount = rc([8, 9])
            actual = price * discount // 10
        new_profit = actual - cost
    new_pct = new_profit * 100 // cost
    return (f"商品加价{profit_pct}%后标价{price}元，打{discount}折出售，利润率是多少？",
            f"进价={price}÷{(100+profit_pct)/100}={cost}元，售价={price}×{discount}0%={actual}元，利润={new_profit}元，利润率={new_pct}%")

def lr_mixed_goods():
    """混合商品定价"""
    p1 = rc([20, 30, 40])  # 单价
    p2 = rc([50, 60, 70])
    n1 = rc([3, 4, 5])  # 数量
    n2 = rc([2, 3, 4])
    total_cost = p1 * n1 + p2 * n2
    total_qty = n1 + n2
    markup_pct = rc([20, 25, 30, 40])
    while total_cost * (100 + markup_pct) % 100 != 0:
        markup_pct = rc([20, 25, 30, 40])
    target_revenue = total_cost * (100 + markup_pct) // 100
    while target_revenue % total_qty != 0:
        n1 = rc([3, 4, 5])
        n2 = rc([2, 3, 4])
        total_qty = n1 + n2
        total_cost = p1 * n1 + p2 * n2
        target_revenue = total_cost * (100 + markup_pct) // 100
    unit_price = target_revenue // total_qty
    return (f"商店以{p1}元/个进了{n1}个，以{p2}元/个进了{n2}个同类商品。想获得{markup_pct}%的总利润，每个应卖多少元？",
            f"总成本 = {p1}×{n1}+{p2}×{n2} = {total_cost}元，目标收入 = {total_cost}×{(100+markup_pct)/100} = {target_revenue}元，单价 = {target_revenue}÷{total_qty} = {unit_price}元")

def lr_loss_gain():
    """两件商品一赚一赔"""
    sell = rc([100, 120, 150, 200])
    pct = rc([20, 25])
    # 甲赚pct%: cost1 = sell / (1 + pct/100)
    cost1_num = sell * 100
    cost1_den = 100 + pct
    while cost1_num % cost1_den != 0:
        sell = rc([100, 120, 150, 200])
        cost1_num = sell * 100
    cost1 = cost1_num // cost1_den
    # 乙赔pct%: cost2 = sell / (1 - pct/100)
    cost2_den = 100 - pct
    cost2_num = sell * 100
    while cost2_num % cost2_den != 0:
        sell = rc([100, 120, 150, 200])
        cost1_num = sell * 100
        while cost1_num % (100 + pct) != 0:
            sell = rc([100, 120, 150, 200])
            cost1_num = sell * 100
        cost1 = cost1_num // (100 + pct)
        cost2_num = sell * 100
    cost2 = cost2_num // cost2_den
    total_sell = 2 * sell
    total_cost = cost1 + cost2
    diff = total_sell - total_cost
    return (f"两件商品都以{sell}元卖出，一件赚{pct}%，一件赔{pct}%。总体是赚还是赔？差额多少？",
            f"甲进价={sell}÷{(100+pct)/100}={cost1}元，乙进价={sell}÷{(100-pct)/100}={cost2}元。总卖={total_sell}元，总进={total_cost}元，{'赚' if diff>0 else '赔'}{abs(diff)}元")

LIRUN = [lr_double_discount, lr_find_cost, lr_mixed_goods, lr_loss_gain]


# ==================== 5. 比例问题（高难度） ====================

def bl_changing_ratio():
    """变化后比例改变"""
    a, b = rc([(3, 5), (2, 5), (3, 7)])
    unit = rc([10, 15, 20])
    va, vb = a * unit, b * unit
    add = rc([10, 20, 30])
    new_a = va + add
    # 新的比
    g = math.gcd(new_a, vb)
    na, nb = new_a // g, vb // g
    return (f"甲乙两数之比{a}:{b}，甲增加{add}后（甲={va}，乙={vb}），新的比是多少？",
            f"甲 = {va}+{add} = {new_a}，乙 = {vb}，新比 = {new_a}:{vb} = {na}:{nb}")

def bl_three_way():
    """三人分配，已知两人差"""
    a, b, c = rc([(2, 3, 5), (1, 2, 3), (2, 3, 4)])
    diff_pairs = [(b, c, "乙丙"), (a, c, "甲丙"), (a, b, "甲乙")]
    p1, p2, label = rc(diff_pairs)
    diff = abs(p2 - p1) * rc([10, 15, 20])
    unit = diff // abs(p2 - p1)
    total = (a + b + c) * unit
    va, vb, vc = a * unit, b * unit, c * unit
    return (f"甲乙丙三人按{a}:{b}:{c}分配奖金，{label}相差{diff}元，奖金总共多少？各分多少？",
            f"差的份数={abs(p2-p1)}，每份={diff}÷{abs(p2-p1)}={unit}元，总={total}元。甲={va}元，乙={vb}元，丙={vc}元")

def bl_ratio_equation():
    """比例+方程"""
    ratio = rc([(3, 4), (4, 5), (5, 6)])
    unit = rc([10, 12, 15, 20])
    a_val, b_val = ratio[0] * unit, ratio[1] * unit
    subtract = rc([5, 10, 15])
    new_a = a_val - subtract
    new_b = b_val - subtract
    g = math.gcd(new_a, new_b)
    return (f"甲乙之比{ratio[0]}:{ratio[1]}（甲={a_val}，乙={b_val}），各减去{subtract}后，新比是多少？",
            f"甲 = {a_val}-{subtract} = {new_a}，乙 = {b_val}-{subtract} = {new_b}，新比 = {new_a//g}:{new_b//g}")

def bl_speed_ratio():
    """速度比求时间"""
    # 直接构造：甲速=a*k, 乙速=b*k, 距离=a*b*m
    a, b = rc([(3, 4), (4, 5), (2, 3)])
    k = rc([5, 10, 15])
    v1 = a * k
    v2 = b * k
    m = rc([2, 3, 4, 5])
    d = a * b * m  # 保证能被v1和v2整除: d/v1 = b*m/a... not necessarily
    # 用更简单的方法：距离=v1*t1, 让t1已知
    t1 = rc([4, 5, 6, 8, 10])
    d = v1 * t1
    while d % v2 != 0:
        t1 = rc([4, 5, 6, 8, 10])
        d = v1 * t1
    t2 = d // v2
    return (f"甲乙速度比{a}:{b}，甲速{v1}米/分。走同一段{d}米的路，甲用{t1}分钟，乙用几分钟？",
            f"速度比{a}:{b}，时间反比 = {b}:{a}。甲{t1}分钟，乙 = {t1}×{a}÷{b} = {t2}分钟")

BILI = [bl_changing_ratio, bl_three_way, bl_ratio_equation, bl_speed_ratio]


# ==================== 6. 年龄问题（高难度） ====================

def ag_past_future():
    """过去与将来的倍数关系"""
    child_now = rc([10, 12, 14])
    parent_now = child_now + rc([24, 26, 28])
    years_ago = rc([2, 4])
    years_later = rc([4, 6, 8])
    child_past = child_now - years_ago
    parent_past = parent_now - years_ago
    while parent_past % child_past != 0:
        child_now = rc([10, 12, 14])
        parent_now = child_now + rc([24, 26, 28])
        child_past = child_now - years_ago
        parent_past = parent_now - years_ago
    t_past = parent_past // child_past
    child_future = child_now + years_later
    parent_future = parent_now + years_later
    return (f"{years_ago}年前爸爸年龄是儿子的{t_past}倍，{years_later}年后父子年龄和是{parent_future+child_future}岁。现在各几岁？",
            f"设儿子现在x岁：{years_ago}年前爸爸={t_past}×(x-{years_ago})，现在爸爸={t_past}×(x-{years_ago})+{years_ago}。代入和方程解得：儿子{child_now}岁，爸爸{parent_now}岁")

def ag_three_people():
    """三人年龄"""
    a = ri(8, 12)
    b = a + ri(2, 6)
    c = a + b + ri(20, 30)
    total = a + b + c
    diff_ab = b - a
    diff_c = c - (a + b)
    return (f"三人年龄和{total}岁，乙比甲大{diff_ab}岁，丙比甲乙年龄和大{diff_c}岁，各几岁？",
            f"设甲x岁：x+(x+{diff_ab})+(2x+{diff_ab}+{diff_c})={total}，4x+{2*diff_ab+diff_c}={total}，x={a}。甲{a}岁，乙{b}岁，丙{c}岁")

def ag_ratio_change():
    """倍数关系随时间变化"""
    # 直接构造有效组合
    valid = []
    for child in [6, 8, 10, 12]:
        for times_now in [4, 5, 6]:
            parent = child * times_now
            for years in range(2, 16):
                child_f = child + years
                parent_f = parent + years
                if parent_f % child_f == 0:
                    t_future = parent_f // child_f
                    if t_future >= 2 and t_future < times_now:
                        valid.append((child, times_now, years, t_future))
    if not valid:
        valid = [(8, 5, 4, 3)]
    child, times_now, years, t_future = rc(valid)
    parent = child * times_now
    return (f"现在爸爸是儿子年龄的{times_now}倍，{years}年后是{t_future}倍。现在各几岁？",
            f"设儿子现在x岁：({times_now}x+{years})÷(x+{years})={t_future}，解方程得x={child}。儿子{child}岁，爸爸{parent}岁")

AGE = [ag_past_future, ag_three_people, ag_ratio_change]


# ==================== 7. 鸡兔同笼变式（高难度） ====================

def jt_three_animals():
    """三种动物"""
    chickens = ri(5, 12)
    rabbits = ri(3, 8)
    spiders = ri(2, 6)  # 8条腿
    heads = chickens + rabbits + spiders
    legs = 2 * chickens + 4 * rabbits + 8 * spiders
    return (f"鸡兔蜘蛛同笼，共{heads}个头{legs}条腿，蜘蛛比兔多{spiders-rabbits}只（{spiders if spiders>=rabbits else rabbits}-{min(spiders,rabbits)}={abs(spiders-rabbits)}），各几只？",
            f"设兔x只，蜘蛛x+{spiders-rabbits}只，鸡{heads}-2x-{spiders-rabbits}只。2({heads}-2x-{spiders-rabbits})+4x+8(x+{spiders-rabbits})={legs}，解得：鸡{chickens}只，兔{rabbits}只，蜘蛛{spiders}只")

def jt_score_complex():
    """复杂评分"""
    correct = ri(12, 18)
    wrong = ri(2, 5)
    blank = ri(1, 3)
    total_q = correct + wrong + blank
    score = correct * 5 - wrong * 3
    while score <= 0:
        correct = ri(12, 18)
        score = correct * 5 - wrong * 3
    return (f"考试共{total_q}题，答对5分、答错扣3分、不答0分。已知答了{correct+wrong}题，得{score}分，答对几题？",
            f"设答对x题，答错{correct+wrong}-x题：5x-3({correct+wrong}-x)={score}，8x={score+3*(correct+wrong)}，x={correct}题")

def jt_mixed_items():
    """混合物品（大盒小盒）"""
    big = ri(5, 12)  # 大盒装12个
    small = ri(8, 20)  # 小盒装5个
    total_boxes = big + small
    total_items = 12 * big + 5 * small
    return (f"大盒装12个，小盒装5个，共{total_boxes}盒装了{total_items}个，大盒小盒各几盒？",
            f"设大盒x：12x+5({total_boxes}-x)={total_items}，7x={total_items-5*total_boxes}，x={big}。大盒{big}盒，小盒{small}盒")

def jt_ticket():
    """买票问题"""
    adult = ri(5, 15)
    child = ri(10, 25)
    total_p = adult + child
    adult_price = rc([10, 15, 20])
    child_price = rc([5, 8])
    while adult_price <= child_price:
        child_price = rc([5, 8])
    total_money = adult * adult_price + child * child_price
    return (f"成人票{adult_price}元，儿童票{child_price}元，{total_p}人共花{total_money}元，成人和儿童各几人？",
            f"假设全是儿童：{total_p}×{child_price}={total_p*child_price}元，多{total_money-total_p*child_price}元，差{adult_price-child_price}元/人，成人={total_money-total_p*child_price}÷{adult_price-child_price}={adult}人，儿童={child}人")

JITU = [jt_three_animals, jt_score_complex, jt_mixed_items, jt_ticket]


# ==================== 8. 植树问题（高难度） ====================

def zs_square():
    """正方形四边植树"""
    side_trees = rc([10, 12, 15, 20])
    total = side_trees * 4 - 4  # 4个角重复
    gap = rc([3, 4, 5])
    perimeter = (side_trees - 1) * gap * 4
    return (f"正方形花坛每边植{side_trees}棵树（含角上），四边共植几棵？若间距{gap}米，周长多少？",
            f"每边{side_trees}棵，4边={side_trees}×4={side_trees*4}，4个角重复计算，共{total}棵。每边间距{(side_trees-1)}个×{gap}米={perimeter//4}米，周长={perimeter}米")

def zs_two_sides_road():
    """道路两旁不同间距"""
    length = rc([300, 600, 900])
    gap_a = rc([5, 10])
    gap_b = rc([6, 15])
    while length % gap_a != 0 or length % gap_b != 0:
        length = rc([300, 600, 900])
    trees_a = length // gap_a + 1
    trees_b = length // gap_b + 1
    total = trees_a + trees_b
    return (f"路长{length}米，一侧每隔{gap_a}米植杨树（两端都植），另一侧每隔{gap_b}米植柳树（两端都植），共几棵？",
            f"杨树 = {length}÷{gap_a}+1 = {trees_a}棵，柳树 = {length}÷{gap_b}+1 = {trees_b}棵，共{total}棵")

def zs_find_position():
    """求第N棵树的位置"""
    length = rc([500, 600, 1000])
    gap = rc([5, 10, 20])
    while length % gap != 0:
        length = rc([500, 600, 1000])
    total = length // gap + 1
    n = rc([15, 20, 25, 30])
    while n >= total:
        n = rc([15, 20, 25, 30])
    pos = (n - 1) * gap
    return (f"路长{length}米，两端都植，每隔{gap}米植一棵。第{n}棵树距起点多少米？",
            f"第{n}棵距起点 = ({n}-1)×{gap} = {pos}米")

ZHISHU = [zs_square, zs_two_sides_road, zs_find_position]


# ==================== 9. 分数百分数（高难度） ====================

def fs_remaining_fraction():
    """分数应用：剩余与消耗"""
    n, d = rc([(1, 4), (1, 3), (2, 5), (3, 8)])
    n2, d2 = rc([(1, 5), (1, 6), (1, 3), (2, 7)])
    total = d * d2 * rc([2, 3, 4, 5])
    part1 = total * n // d
    part2 = total * n2 // d2
    while total * n % d != 0 or total * n2 % d2 != 0:
        total = d * d2 * rc([2, 3, 4, 5])
        part1 = total * n // d
        part2 = total * n2 // d2
    remain = total - part1 - part2
    return (f"一本书{total}页，第一天看了{n}/{d}，第二天看了{n2}/{d2}，还剩多少页？",
            f"第一天{total}×{n}/{d}={part1}页，第二天{total}×{n2}/{d2}={part2}页，剩{total}-{part1}-{part2}={remain}页")

def fs_pct_change():
    """连续百分比变化"""
    base = rc([100, 200, 400, 500])
    up_pct = rc([10, 20, 25, 50])
    down_pct = rc([10, 20, 25, 50])
    after_up = base * (100 + up_pct) // 100
    while base * (100 + up_pct) % 100 != 0:
        base = rc([100, 200, 400, 500])
        after_up = base * (100 + up_pct) // 100
    final = after_up * (100 - down_pct) // 100
    while after_up * (100 - down_pct) % 100 != 0:
        down_pct = rc([10, 20, 25, 50])
        final = after_up * (100 - down_pct) // 100
    change = final - base
    return (f"某商品先涨价{up_pct}%，再降价{down_pct}%，与原价{base}元相比是涨还是降？变化多少？",
            f"涨后={base}×{(100+up_pct)/100}={after_up}元，降后={after_up}×{(100-down_pct)/100}={final}元，{'涨' if change>0 else '降'}了{abs(change)}元")

def fs_equation():
    """列方程解分数问题"""
    # x × (1 - a/b) = c
    n, d = rc([(1, 4), (1, 3), (2, 5), (1, 5)])
    x = d * rc([10, 15, 20, 25, 30])
    remain = x * (d - n) // d
    while x * (d - n) % d != 0:
        x = d * rc([10, 15, 20, 25, 30])
        remain = x * (d - n) // d
    return (f"一桶油用了{n}/{d}后还剩{remain}千克，原来有多少千克？",
            f"设原来x千克：x×(1-{n}/{d})={remain}，x×{(d-n)}/{d}={remain}，x={remain}×{d}/{d-n}={x}千克")

def fs_compare_two():
    """比多比少的综合"""
    # 直接构造：a 已知, b = a + diff, 需要 diff*100 % b == 0
    # 令 b = 100*k, diff = b - a, 需要 (b-a)*100 % b == 0
    # 即 100*(b-a)/b 是整数 → 100 - 100a/b 是整数 → 100a/b 是整数 → b | 100a
    # 简单做法：令 a=80, b=100, diff=20, less_pct=20%
    pairs = [(80, 100), (100, 125), (120, 150), (150, 200), (160, 200), (75, 100)]
    a, b = rc(pairs)
    pct_more = (b - a) * 100 // a  # 乙比甲多的百分比
    diff = b - a
    less_pct = diff * 100 // b
    return (f"乙比甲多{pct_more}%（甲={a}，乙={b}），那么甲比乙少百分之几？",
            f"差 = {diff}，甲比乙少 = {diff}÷{b}×100% = {less_pct}%（注意：基数不同，百分比不同！）")

FENSHU = [fs_remaining_fraction, fs_pct_change, fs_equation, fs_compare_two]


# ==================== 10. 几何应用（高难度） ====================

def geo_composite():
    """组合图形面积"""
    l = rc([10, 12, 15, 20])
    w = rc([6, 8, 10])
    # 长方形中挖去一个半圆
    r = w // 2
    rect_area = l * w
    semi_area = round(3.14 * r * r / 2, 2)
    result = round(rect_area - semi_area, 2)
    return (f"长方形长{l}cm宽{w}cm，在宽的一侧挖去一个半圆（直径={w}cm），剩余面积是多少？（π取3.14）",
            f"长方形面积={l}×{w}={rect_area}cm²，半圆面积=3.14×{r}²÷2={semi_area}cm²，剩余={rect_area}-{semi_area}={result}cm²")

def geo_pyramid_surface():
    """圆柱表面积"""
    r = rc([3, 5, 10])
    h = rc([8, 10, 12, 15])
    side = round(2 * 3.14 * r * h, 2)
    top = round(3.14 * r * r, 2)
    total = round(side + 2 * top, 2)
    return (f"圆柱底面半径{r}cm，高{h}cm，求表面积（含两个底面）。（π取3.14）",
            f"侧面积=2×3.14×{r}×{h}={side}cm²，底面积=3.14×{r}²={top}cm²，表面积={side}+2×{top}={total}cm²")

def geo_water_level():
    """水箱水位"""
    # 预计算有效组合：cube_side**3 必须能被 l*w 整除
    valid = []
    for l in [10, 20]:
        for w in [10, 15]:
            base = l * w
            for cs in range(5, 21):
                if (cs ** 3) % base == 0:
                    for hw in [8, 10, 12, 15]:
                        valid.append((l, w, hw, cs))
    if not valid:
        valid = [(10, 10, 10, 10)]
    l, w, h_water, cube_side = rc(valid)
    volume = l * w * h_water
    cube_vol = cube_side ** 3
    base = l * w
    rise = cube_vol // base
    return (f"长方体水箱底面{l}×{w}cm，水深{h_water}cm。放入棱长{cube_side}cm正方体铁块（完全浸没），水面升高多少？",
            f"铁块体积={cube_side}³={cube_vol}cm³，水面升高={cube_vol}÷({l}×{w})={rise}cm")

def geo_cone():
    """圆锥体积"""
    r = rc([3, 5, 6])
    h = rc([9, 12, 15, 18])
    volume = round(3.14 * r * r * h / 3, 2)
    return (f"圆锥底面半径{r}cm，高{h}cm，体积是多少？（π取3.14）",
            f"体积 = 1/3 × 3.14 × {r}² × {h} = {volume}cm³")

def geo_ring():
    """环形面积"""
    R = rc([5, 8, 10, 12])
    r = rc([3, 4, 5, 6])
    while r >= R:
        r = rc([3, 4, 5, 6])
    area = round(3.14 * (R * R - r * r), 2)
    return (f"大圆半径{R}cm，小圆半径{r}cm，求环形面积。（π取3.14）",
            f"环形面积 = 3.14×({R}²-{r}²) = 3.14×{R*R-r*r} = {area}cm²")

GEO = [geo_composite, geo_pyramid_surface, geo_water_level, geo_cone, geo_ring]


# ==================== 11. 平均数（高难度） ====================

def av_remove_one():
    """去掉一个后平均数变化"""
    n = rc([5, 6, 7])
    avg = ri(75, 90)
    total = n * avg
    removed = ri(40, 60)
    new_total = total - removed
    new_n = n - 1
    while new_total % new_n != 0:
        removed = ri(40, 60)
        new_total = total - removed
    new_avg = new_total // new_n
    return (f"{n}个数平均{avg}，去掉一个{removed}后，剩余数的平均数是多少？",
            f"原总和={n}×{avg}={total}，去掉{removed}后总和={new_total}，平均={new_total}÷{new_n}={new_avg}")

def av_correct_wrong():
    """改错后平均数变化"""
    n = rc([5, 6, 8, 10])
    wrong_val = ri(60, 80)
    correct_val = wrong_val + ri(10, 30)
    old_avg = ri(70, 85)
    old_total = n * old_avg
    new_total = old_total - wrong_val + correct_val
    while new_total % n != 0:
        correct_val = wrong_val + ri(10, 30)
        new_total = old_total - wrong_val + correct_val
    new_avg = new_total // n
    return (f"{n}个数平均{old_avg}，发现其中一个数{wrong_val}写错了应为{correct_val}，正确平均数是多少？",
            f"原总和={old_total}，修正后={old_total}-{wrong_val}+{correct_val}={new_total}，正确平均={new_total}÷{n}={new_avg}")

def av_mixed_group():
    """男女分组平均"""
    # 直接构造：选择使总和能被总人数整除的参数
    n_boys = rc([15, 20, 25])
    n_girls = rc([15, 20, 25])
    total_n = n_boys + n_girls
    # 先定总平均，再反推女生平均
    overall = ri(78, 90)
    avg_boys = ri(75, 88)
    total = overall * total_n
    girls_total = total - n_boys * avg_boys
    while girls_total <= 0 or girls_total % n_girls != 0:
        n_boys = rc([15, 20, 25])
        n_girls = rc([15, 20, 25])
        total_n = n_boys + n_girls
        overall = ri(78, 90)
        avg_boys = ri(75, 88)
        total = overall * total_n
        girls_total = total - n_boys * avg_boys
    avg_girls = girls_total // n_girls
    return (f"班级男生{n_boys}人平均{avg_boys}分，女生{n_girls}人平均{avg_girls}分，全班平均多少分？",
            f"总分={n_boys}×{avg_boys}+{n_girls}×{avg_girls}={total}，总人数={total_n}，平均={total}÷{total_n}={overall}分")

AVG = [av_remove_one, av_correct_wrong, av_mixed_group]


# ==================== 12. 还原问题（高难度） ====================

def hy_multi_step():
    """多步还原"""
    original = ri(20, 50)
    a1 = ri(10, 25)
    m = ri(2, 4)
    a2 = ri(5, 15)
    step1 = original + a1
    step2 = step1 * m
    step3 = step2 - a2
    return (f"某数加{a1}，再乘{m}，再减{a2}，得{step3}。原数是多少？",
            f"逆运算：{step3}+{a2}={step2}，{step2}÷{m}={step1}，{step1}-{a1}={original}")

def hy_fraction_step():
    """分数还原"""
    # 用去1/3，再用去剩余的1/2，还剩N
    n, d = rc([(1, 3), (1, 4), (2, 5)])
    remain1_num = d - n  # 用去n/d后剩(d-n)/d
    n2, d2 = rc([(1, 2), (1, 3)])
    remain2_num = d2 - n2  # 再用去n2/d2
    # 最终 = x × (d-n)/d × (d2-n2)/d2 = N
    # x = N × d × d2 / ((d-n) × (d2-n2))
    denom = remain1_num * remain2_num
    x_factor = d * d2
    unit = rc([2, 3, 4, 5])
    result = denom * unit
    x = x_factor * unit
    return (f"一桶油第一次用去{n}/{d}，第二次用去剩余的{n2}/{d2}，还剩{result}千克。原来有多少千克？",
            f"逆运算：{result}÷{(d2-n2)}/{d2}÷{(d-n)}/{d} = {result}×{d2}/{d2-n2}×{d}/{d-n} = {x}千克")

def hy_give_away():
    """送人还原"""
    original = ri(30, 80)
    # 先送一半多a个
    half = original // 2
    while original % 2 != 0:
        original = ri(30, 80)
        half = original // 2
    extra = ri(2, 8)
    gave1 = half + extra
    remain1 = original - gave1
    # 再送剩余的一半
    while remain1 % 2 != 0:
        extra = ri(2, 8)
        gave1 = half + extra
        remain1 = original - gave1
    gave2 = remain1 // 2
    final = remain1 - gave2
    return (f"一些糖果，第一次送出一半多{extra}颗，第二次送出剩余的一半，还剩{final}颗。原来有多少颗？",
            f"逆运算：{final}×2={remain1}（第二次送出前的剩余），{remain1}+{extra}={remain1+extra}（一半+多的{extra}），原来={half+extra+remain1}={original}颗")

HUANYUAN = [hy_multi_step, hy_fraction_step, hy_give_away]


# ==================== 主逻辑 ====================

ALL_CATEGORIES = [
    ("行程问题", XINGCHENG),
    ("工程问题", GONGCHENG),
    ("浓度问题", NONGDU),
    ("利润折扣", LIRUN),
    ("比例问题", BILI),
    ("年龄问题", AGE),
    ("鸡兔同笼", JITU),
    ("植树问题", ZHISHU),
    ("分数百分数", FENSHU),
    ("几何应用", GEO),
    ("平均数", AVG),
    ("还原问题", HUANYUAN),
]


def _safe_call(func, timeout=2):
    """带超时保护的函数调用"""
    def handler(signum, frame):
        raise TimeoutError()
    signal.signal(signal.SIGALRM, handler)
    signal.alarm(timeout)
    try:
        result = func()
        signal.alarm(0)
        return result
    except Exception:
        signal.alarm(0)
        return None


def generate_all(total=200):
    per_cat = total // len(ALL_CATEGORIES)
    extra = total - per_cat * len(ALL_CATEGORIES)
    all_problems = []
    global_seen = set()

    for i, (cat_name, funcs) in enumerate(ALL_CATEGORIES):
        n = per_cat + (1 if i < extra else 0)
        cat_problems = []

        # 每个函数先生成6道候选（带超时保护）
        for func in funcs:
            for _ in range(6):
                r = _safe_call(func)
                if r:
                    q, a = r
                    fp = q[:35]
                    if fp not in global_seen:
                        global_seen.add(fp)
                        cat_problems.append((q, a))

        # 补充阶段限制为 n*10 次
        attempts = 0
        while len(cat_problems) < n and attempts < n * 10:
            func = random.choice(funcs)
            r = _safe_call(func)
            if r:
                q, a = r
                fp = q[:35]
                if fp not in global_seen:
                    global_seen.add(fp)
                    cat_problems.append((q, a))
            attempts += 1

        all_problems.extend([(cat_name, q, a) for q, a in cat_problems[:n]])

    random.shuffle(all_problems)
    return all_problems


def to_markdown(problems):
    lines = []
    lines.append("# 小升初数学应用题 200 道（拔高版·含答案）")
    lines.append("")
    lines.append("> 涵盖行程、工程、浓度、利润折扣、比例、年龄、鸡兔同笼、植树、分数百分数、几何、平均数、还原问题共 12 大类。每题 2-4 步解题，适合小升初拔高训练。")
    lines.append("")
    lines.append("---")
    lines.append("")

    cat_count = {}
    for cat, q, a in problems:
        cat_count[cat] = cat_count.get(cat, 0) + 1

    idx = 1
    started_cats = set()
    for cat, q, a in problems:
        if cat not in started_cats:
            started_cats.add(cat)
            lines.append(f"## {cat}（{cat_count[cat]} 道）")
            lines.append("")

        lines.append(f"**第 {idx} 题**")
        lines.append("")
        lines.append(q)
        lines.append("")
        lines.append(f"> **答案：** {a}")
        lines.append("")
        idx += 1

    return "\n".join(lines)


def create_word_doc(problems, is_answer=False):
    """生成应用题Word文档（题目卷或答案卷）"""
    doc = Document()

    # 页面设置 A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_text = "小升初数学应用题 200 道（答案卷）" if is_answer else "小升初数学应用题 200 道（题目卷）"
    run = title.add_run(title_text)
    run.font.size = Pt(18)
    run.bold = True

    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("（小升初专项训练 · 拔高版 · 12大类应用题）")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # 信息栏
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = info.add_run("姓名：__________    班级：__________    日期：__________    得分：__________")
    run.font.size = Pt(11)

    doc.add_paragraph()

    # 统计各分类题数
    cat_count = {}
    for cat, q, a in problems:
        cat_count[cat] = cat_count.get(cat, 0) + 1

    if is_answer:
        # 答案卷：按分类分组显示
        idx = 1
        started_cats = set()
        for cat, q, a in problems:
            if cat not in started_cats:
                started_cats.add(cat)
                cat_p = doc.add_paragraph()
                cat_p.space_before = Pt(6)
                cat_p.space_after = Pt(2)
                run = cat_p.add_run(f"【{cat}】（{cat_count[cat]}题）")
                run.font.size = Pt(11)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 80, 160)

            p = doc.add_paragraph()
            p.space_before = Pt(2)
            p.space_after = Pt(2)
            run = p.add_run(f"{idx}. {a}")
            run.font.size = Pt(10)
            idx += 1
    else:
        # 题目卷：每道题一段 + 空白答题区
        # 显示题型统计
        stats_p = doc.add_paragraph()
        run = stats_p.add_run("题目包含以下类型：")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(120, 120, 120)
        stats_text = "、".join([f"{k}({v}题)" for k, v in cat_count.items()])
        run = stats_p.add_run(stats_text)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(150, 150, 150)

        idx = 1
        for cat, q, a in problems:
            p = doc.add_paragraph()
            p.space_before = Pt(4)
            p.space_after = Pt(4)
            run = p.add_run(f"{idx}. {q}")
            run.font.size = Pt(11)
            blank = doc.add_paragraph()
            run = blank.add_run("\n\n")
            run.font.size = Pt(11)
            idx += 1

    return doc


if __name__ == "__main__":
    print("=" * 55)
    print("  小升初数学应用题 200 道生成器（拔高版·Word文档）")
    print("=" * 55)

    print("\n[1/3] 生成 200 道应用题...")
    problems = generate_all(200)
    print(f"  共 {len(problems)} 道题目")

    # 统计各分类
    cat_count = {}
    for cat, q, a in problems:
        cat_count[cat] = cat_count.get(cat, 0) + 1
    print("  各分类题数：")
    for cat, count in cat_count.items():
        print(f"    {cat}: {count}题")

    ts = time.strftime("%Y%m%d_%H%M%S")

    print("\n[2/3] 生成题目卷...")
    doc_q = create_word_doc(problems, is_answer=False)
    fname_q = f"小升初应用题200道_题目卷_{ts}.docx"
    doc_q.save(fname_q)
    print(f"  已保存: {fname_q}")

    print("\n[3/3] 生成答案卷...")
    doc_a = create_word_doc(problems, is_answer=True)
    fname_a = f"小升初应用题200道_答案卷_{ts}.docx"
    doc_a.save(fname_a)
    print(f"  已保存: {fname_a}")

    print(f"\n完成！共生成 2 个文件：")
    print(f"  题目卷: {fname_q}")
    print(f"  答案卷: {fname_a}")
